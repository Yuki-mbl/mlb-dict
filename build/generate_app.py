#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate MLB Baseball Dictionary HTML App"""
import sys, json, re, base64, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from lxml import etree as _ET
import pykakasi

_SOFT_BR = '\x1f'  # ソフト改行の代替マーカー（テキスト中に出現しない文字）

def _para_text(p):
    """p.text の代替: ソフト改行 (w:br) を _SOFT_BR に変換して全テキストを返す"""
    parts = []
    for el in p._element.iter():
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 't':
            parts.append(el.text or '')
        elif tag == 'br':
            parts.append(_SOFT_BR)
    return ''.join(parts)
_kks = pykakasi.kakasi()
def to_yomi(text):
    """漢字を含む文字列の読みをカタカナで返す"""
    result = _kks.convert(text)
    return ''.join(item['kana'] for item in result)

# ── Filtering ─────────────────────────────────────────────────────────────────

STADIUM_KW = ['Stadium', 'Field', 'Park', 'Ballpark', 'Arena', 'Dome',
               'Coliseum', 'Centre', 'Complex', 'Grounds', 'Bowl']
TEAM_NAMES_EN = [
    'Diamondbacks', 'Braves', 'Orioles', 'Red Sox', 'Cubs', 'White Sox',
    'Reds', 'Indians', 'Guardians', 'Rockies', 'Tigers', 'Marlins', 'Astros',
    'Royals', 'Angels', 'Dodgers', 'Brewers', 'Twins', 'Expos', 'Mets',
    'Yankees', 'Athletics', 'Phillies', 'Pirates', 'Padres', 'Giants',
    'Mariners', 'Cardinals', 'Rays', 'Rangers', 'Blue Jays', 'Nationals',
    "D'backs", 'Snakes',
]
TEAM_ABBREVS = {
    'ARI','ATL','BAL','BOS','CHC','CHW','CIN','CLE','COL',
    'DET','FLA','HOU','KCR','LAA','LAD','MIL','MIN','MON',
    'NYM','NYY','OAK','PHI','PIT','SDP','SFG','SEA','STL',
    'TBD','TEX','TOR','WSN','ANA','ALT',
}
TEAM_NAMES_JA = [
    'レッドソックス','ヤンキース','ドジャース','カブス','ジャイアンツ',
    'アスレチックス','エンジェルス','エンゼルス','パドレス','ブリュワーズ',
    'ブレーブス','オリオールズ','タイガース','インディアンス','マリナーズ',
    'レンジャーズ','アストロズ','ロイヤルズ','ホワイトソックス','ツインズ',
    'ブルージェイズ','レイズ','フィリーズ','メッツ','カージナルス',
    'ロッキーズ','マーリンズ','エクスポズ','パイレーツ','レッズ',
    'ダイヤモンドバックス','ダイアモンドバックス','ガーディアンズ',
]
JA_VENUE_KW = ['スタジアム','パーク','フィールド','ドーム','アリーナ','ボールパーク','球場']

def should_filter(term):
    t = str(term).strip()
    for kw in STADIUM_KW:
        if kw in t: return True
    for nm in TEAM_NAMES_EN:
        if nm in t: return True
    first = t.split()[0].rstrip('=').strip() if t.split() else ''
    if first in TEAM_ABBREVS: return True
    for nm in TEAM_NAMES_JA:
        if nm in t: return True
    for kw in JA_VENUE_KW:
        if t.endswith(kw) and len(t) > len(kw) + 2: return True
        if kw == '球場' and kw in t: return True
    return False

_FULLWIDTH_DIGITS = str.maketrans(
    '０１２３４５６７８９'
    'ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺ'
    'ａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ'
    '＝',
    '0123456789'
    'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
    'abcdefghijklmnopqrstuvwxyz'
    '='
)

def _is_cjk(c):
    if not c: return False
    o = ord(c)
    return (0x3000 <= o <= 0x9FFF) or (0xFF00 <= o <= 0xFFEF)

def smart_join(a, b):
    """Join two text fragments: no space at Japanese/CJK boundaries, hyphen boundaries."""
    if not a: return b
    if not b: return a
    a = a.rstrip()
    last, first = (a[-1] if a else ''), b[0]
    if _is_cjk(last) or _is_cjk(first) or last == '-':
        return a + b
    return a + ' ' + b

# Dashes that appear as corrupted long-vowel marks between kana chars
_KANA_RANGE = re.compile(r'[ぁ-んァ-ヶ]')
_BAD_DASH = re.compile(r'[–—―－ｰ]')

def _fix_kana_dash(text):
    """Replace non-standard dashes with ー when preceded by a kana character."""
    result = []
    for i, ch in enumerate(text):
        if _BAD_DASH.match(ch) and i > 0 and _KANA_RANGE.match(text[i-1]):
            result.append('ー')
        else:
            result.append(ch)
    return ''.join(result)

# Sort-key overrides for entries the original dictionary filed by their English meaning
# 2006年発行時から球団名・球場名が変わったため除外
EJ_EXCLUDE = {
    # 削除対象エントリ
    'handy man',
    # 「Take Me Out To The Ball Game」歌詞の断片（誤パース防止）
    'at the old ball game',
    'take me out with the crowd',
    'buy me some peanuts and cracker jack',
    "i don't care if i never get back",
    'let me root, root, root for the home team',
    "if they don't win, it's a shame",
    "for it's one, two, three strikes, you're out,",
    # 選手紹介文の断片
    'tops in the majors.',
    'lef-handed hitters even though he\'s righty.',
    # cf./e.g. 参照が誤って見出しになったもの
    'cf.',
    'cf. make a great/spectacular play',
    'cf. sidearmer',
    'e.g.',
    # 球場名（旧称含む）
    'ameriquest field in arlington',
    'angel stadium of anaheim',
    'bank one ballpark',
    'busch(memorial)stadium',
    'citizens bank park',
    'comerica park',
    'coors field',
    'dodger stadium',
    'fenway park',
    'fenway park single',
    'great american ballpark',
    'jacobs field',
    'kauffman stadium',
    'miller park',
    'minute maid park',
    'network associates coliseum (the)',
    'oriole park at camden yards',
    'pnc park',
    'pacbell park',
    'pacific bell park',
    'petco park',
    'pro player stadium',
    'rfk stadium',
    'riverfront stadium',
    'safeco field',
    'shea stadium',
    'three-rivers stadium',
    'tropicana field',
    'turner field',
    'u.s. cellular field',
    'wrigley field',
    'yankee stadium',
}

JE_EXCLUDE = {
    'アナハイム・エンジェルス',
    'タンパベイ・デビル・レイズ',
    'フロリダ・マーリンズ',
    'アメリクエスト・フィールド・イン・アーリントン',
    'アールエフケイ・スタジアム',
    'セーフコ・フィールド',
    'ネットワーク・アソシエイツ・コロシアム',
    'ネットワーク・アソシーツ・コロシアム',
    'バンクワン・ボールパーク',
    'ターナー・フィールド',
    'プロプレヤー・スタジアム',
    'ブッシュ・メモリアル・スタジアム',
}

EJ_SORT_OVERRIDES = {
    '-plus innings': 'plus innings',
    '1B, 1b':        'one base',
    '2B, 2b':        'two base',
    '3B':            'three base',
}

def clean(text):
    if not text: return ''
    text = str(text)
    text = text.replace('大リーグ', 'メジャーリーグ')
    text = _fix_kana_dash(text)
    text = text.translate(_FULLWIDTH_DIGITS)
    # 英語用語中の全角括弧・全角スペースを半角に（日本語定義内は後続処理で保持）
    text = text.replace('（', '(').replace('）', ')').replace('　', ' ')
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = re.sub(r'　{2,}', '　', text)
    # Wordのソフト改行等で文中に混入した\r\n/\rを正規化してから除去
    text = text.replace('\r\n', '\n').replace('\r', '')
    text = re.sub(r'(?<![。\n])\n(?!\n)', '', text)
    return text.strip('　 \t\n\r')

# ── Parse EJ ─────────────────────────────────────────────────────────────────

def is_valid_ej_term(term):
    """English headwords must be primarily ASCII — filter out Japanese fragments."""
    t = term.strip()
    if not t:
        return False
    # Must start with letter, hyphen, or quote
    if not re.match(r'^[-A-Za-z#(\'"]', t):
        # Allow digit-starting abbreviations like 1B, 2B, 3B, 40-man
        if not re.match(r'^\d+[A-Za-z\-]', t):
            return False
    # Term must not contain Japanese (kana/kanji) — those are definition fragments
    # ただし全角括弧（）はアルファベット用語に使われるため許可
    if re.search(r'[ぁ-ん゛゜ァ-ヶー一-龯]', t):
        return False
    if re.search(r'[！-＇＊-｠]', t.replace('～', '')):  # 全角括弧・～以外の全角記号は除外
        return False
    # Excessively long "terms" are sentence fragments
    if len(t) > 80:
        return False
    return True

def parse_ej(filepath):
    doc = Document(filepath)
    lines = [_para_text(p) for p in doc.paragraphs]
    si = 0
    for i, ln in enumerate(lines):
        if re.match(r'^[A-Za-z0-9]', ln) and re.search(r'[぀-鿿゠-ヿ一-鿿]', ln):
            si = i; break
    entries = []
    term, defn = None, None
    prev_term, prev_defn = None, None   # 番号付き継続のために前エントリーを保持
    cont_to_prev = False                # 継続行を前エントリーへ振り向けるフラグ
    for ln in lines[si:]:
        # 全角スペース(U+3000)を半角2スペースに正規化（区切り文字として使われる場合の対応）
        ln = re.sub(r'(?<=[A-Za-z0-9\)\]～])　', '  ', ln)
        s = ln.lstrip().replace(_SOFT_BR, ' ')
        if not s.strip(): continue
        new_term = new_defn = None

        # ── ソフト改行処理 ────────────────────────────────────────────────────
        # 段落内ソフト改行 (w:br) が _SOFT_BR で記録されている
        # 第1セグメントが有効な見出し語 → 見出し＋定義として処理
        # それ以外 → 継続テキストとして既存定義に結合
        if _SOFT_BR in ln:
            segs = [seg.strip() for seg in ln.split(_SOFT_BR) if seg.strip()]
            if segs:
                head_seg = segs[0]
                rest_seg = ' '.join(segs[1:])
                ct_head = clean(head_seg)
                if (is_valid_ej_term(ct_head)
                        and not re.search(r'[ぁ-んァ-ヶー一-龯]', head_seg)
                        and not re.search(r'\s{2,}', head_seg)):
                    # ソフト改行前が見出し語
                    new_term = head_seg.strip()
                    new_defn = rest_seg
                else:
                    # ソフト改行が定義内部にある → スペースで結合して継続
                    combined = ' '.join(segs)
                    if defn is not None:
                        defn = smart_join(defn, combined)
                    elif cont_to_prev and entries:
                        entries[-1]['ja'] = smart_join(entries[-1]['ja'], combined)
            # ソフト改行行の処理完了 → new_term/new_defn の保存ロジックへ
            if new_term is None and new_defn is None:
                continue
        # ── 通常行処理 ─────────────────────────────────────────────────────────
        elif not ln[0:1].isspace():
            m = re.match(r'^(\S.*?)\s{2,}(.+)', ln)
            if m:
                pt, pd = m.group(1).strip(), m.group(2).strip()
                ct = clean(pt)
                if re.search(r'[ぁ-ん゛゜ァ-ヶー一-龯]', ct) or re.search(r'[！-＇＊-｠]', ct.replace('～', '')):
                    # Japanese leaked into term → single-space separator; resplit at first JP char
                    m2 = re.match(
                        r'^([A-Za-z0-9][\w\s\-\/\(\)\.\'\!\,\;\:\&\#～=]*?)\s+'
                        r'(?=[ぁ-んァ-ヶー一-龯１-９（])', ln)
                    if m2:
                        new_term = m2.group(1).strip()
                        nd = ln[m2.end():].strip()
                        # 番号付き継続（２．など）は前エントリーへ振り向け、後続継続行もそちらへ
                        if re.match(r'^[2-9２-９][．.]', nd):
                            if defn is not None: defn = smart_join(defn, nd)
                            new_defn = ''
                            cont_to_prev = True  # 続く継続行も前エントリー(entries[-1])へ
                        else:
                            cont_to_prev = False
                            new_defn = nd
                            # "TERM =equiv" アーティファクト修正
                            _eq = re.match(r'^(.*?)\s+=(\S.*)$', new_term)
                            if _eq:
                                new_term = _eq.group(1).strip()
                                new_defn = '=' + _eq.group(2) + ' ' + new_defn
                            elif new_term.rstrip().endswith('='):
                                new_term = new_term.rstrip('= ').strip()
                    # else: fall through to continuation
                elif re.match(r'^[2-9２-９][．.]', pd):
                    # 定義が２．以降で始まる → 前エントリーの複数義続き（termにJPなし）
                    pass  # treat as continuation below
                else:
                    # "TERM = 1. ..." 番号付き定義が見出しに混入 → 見出しだけ抽出
                    _ndf_m = re.match(r'^(.*?)\s*=\s*[0-9][.．]', pt)
                    if _ndf_m:
                        _real_term = _ndf_m.group(1).strip()
                        _remainder = pt[_ndf_m.end():].strip()  # "= stolen base" 等
                        new_term = _real_term
                        new_defn = (_remainder + ' ' if _remainder else '') + m.group(2)
                        cont_to_prev = False
                    else:
                        # 文章断片の除外（7語超+文末ピリオド、カンマ2個以上、冠詞・前置詞始まり）
                        _pt_bad = (
                            (pt.endswith('.') and len(pt.split()) > 6) or
                            pt.rstrip().endswith('=') or
                            re.search(r',.*,', pt) or
                            re.match(r'^(cf\.|e\.g\.)', pt.lower()) or
                            re.search(r'\b(even though|although|because|unless|whereas|despite)\b', pt.lower()) or
                            (re.match(r'^(a |an |the |with |in |on |of |for |and |but |by |from |to )', pt.lower())
                             and (len(pt.split()) >= 4 or pt[0].islower()))
                        )
                        if not _pt_bad:
                            new_term, new_defn = pt, m.group(2)
                            cont_to_prev = False
            elif re.search(r'[ぁ-んァ-ヶー一-龯]', ln):
                # 1スペース区切りエントリーのフォールバック（厳格なヒューリスティック）
                m2 = re.match(
                    r'^([A-Za-z][A-Za-z0-9 \-\/\(\)\.\=\'\,\&\#～/]*?)\s+'
                    r'(?=[ぁ-んァ-ヶー一-龯１-９（])', ln)
                if m2:
                    t = m2.group(1).strip()
                    words = t.split()
                    # 断片行の除外条件
                    _bad = (
                        len(words) >= 7 or          # 長すぎる（文の断片）
                        t.endswith('.') or           # 文末ピリオド
                        re.search(r'\d$', t) or      # 数字で終わる（断片）
                        re.search(r',.*,', t) or     # カンマ2個以上（リスト断片）
                        re.match(r'^(cf\.|e\.g\.)', t.lower()) or
                        re.match(r'^(a |an |the |with |in |on |of |for |and |but |by |from |to )', t.lower())
                    )
                    if not _bad:
                        new_term = t
                        new_defn = ln[m2.end():].strip()
                        # "TERM =equiv" → equiv は定義の一部
                        _eq = re.match(r'^(.*?)\s+=(\S.*)$', new_term)
                        if _eq:
                            new_term = _eq.group(1).strip()
                            new_defn = '=' + _eq.group(2) + ' ' + new_defn
                        elif new_term.rstrip().endswith('='):
                            # "TERM =" (trailing =) → strip =
                            new_term = new_term.rstrip('= ').strip()
                        cont_to_prev = False
            # ── スペースなしEN[JP]境界: "TERM = EQUIVjapanese" ───────────────
            # 例: smoke = smoker豪速球、BOS = Boston Redsoxの略記号
            # m2(1スペース)が不成立かつ行に = がある場合のみ試みる
            if new_term is None and re.search(r'[=＝]', ln):
                m3 = re.match(
                    r'^([A-Za-z][A-Za-z0-9 \-\/\(\)\.\'\,\&\#～/（）]*?)'
                    r'\s*[=＝]\s*'
                    r'([A-Za-z][A-Za-z0-9 \-\/\(\)\.\'\,\&\#～/]*?)'
                    r'(?=[ぁ-んァ-ヶー一-龯])',
                    ln)
                if m3:
                    _t3 = m3.group(1).strip()
                    _eq3 = m3.group(2).strip()
                    _full3 = _t3 + ' = ' + _eq3
                    _bad3 = (
                        len(_full3.split()) >= 12 or
                        _t3.endswith('.') or
                        re.search(r',.*,', _t3)
                    )
                    if not _bad3:
                        new_term = _full3  # dedup loop が " = " で分割
                        new_defn = ln[m3.end():].strip()
                        cont_to_prev = False
        if new_term is not None:
            if term is not None and defn is not None:
                d = clean(defn)
                ct = clean(term)
                if d and is_valid_ej_term(ct):
                    entry = {'en': ct, 'ja': d}
                    if ct in EJ_SORT_OVERRIDES:
                        entry['sk'] = EJ_SORT_OVERRIDES[ct]
                    entries.append(entry)
            prev_term, prev_defn = term, defn
            term, defn = new_term, new_defn
        else:
            if cont_to_prev and entries:
                # 番号付き継続の後続行 → 直前に保存されたエントリーへ追記
                entries[-1]['ja'] = smart_join(entries[-1]['ja'], s)
            else:
                if defn is not None: defn = smart_join(defn, s)
    if term and defn:
        d = clean(defn)
        ct = clean(term).replace('（','(').replace('）',')').replace('　',' ')
        if d and is_valid_ej_term(ct):
            entry = {'en': ct, 'ja': d}
            if ct in EJ_SORT_OVERRIDES:
                entry['sk'] = EJ_SORT_OVERRIDES[ct]
            entries.append(entry)
    return entries

# ── Parse JE ─────────────────────────────────────────────────────────────────

def is_valid_je_term(term, en_eq):
    if re.search(r'[。、！？]', term): return False
    if any(x in term for x in ['例：','cf.','e.g.','＜印','印の','表記','省略','参照']): return False
    if not re.match(r'^[A-Za-z]', en_eq): return False
    if len(term) > 30: return False
    return True

def parse_je(filepath):
    doc = Document(filepath)
    lines = [_para_text(p) for p in doc.paragraphs]
    si = 0
    found_abbrev = False
    for i, ln in enumerate(lines):
        if '略語表記' in ln or 'e.g. =' in ln or 'NPB  =' in ln:
            found_abbrev = True
        if found_abbrev and re.match(r'^[぀-鿿゠-ヿ一-鿿]', ln) and '＜' in ln:
            if '。' not in ln.split('＜')[0] and '例：' not in ln:
                si = i; break
    entries = []
    term, en_eq, defn = None, None, None
    for ln in lines[si:]:
        s = ln.lstrip()
        if not s: continue
        # A new entry line must start with Japanese AND contain ＜ (the en-eq separator).
        # Continuation lines that happen to start with Japanese (e.g. "に振ること。") must NOT
        # be treated as new entries — they should be appended to the current definition.
        is_new = bool(re.match(r'^[぀-鿿゠-ヿ一-鿿]', ln)) and '＜' in ln
        if is_new:
            if term and en_eq and not should_filter(term) and is_valid_je_term(term, en_eq):
                entries.append({'ja': clean(term), 'en': clean(en_eq), 'def': clean(defn or '')})
            left, right = ln.split('＜', 1)
            term = left.strip()
            right = re.sub(r'^[和古]\s*', '', right.strip())
            em = re.match(r'^([A-Za-z0-9][A-Za-z0-9\s\-\/\(\)\.\'\!\,\;\:\&]*?)(?:\s{2,}|(?=[　-鿿぀-ヿ一-鿿]))', right)
            if em:
                en_eq = em.group(1).strip(); defn = right[len(em.group(0)):].strip()
                # defnの先頭が「、英語語句」なら en_eq に追加（例: leg guards、shinguards）
                _more = re.match(r'^[、,]\s*([A-Za-z][A-Za-z0-9\s\-\/\(\)\'\.]*?)(?=[　、,。！？\s]*[ぁ-ん゛゜ァ-ヶー一-龯！-｠]|$)', defn)
                while _more:
                    en_eq = en_eq.rstrip(', ') + ', ' + _more.group(1).strip().rstrip(', ')
                    defn = defn[len(_more.group(0)):].strip()
                    _more = re.match(r'^[、,]\s*([A-Za-z][A-Za-z0-9\s\-\/\(\)\'\.]*?)(?=[　、,。！？\s]*[ぁ-ん゛゜ァ-ヶー一-龯！-｠]|$)', defn)
            else:
                sp = right.find('　')
                if sp > 0: en_eq = right[:sp].strip(); defn = right[sp+1:].strip()
                else:
                    # 全角スペースがない場合：最初の日本語文字の位置で分割
                    m_ja = re.search(r'[ぁ-ん゛゜ァ-ヶー一-龯！-｠]', right)
                    if m_ja:
                        en_eq = right[:m_ja.start()].strip(); defn = right[m_ja.start():].strip()
                    else: en_eq = right.strip()[:100]; defn = ''
        else:
            if defn is not None: defn = smart_join(defn, s)
    if term and en_eq and not should_filter(term) and is_valid_je_term(term, en_eq):
        entries.append({'ja': clean(term), 'en': clean(en_eq), 'def': clean(defn or '')})
    return entries

# ── Extract & deduplicate ─────────────────────────────────────────────────────

# ── データ素材：スナップショットJSONがあれば docx 不要でそれを使う ──────────────
# ej_data.json / je_data.json（完成データ）が同じフォルダにあれば、元docxを読まずに
# それを最終データとして採用する（下流の解析・加筆は空リストで素通り→最後に上書き）。
_SNAP_DIR = os.path.dirname(os.path.abspath(__file__))
_EJ_SNAP = os.path.join(_SNAP_DIR, 'ej_data.json')
_JE_SNAP = os.path.join(_SNAP_DIR, 'je_data.json')
_USE_SNAP = os.path.exists(_EJ_SNAP) and os.path.exists(_JE_SNAP)
if _USE_SNAP:
    print("Using JSON snapshot (ej_data.json / je_data.json) — docx not needed", file=sys.stderr)
    ej1 = []; ej2 = []; je = []
else:
    print("Parsing A–L …", file=sys.stderr)
    ej1 = parse_ej(r'C:\Users\s3104\Downloads\野球用語英和小辞典(A to L ).docx')
    print(f"  {len(ej1)} raw entries", file=sys.stderr)
    print("Parsing M–Z …", file=sys.stderr)
    ej2 = parse_ej(r'C:\Users\s3104\Downloads\野球用語英和小辞典( M to Z ).docx')
    print(f"  {len(ej2)} raw entries", file=sys.stderr)
    print("Parsing 和英 …", file=sys.stderr)
    je = parse_je(r'C:\Users\s3104\Downloads\waei_converted.docx')
    print(f"  {len(je)} raw entries", file=sys.stderr)

seen = set(); ej_all = []
for e in ej1 + ej2:
    k = e['en'].lower()
    if k in EJ_EXCLUDE: continue
    # 見出しの "TERM = equivalent" → en='TERM', ja='= equivalent ...'
    if ' = ' in e['en']:
        _sp = e['en'].split(' = ', 1)
        e['en'] = _sp[0].strip()
        e['ja'] = ('= ' + _sp[1].strip() + ' ' + e['ja']).strip()
        k = e['en'].lower()  # キーを再計算
        if k in EJ_EXCLUDE: continue
    if k not in seen and e['ja']: seen.add(k); ej_all.append(e)
# 2006年以降に新設・改定されたルールを追加
EJ_NEW_RULES = [
    {'en': 'K(reversed)', 'ja': '(Kという文字を裏返して書いて)見逃しの三振 (called strikeout) の略。スコアカードやボックススコア上の略記号。cf. KC, strike out looking'},
    {'en': 'ABS Challenge System',        'ja': '選手がストライク・ボールの判定に異議を申し立て、自動ボールストライク判定システム（トラックマン等を使用した電子ストライクゾーン）による再審査を要求できる制度。各チームに1試合あたり一定回数のチャレンジ権が与えられ、チャレンジが成功した場合は権利が返還される。球審の判定を完全に自動化するのではなく、人間の審判を主体としつつ明らかな誤審を救済するハイブリッド方式を採用している。マイナーリーグでの試験運用を経て、2026年シーズンからMLBで正式に導入された。'},
    {'en': 'Automatic Intentional Walk', 'ja': '監督がベンチから審判に申告するだけで故意四球が成立し、投手が4球を実際に投げる必要がなくなった。それまでは意図的に外れた球を4球投げる必要があり、その間に暴投や打者がバットに当てるなどのアクシデントが稀に起きることもあった。試合時間短縮策の一環として2017年に導入。'},
    {'en': 'Disengagement Limit',         'ja': '投手が1打者との対戦中に牽制球を投げる、またはプレートを外す行為は合計2回までに制限される。3回目の牽制やステップオフで走者をアウトにできなかった場合は自動的にボークとなり、全走者が1つ進塁する。ピッチクロックと連動した制度であり、投手が時間稼ぎや打者のタイミング外しを目的に過剰な牽制をすることを防ぐ狙いがある。2023年の導入後、盗塁数が大幅に増加した一因とされている。'},
    {'en': 'Expanded Instant Replay',     'ja': '審判の判定に対し、監督がビデオ検証（チャレンジ）を要求できる制度。ニューヨークのリプレイ・オペレーション・センターで専任スタッフが映像を確認し、判定を覆すかどうかを決定する。対象はホームラン判定、タッグプレー、フォースプレー、フェア・ファウル、捕球の可否など多岐にわたる。ただしボール・ストライクの判定はABSチャレンジとは別の制度であり、通常のチャレンジの対象外。チャレンジは前半は1回、判定が覆れば追加で使用できる。'},
    {'en': 'Extra-Innings Tiebreaker Rule','ja': '延長戦の各回開始時に、その回の先頭打者の前の打順にあたる選手が二塁走者として自動的に置かれる。得点が入りやすくなるため試合が長引くリスクが減り、投手の消耗も抑えられる。2020年にコロナ禍での特別ルールとして導入され、賛否両論があったものの試合時間の大幅な短縮効果が認められ2023年に正式採用となった。従来のMLBは「引き分けなし」を原則としており、過去には18回や20回を超える試合も珍しくなかった。'},
    {'en': 'Foreign Substance Crackdown', 'ja': '審判が試合中に投手の手、指、グラブ、帽子などを検査できる。スパイダータックと呼ばれる強力な粘着物質をはじめとする異物をボールに塗ることで、回転数を人工的に高めて変化球の切れを増す行為が2010年代後半に広まり、打者との極端な不均衡が問題視されるようになった。違反が発覚した場合は即時退場および10試合の出場停止処分となる。2021年から取り締まりが大幅に強化された。なお松脂（ロジン）の使用は従来通り認められている。'},
    {'en': 'Home Plate Collision Rule',   'ja': '走者が捕手に故意に体当たりすることを原則禁止。捕手側もボールを持っていない状態で走者の走路をブロックすることが禁止される。捕手はボールを待ちながら本塁の一部を空けておかなければならない。違反した場合、走者はアウト、あるいは捕手の妨害としてセーフの判定が下される。それまでの野球では本塁でのクロスプレーに体当たりは当然のように行われており、捕手が大きなケガを負うケースも多かったことから2014年に導入された。'},
    {'en': 'Larger Bases',                'ja': '一・二・三塁のベースが従来の15インチ（約38cm）四方から18インチ（約46cm）四方に拡大された。主な目的は一塁付近での接触事故の減少だが、ベースが大きくなることで一塁到達距離がわずかに短くなり、内野安打や盗塁の成功率が若干上がる効果もある。2023年の導入後、実際に盗塁数が増加傾向となり、ピッチクロックや牽制制限との相乗効果で走塁が活発化した。'},
    {'en': 'Ohtani Rule (Two-Way Player Exception)', 'ja': 'DHを兼任している投手が、マウンドを降りた後も指名打者としてその試合に打席に立ち続けられる特例ルール。通常、投手がマウンドを降りると同時にDHも消滅し、投手の打順にはそのまま後続の打者が入らなければならないが、2022年のこの特例によって大谷翔平のような二刀流選手が投手として降板した後もDHとして継続出場できるようになった。2022年のユニバーサルDH導入と同時に設けられ、事実上大谷のために作られたルールとして「大谷ルール」と呼ばれる。'},
    {'en': 'Pitch Clock',                 'ja': '投手は走者がいない場面では15秒以内、走者がいる場面では20秒以内に投球動作を開始しなければならない。打者側にも、投球の8秒前までに打席に入り構えを完了する義務がある。投手の違反は自動的にボール、打者の違反は自動的にストライクが宣告される。さらに投手がタイムを要求してプレートを外す行為（牽制を含む）は1打席に2回までに制限される（牽制回数制限と連動）。2023年の導入初年度に平均試合時間が約30分短縮されるという劇的な効果をもたらした。'},
    {'en': 'Shift Restriction',           'ja': '投球時に内野手4人全員がダート（土の部分）に位置し、二塁ベースを挟んで左右に2人ずつ配置しなければならない。それまでは打者の打球傾向に応じて内野手を極端に偏らせる「シフト」が広く使われており、特に左の長距離打者に対して三塁側に3〜4人を置くシフトが常態化していた。これが安打数の減少や得点力低下につながるとして批判され、2023年に禁止された。外野手の配置については現時点では制限がない。'},
    {'en': 'Slide Rule (Takeout Slide Ban)', 'ja': '走者は塁に到達することを目的とした正当なスライディングをしなければならない。ゲッツー崩しを目的として野手に向かって体を投げ出したり、故意に接触しにいく行為は禁止。具体的には、スライディングの際に手や足がベースに届く範囲内に体が収まっていること、かつ野手へのコンタクトを意図していないことが条件とされる。違反した場合は走者アウトに加え、打者走者も併殺となる。2016年のポストシーズンにおけるチェース・ユーティリスの二塁タックルプレーが物議を醸したことが導入の直接的なきっかけとなった。'},
    {'en': 'Three-Batter Minimum Rule',   'ja': 'リリーフ投手は最低3人の打者に対して投げるか、そのイニングが終了するまで交代できない。それまでは左打者一人だけに対して登板してすぐ交代する「ワンポイントリリーフ」が頻繁に使われており、試合のテンポを著しく損なうとして批判されていた。ただし負傷や疾病による交代はこのルールの例外とされる。2020年に導入され、ブルペンの起用戦略に大きな変化をもたらした。'},
    {'en': 'Universal Designated Hitter', 'ja': '指名打者制が両リーグに適用され、投手が打席に立つことがなくなった。もともとアメリカンリーグは1973年から採用していたが、ナショナルリーグは「投手が打つ」という伝統を長年守り続けており、両リーグの違いが戦略面での面白さともなっていた。2020年にコロナ禍の特別措置として一時的に両リーグに導入し、2022年に正式に統一採用となった。これによりダブルスイッチの本来の意義が薄れるなど、監督の采配にも影響が出ている。'},
    {'en': 'last ten games', 'ja': '直近10試合。略記号 = L10\ne.g. This team has been on a roll — just check their record in the last ten games.（このチームは好調が続いている — 直近10試合の成績を見れば一目瞭然だ）'},
    {'en': 'loser', 'ja': '敗戦投手。：cf. L'},
    {'en': 'ALCS', 'ja': '= American League Championship Series の略。頭文字語として用いる。cf. NLCS'},
    {'en': 'NLCS', 'ja': '= National League Championship Series の略。cf. ALCS'},
    {'en': 'American League Championship Series', 'ja': 'アメリカンリーグ選手権シリーズ。7回戦制で行われ、4勝した方が優勝。cf. ALCS, NLCS'},
    {'en': 'National League Championship Series', 'ja': 'ナショナル・リーグ選手権シリーズ。7回戦制で行われ、4勝した方が優勝。cf. ALCS, NLCS'},
    # 野球用語
{'en': 'Fenway Park single', 'ja': 'フェンウェイ・パーク・ヒット。：レフト後方の高いフェンス（モンスター）に勢いよく打球が跳ね返りすぎて長打にならずに単打になったもの。'},
    {'en': 'Take Me Out To The Ball Game', 'ja': '「私を球場に連れてって」という歌のタイトル。作詞：Jack Norworth、作曲：Albert Von Tilzer。この歌はSeventh Inning Stretchの時に観客全員が大合唱する。cf. seventh inning stretch'},
    {'en': 'trip', 'ja': '遠征（試合）。e.g. The club went 10-0 on a swing in April.（その球団は4月の遠征では10勝0敗であった）'},
    {'en': 'rubber (the)', 'ja': '1. = pitcher\'s plate/pitching plate 投手の踏み板。cf. slab, step off the rubber\n2. = rubber game/match（the）シリーズの相星決戦。'},
    # 統計略語
    {'en': 'OPS',   'ja': 'On-base Plus Slugging',              'def': 'どれくらい塁に出て、どれくらい長打を打つかを1つにまとめた打撃の総合力', 'cat': 'abbr'},
    {'en': 'WAR',   'ja': 'Wins Above Replacement',             'def': 'その選手がいることで、チームが何勝増えるかを表す総合貢献度', 'cat': 'abbr'},
    {'en': 'wRC+',  'ja': 'Weighted Runs Created Plus',         'def': '平均選手=100として、どれだけ多く点を生み出すかを示す打撃力', 'cat': 'abbr'},
    {'en': 'FIP',   'ja': 'Fielding Independent Pitching',      'def': '守備に左右されず、投手自身の力だけでどれだけ抑えたかを見る指標', 'cat': 'abbr'},
    {'en': 'xBA',   'ja': 'Expected Batting Average',           'def': '打球の質から見た「本来これくらい打てるはずの打率」', 'cat': 'abbr'},
    {'en': 'xSLG',  'ja': 'Expected Slugging Percentage',       'def': '打球の質から見た「本来これくらい長打が出るはず」という長打力', 'cat': 'abbr'},
    {'en': 'BABIP', 'ja': 'Batting Average on Balls In Play',   'def': 'フェアグラウンドに飛んだ打球がヒットになる割合（運の影響が大きい）', 'cat': 'abbr'},
    {'en': 'OBP',   'ja': 'On-Base Percentage',                 'def': 'ヒット・四球など全部含めてどれだけアウトにならず塁に出たか', 'cat': 'abbr'},
    {'en': 'SLG',   'ja': 'Slugging Percentage',                'def': 'ヒットの質（単打より長打を重く評価した打率）', 'cat': 'abbr'},
    {'en': 'ERA',   'ja': 'Earned Run Average',                 'def': '9イニングあたり何点取られたか（昔からある投手評価）', 'cat': 'abbr'},
    {'en': 'WHIP',  'ja': 'Walks plus Hits per Inning Pitched', 'def': '1イニングで何人ランナーを出したか（少ないほど良い）', 'cat': 'abbr'},
    {'en': 'K/9',   'ja': 'Strikeouts per 9 Innings',           'def': '9イニングでどれだけ三振を取るか', 'cat': 'abbr'},
    {'en': 'BB/9',  'ja': 'Walks per 9 Innings',                'def': '9イニングでどれだけ四球を出すか（少ないほど良い）', 'cat': 'abbr'},
    {'en': 'HR/9',  'ja': 'Home Runs per 9 Innings',            'def': '9イニングでどれだけホームランを打たれるか', 'cat': 'abbr'},
    {'en': 'wOBA',  'ja': 'Weighted On-Base Average',           'def': 'ヒットや四球などに本当の得点価値をつけた出塁力指標', 'cat': 'abbr'},
    {'en': 'ISO',   'ja': 'Isolated Power',                     'def': '長打力だけを切り出したシンプルなパワー指標', 'cat': 'abbr'},
    {'en': 'xwOBA', 'ja': 'Expected Weighted On-Base Average',  'def': '打球内容から見た本来の出塁力', 'cat': 'abbr'},
    {'en': 'fWAR',  'ja': 'FanGraphs Wins Above Replacement',   'def': 'FanGraphs方式のWAR（守備より分析寄り）', 'cat': 'abbr'},
    {'en': 'bWAR',  'ja': 'Baseball-Reference WAR',             'def': 'Baseball-Reference方式のWAR（失点ベースで計算）', 'cat': 'abbr'},
    # リーグ区分
    {'en': 'AA', 'ja': 'マイナーリーグの第2番目のレベル。ダブルAとかAAクラスといわれる。cf. Double A, AAA, Triple A'},
    {'en': 'AAA', 'ja': 'マイナーリーグの第1番目のレベル。トリプルAとかAAAクラスといわれる。cf. Triple A, AA, Double A'},
    {'en': 'Single A', 'ja': 'Aクラス（の）：マイナーリーグの第3番目のレベル。cf. AA, AAA, Double A, Triple A'},
    {'en': 'Double A', 'ja': '2Aクラス（の）：マイナーリーグの第2番目のレベル。cf. AA, Single A, AAA, Triple A'},
    {'en': 'Triple A', 'ja': '3Aクラス（の）：マイナーリーグの最上位レベル。cf. AAA, Double A, Single A'},
    {'en': 'skipper', 'ja': '= manager 監督の俗称。：cf. field boss, field manager, skip'},
    {'en': 'advance a runner ( to ～ )', 'ja': '走者を（～へ）進塁させる'},
    {'en': 'advance ( to ～ )', 'ja': '（走者が～へ）進塁する'},
    {'en': 'board (the)', 'ja': '= scoreboard スコアボード。cf. put a team on the board'},
    {'en': 'forty-forty man', 'ja': '= 40-40 man 1シーズンに40本塁打、40盗塁した選手。e.g. Jose Canseco was the first forty-forty man in Major League history.（ホセ・カンセコはメジャーリーグ史上初の40-40選手だった）'},
]
for r in EJ_NEW_RULES:
    k = r['en'].lower()
    if k not in seen:
        seen.add(k)
        ej_all.append(r)

# 全30球団の本拠地球場（2025年時点）
_STADIUMS = [
    # AL East
    {'en': 'Oriole Park at Camden Yards', 'ja': 'オリオール・パーク・アット・カムデン・ヤーズ。ボルティモア・オリオールズの本拠地。メリーランド州ボルティモア。', 'def': ''},
    {'en': 'Fenway Park',                  'ja': 'フェンウェイ・パーク。ボストン・レッドソックスの本拠地。マサチューセッツ州ボストン。1912年開場の最も歴史ある現役球場のひとつ。グリーン・モンスター（左翼フェンス）で有名。', 'def': ''},
    {'en': 'Yankee Stadium',               'ja': 'ヤンキー・スタジアム。ニューヨーク・ヤンキースの本拠地。ニューヨーク州ブロンクス。2009年開場。', 'def': ''},
    {'en': 'Tropicana Field',              'ja': 'トロピカーナ・フィールド。タンパベイ・レイズの本拠地。フロリダ州セントピーターズバーグ。屋根付きドーム球場。', 'def': ''},
    {'en': 'Rogers Centre',                'ja': 'ロジャーズ・センター。トロント・ブルージェイズの本拠地。カナダ・オンタリオ州トロント。開閉式屋根を持つ。', 'def': ''},
    # AL Central
    {'en': 'Guaranteed Rate Field',        'ja': 'ギャランティード・レート・フィールド。シカゴ・ホワイトソックスの本拠地。イリノイ州シカゴ。', 'def': ''},
    {'en': 'Progressive Field',            'ja': 'プログレッシブ・フィールド。クリーブランド・ガーディアンズの本拠地。オハイオ州クリーブランド。', 'def': ''},
    {'en': 'Comerica Park',                'ja': 'コメリカ・パーク。デトロイト・タイガースの本拠地。ミシガン州デトロイト。', 'def': ''},
    {'en': 'Kauffman Stadium',             'ja': 'カウフマン・スタジアム。カンザスシティ・ロイヤルズの本拠地。ミズーリ州カンザスシティ。', 'def': ''},
    {'en': 'Target Field',                 'ja': 'ターゲット・フィールド。ミネソタ・ツインズの本拠地。ミネソタ州ミネアポリス。', 'def': ''},
    # AL West
    {'en': 'Minute Maid Park',             'ja': 'ミニット・メイド・パーク。ヒューストン・アストロズの本拠地。テキサス州ヒューストン。開閉式屋根を持つ。', 'def': ''},
    {'en': 'Angel Stadium',                'ja': 'エンジェル・スタジアム。ロサンゼルス・エンゼルスの本拠地。カリフォルニア州アナハイム。', 'def': ''},
    {'en': 'Sutter Health Park',           'ja': 'サター・ヘルス・パーク。オークランド・アスレティックスが新球場建設中に使用する暫定本拠地。カリフォルニア州サクラメント。', 'def': ''},
    {'en': 'T-Mobile Park',                'ja': 'T-モバイル・パーク。シアトル・マリナーズの本拠地。ワシントン州シアトル。開閉式屋根を持つ。', 'def': ''},
    {'en': 'Globe Life Field',             'ja': 'グローブ・ライフ・フィールド。テキサス・レンジャーズの本拠地。テキサス州アーリントン。2020年開場の開閉式屋根付き球場。', 'def': ''},
    # NL East
    {'en': 'Truist Park',                  'ja': 'トゥルーイスト・パーク。アトランタ・ブレーブスの本拠地。ジョージア州カンバーランド。', 'def': ''},
    {'en': 'loanDepot park',               'ja': 'ローンデポ・パーク。マイアミ・マーリンズの本拠地。フロリダ州マイアミ。開閉式屋根を持つ。', 'def': ''},
    {'en': 'Citi Field',                   'ja': 'シティ・フィールド。ニューヨーク・メッツの本拠地。ニューヨーク州クイーンズ。2009年開場。', 'def': ''},
    {'en': 'Citizens Bank Park',           'ja': 'シティズンズ・バンク・パーク。フィラデルフィア・フィリーズの本拠地。ペンシルベニア州フィラデルフィア。', 'def': ''},
    {'en': 'Nationals Park',               'ja': 'ナショナルズ・パーク。ワシントン・ナショナルズの本拠地。ワシントンD.C.。', 'def': ''},
    # NL Central
    {'en': 'Wrigley Field',                'ja': 'リグレー・フィールド。シカゴ・カブスの本拠地。イリノイ州シカゴ。1914年開場の歴史的球場。ツタに覆われた外野フェンスで有名。', 'def': ''},
    {'en': 'Great American Ball Park',     'ja': 'グレート・アメリカン・ボールパーク。シンシナティ・レッズの本拠地。オハイオ州シンシナティ。', 'def': ''},
    {'en': 'American Family Field',        'ja': 'アメリカン・ファミリー・フィールド。ミルウォーキー・ブルワーズの本拠地。ウィスコンシン州ミルウォーキー。開閉式屋根を持つ。', 'def': ''},
    {'en': 'PNC Park',                     'ja': 'PNCパーク。ピッツバーグ・パイレーツの本拠地。ペンシルベニア州ピッツバーグ。アレゲニー川沿いに立地する美しい球場として名高い。', 'def': ''},
    {'en': 'Busch Stadium',                'ja': 'ブッシュ・スタジアム。セントルイス・カーディナルスの本拠地。ミズーリ州セントルイス。', 'def': ''},
    # NL West
    {'en': 'Chase Field',                  'ja': 'チェース・フィールド。アリゾナ・ダイヤモンドバックスの本拠地。アリゾナ州フェニックス。開閉式屋根と天然芝を持つ。', 'def': ''},
    {'en': 'Coors Field',                  'ja': 'クアーズ・フィールド。コロラド・ロッキーズの本拠地。コロラド州デンバー。標高1600m超の高地にあり打球が飛びやすい球場として有名。', 'def': ''},
    {'en': 'Dodger Stadium',               'ja': 'ドジャー・スタジアム。ロサンゼルス・ドジャースの本拠地。カリフォルニア州ロサンゼルス。1962年開場。MLB最大収容人数を誇る。', 'def': ''},
    {'en': 'Petco Park',                   'ja': 'ペトコ・パーク。サンディエゴ・パドレスの本拠地。カリフォルニア州サンディエゴ。', 'def': ''},
    {'en': 'Oracle Park',                  'ja': 'オラクル・パーク。サンフランシスコ・ジャイアンツの本拠地。カリフォルニア州サンフランシスコ。マッコビー・コーブに面し、本塁打が湾に飛び込むことで有名。', 'def': ''},
]
for _s in _STADIUMS:
    k = _s['en'].lower()
    if k not in seen:
        seen.add(k)
        ej_all.append(_s)

# 現行チーム名（EJ）- cf.リンク先として必要
_CURRENT_TEAMS_EJ = [
    {'en': 'Miami Marlins (the)',           'ja': 'マイアミ・マーリンズ（ナショナル・リーグ東地区）：フロリダ・マーリンズが2012年に改称。本拠地はローンデポ・パーク。', 'def': ''},
    {'en': 'Tampa Bay Rays (the)',          'ja': 'タンパベイ・レイズ（アメリカン・リーグ東地区）：タンパベイ・デビルレイズが2008年に改称。本拠地はトロピカーナ・フィールド。', 'def': ''},
    {'en': 'Cleveland Guardians (the)',     'ja': 'クリーブランド・ガーディアンズ（アメリカン・リーグ中地区）：クリーブランド・インディアンスが2022年に改称。本拠地はプログレッシブ・フィールド。', 'def': ''},
    {'en': 'Los Angeles Angels (the)',      'ja': 'ロサンゼルス・エンゼルス（アメリカン・リーグ西地区）：アナハイム・エンゼルス→ロサンゼルス・エンゼルス・オブ・アナハイムを経て2016年に改称。本拠地はエンジェル・スタジアム。', 'def': ''},
    {'en': 'Los Angeles Angels of Anaheim (the)', 'ja': 'ロサンゼルス・エンゼルス・オブ・アナハイム（アメリカン・リーグ西地区）：アナハイム・エンゼルスが2005年に改称。2016年にロサンゼルス・エンゼルスへさらに改称。（追記）現在は無い。cf. Los Angeles Angels', 'def': ''},
    {'en': 'Athletics (the)',                  'ja': 'アスレチックス（アメリカン・リーグ西地区）：オークランド・アスレチックスが2025年に改称。2024年シーズン終了をもってオークランドを撤退し、サクラメントを暫定本拠地とする。ラスベガスに建設中の新球場完成後（2028年予定）移転予定。', 'def': ''},
]
for _t in _CURRENT_TEAMS_EJ:
    k = _t['en'].lower()
    if k not in seen:
        seen.add(k)
        ej_all.append(_t)

# 旧球団名エントリ（EJ）
_OLD_TEAMS_EJ = [
    {'en': 'Montreal Expos (the)',          'ja': 'モントリオール・エクスポズ（ナショナル・リーグ東地区）：1969年創設。カナダ・モントリオールを本拠地としたMLB唯一のカナダ球団（当時）。（追記）現在は無い。2005年にワシントンへ移転し cf. Washington Nationals となった。', 'def': ''},
    {'en': 'Florida Marlins (the)',         'ja': 'フロリダ・マーリンズ（ナショナル・リーグ東地区）：1993年創設。フロリダ州マイアミを本拠地とした球団。（追記）現在は無い。2012年に cf. Miami Marlins へ改称。', 'def': ''},
    {'en': 'Tampa Bay Devil Rays (the)',    'ja': 'タンパベイ・デビルレイズ（アメリカン・リーグ東地区）：1998年創設。フロリダ州セントピーターズバーグを本拠地とした球団。（追記）現在は無い。2008年に cf. Tampa Bay Rays に改称。', 'def': ''},
    {'en': 'Cleveland Indians (the)',       'ja': 'クリーブランド・インディアンス（アメリカン・リーグ中地区）：1901年創設。オハイオ州クリーブランドを本拠地とした球団。（追記）現在は無い。2022年に cf. Cleveland Guardians に改称。', 'def': ''},
    {'en': 'Los Angeles Angels of Anaheim (the)', 'ja': 'ロサンゼルス・エンゼルス・オブ・アナハイム（アメリカン・リーグ西地区）：アナハイム・エンゼルスが2005年に改称。（追記）現在は無い。2016年に cf. Los Angeles Angels に改称。', 'def': ''},
]
for _t in _OLD_TEAMS_EJ:
    k = _t['en'].lower()
    if k not in seen:
        seen.add(k)
        ej_all.append(_t)

# EJエントリへの追記（定義末尾に追加）
EJ_ADDENDA = {
    'W':         '（追記）現状は、依然存在するが、評価指標としての重要度が激減。FIP・WAR等に取って代わられた。',
    'CG':        '（追記）記録としては残るが、現代では1シーズンに数えるほどしか起きない。',
    'SHO':       '（追記）現代の分業制では非常にまれ。',
    'GS':        '（追記）現状は、統計としては残るが、以前ほど重視されない。',
    'SB':        '（追記）現状は、2023年以降のルール改正（ピッチクロック・大きなベース）で盗塁数・文脈が変化。',
    'RBI,/rbi':  '（追記）依然広く使われるが、文脈依存統計として信頼性が低いとされ、セイバーメトリクスでは軽視される。',
}
for _e in ej_all:
    _note = EJ_ADDENDA.get(_e['en'])
    if _note:
        if _e.get('def'):
            _e['def'] = _e['def'].rstrip('。') + '。\n' + _note
        else:
            _e['ja'] = (_e.get('ja') or '').rstrip('。') + '。\n' + _note

# 定義文の上書き（DOCXの内容より優先）
EJ_DEF_OVERRIDE = {
    'role player': '特定の役割に特化した選手。：cf. utility man/player',
    'bunt': 'バントする。バットをスイングせずに当てるだけで打つ技術。送りバント（sacrifice bunt）・セーフティバント（squeeze bunt）などがある。cf. sacrifice bunt, squeeze play, drag bunt',
    'a': '1．マイナーリーグの第3番目のレベル；シングルAまたは Aクラスと言われる。 , Single A, AA, Double A , AAA , Triple A（追記）現状は、シングルAに統一。\n2．「補殺」のスコアカードやボックススコア上の略記号。 cf. assist, scorecard\n3．= attendance入場者数のスコアカードやボックススコア上の略記号。 cf. attendance, box score（追記）メディアによっては「ATT」と略されることがある。',
}
for e in ej_all:
    k = e['en'].lower()
    if k in EJ_DEF_OVERRIDE:
        e['ja'] = EJ_DEF_OVERRIDE[k]

# 統計略語として扱うべき既存エントリーにcatフラグを付与
_ABBR_TAG_KEYS = {'era', 'obp', 'whip', 'slg', 'ops', 'war', 'wrc+', 'fip', 'xba', 'xslg',
                  'babip', 'k/9', 'bb/9', 'hr/9', 'woba', 'iso', 'xwoba', 'fwar', 'bwar'}
# 辞書中の略語エントリを自動検出してabbr扱いにする
# 条件: 見出し語が略語らしい形 かつ 定義に「の略」「略称」等を含む
_abbr_en_pat = re.compile(r'^[A-Z][A-Z0-9/\-().,+]{0,9}$|^[a-z]{1,4}/[A-Z]{1,4}$')
_abbr_def_pat = re.compile(r'の略(?:称|語)?|略語|略称')
# チーム略称・リーグ名略称は略語辞典に含めない
_ABBR_EXCLUDE_KEYS = {
    'al','nl','ml','mlb','alcs','nlcs','alds','nlds','ws',
    'ana','ari','atl','alt','bal','bos','chc','chw','cin','cle','col',
    'det','fla','hou','kc','la','laa','lad','mia','mil','min','mln',
    'nym','nyy','oak','oka','phi','pit','sd','sea','sf','sfg',
    'stl','tb','tex','tor','was','wsh',
    # 時代とともに重要度が変化した略語（辞書には残すが略語辞典からは除外）
    'w','cg','sho','gs','sb','rbi,/rbi',
}
for _e in ej_all:
    if _e['en'].lower() in _ABBR_TAG_KEYS and re.match(r'^[A-Z0-9/+]+$', _e['en']):
        _e['cat'] = 'abbr'
    elif _e.get('cat') != 'abbr':
        _en = _e['en'].split(',')[0].split('/')[0].strip()
        if _en.lower() in _ABBR_EXCLUDE_KEYS:
            continue
        _text = (_e.get('ja') or '') + ' ' + (_e.get('def') or '')
        # チーム名・固有名詞の略称は除外（「= ○○ の略称」パターン）
        if re.match(r'^= [A-Z][a-zA-Z]', _text):
            continue
        if _abbr_en_pat.match(_en) and _abbr_def_pat.search(_text):
            _e['cat'] = 'abbr'

ej_all.sort(key=lambda x: (x.get('sk') or x['en']).lower())

# 見出し表記の正規化（大文字略語に小文字バリアントを付加）
_EN_ALIASES = {'3B': '3B, 3b'}
for e in ej_all:
    if e['en'] in _EN_ALIASES:
        e['en'] = _EN_ALIASES[e['en']]

# Merge -plus innings content into play third base (they were one entry in the original)
_play_idx = next((i for i,e in enumerate(ej_all) if e['en'] == 'play third base'), None)
_plus_idx  = next((i for i,e in enumerate(ej_all) if e['en'] == '-plus innings'), None)
if _play_idx is not None and _plus_idx is not None:
    ej_all[_play_idx]['ja'] += '\n-plus innings: ' + ej_all[_plus_idx]['ja']
    ej_all.pop(_plus_idx)

# 追記：2006年以降のルール変更
EJ_ADDENDA = {
    'extra innings':                    '（追記）2020年よりMLBでタイブレーカー制（延長10回から無死二塁スタート）が導入された。',
    'designated hitter':                '（追記）2022年よりMLBは両リーグで指名打者制（ユニバーサルDH）を採用。',
    'national league (the)':            '（追記）2022年以降は両リーグDH制のため、指名打者制のないリーグという説明は現在当てはまらない。',
    'double switch':                    '（追記）現在は両リーグDH制のため、「指名打者制のないナショナル・リーグで」という前提は実質的に意味が変化している。\nダブルスイッチとは、もともとナショナルリーグ（投手が打席に立つリーグ）特有の継投策でした。投手交代の際に、同時に野手も交代させ、打順の位置を入れ替えることで「新しい投手がすぐに打席に回ってこないようにする」テクニックです。たとえば投手が打順の2番にいる場面で交代させたいとき、そのまま替えると新投手もすぐ2番の打順が来てしまう。そこで9番の野手とセットで交代し、新投手を9番の打順に置けば、しばらく打席が回ってこない、という策です。',
    'wild card (the)':                  '（追記）2022年よりワイルドカードは各リーグ3チームに拡大され、ワイルドカードシリーズ（3回戦制）が新設された。',
    'playoffs (the)':                   '（追記）現在は各リーグ6チームによる新方式。',
    'division series (the)':            '（追記）現在は各リーグ6チームによる新方式。',
    'standings':                        '（追記）タイブレーク制の導入でゲーム構造が変化している。',
    'waiver':                           '（追記）NPBについて「逆指名制度」とあるが、逆指名制は2007年に廃止され、現在は完全ウェーバー制（育成ドラフトを含む）。',
    'babe ruth':                        '（追記）「現在はバリー・ボンズの年間71本が最高記録」とあるが、この記録は更新されていない。一方、バリー・ボンズは薬物問題で名誉棄損的扱いを受け、殿堂入りも見送られている（2023年に改選から除外）。また、バリー・ボンズは本書の例文に多数登場するが、現在はステロイド使用問題で評価が大きく変化している。',
    'curfew':                           '（追記）「アメリカンリーグでは午前1時を過ぎて新しいイニングに入れない」という規則は現在廃止されている。',
    'curfew rule':                      '（追記）「アメリカンリーグでは午前1時を過ぎて新しいイニングに入れない」という規則は現在廃止されている。',
    'spit ball':                        '（追記）辞典時点ですでに禁止と記載されているが、2021年の粘着物質取締り強化でさらに厳格化された。',
    'aa':                               '（追記）現状は、ダブルAに統一。',
    'aaa':                              '（追記）現状は、トリプルAに統一。',
    'anaheim angels (the)':             '（追記）現在は無い。2005年より cf. Los Angeles Angels of Anaheim へ改称、さらに2016年より cf. Los Angeles Angels へ改称。',
    'ana':                              '（追記）現在はロサンゼルス・エンゼルスの略称 LAA を使用。',
    'florida marlins (the)':            '（追記）現在は無い。2012年に cf. Miami Marlins に改称。',
    'montreal expos (the)':             '（追記）現在は無い。2005年にワシントンDCへ移転し cf. Washington Nationals となった。',
    'tampa bay devil rays (the)':       '（追記）現在は無い。2008年より cf. Tampa Bay Rays へ改称。',
    'cleveland indians (the)':          '（追記）現在は無い。2022年より cf. Cleveland Guardians へ改称。',
    'oakland athletics':                '（追記）2024年シーズン終了をもってオークランドを撤退。2025年よりチーム名を cf. Athletics に変更し、サクラメントを暫定本拠地とする。ラスベガスに新球場を建設中（2028年開場予定）。',
    'houston astros':                   '（追記）2013年よりアメリカン・リーグ西地区に移籍（ナショナル・リーグ中地区から変更）。',
    'washington nationals':             '（追記）本拠地球場は2008年よりナショナルズ・パーク（ワシントンD.C.）に移転。RFKスタジアムでの使用は終了。',
    'milwaukee brewers (the)':          '（追記）本拠地ミラー・パークは2021年にアメリカン・ファミリー・フィールドに改称。',
    'oakland athletics (the)':          '（追記）歴代の本拠地及び（チーム名）\n1901～1954：Philadelphia（Philadelphia Athletics）\n1955～1967：Kansas City（Kansas City Athletics）\n1968～2024：オークランド（Oakland Athletics）\n2025～2027：West Sacramento（Athletics）\n2028～：Las Vegasへ移転予定',
    "a's":                              '（追記）歴代の本拠地及び（チーム名）\n1901～1954：Philadelphia（Philadelphia Athletics）\n1955～1967：Kansas City（Kansas City Athletics）\n1968～2024：オークランド（Oakland Athletics）\n2025～2027：West Sacramento（Athletics）\n2028～：Las Vegasへ移転予定',
}

JE_ADDENDA = {
    '延長（戦）':          '（追記）タイブレーク制の導入：2020年にコロナ禍の特別ルールとして導入、2023年に正式採用 cf.タイブレーク制',
    '延長(戦)':           '（追記）タイブレーク制の導入：2020年にコロナ禍の特別ルールとして導入、2023年に正式採用 cf.タイブレーク制',
    'アメリカンリーグ':    '（追記）「午前1時を過ぎると新しいイニングには入れない」というカーファールールは現在廃止されている。',
    '指名打者':           '（追記）2022年よりMLBは両リーグで指名打者制（ユニバーサルDH）を採用。',
    'ナショナルリーグ':    '（追記）2022年以降は両リーグDH制のため、指名打者制のないリーグという説明は現在当てはまらない。ユニバーサルDH導入（2022年）により、ナショナルリーグにおける投手の打席は完全に消滅した。',
    'ワイルドカード':      '（追記）2022年よりワイルドカードは各リーグ3チームに拡大され、ワイルドカードシリーズ（3回戦制）が新設された。辞典が説明する1チーム制・4チームによるプレーオフという形式は現在存在しない。',
    'プレーオフ':          '（追記）現在は各リーグ6チームによる新方式。',
    'ディビジョン・シリーズ': '（追記）現在は各リーグ6チームによる新方式。',
    '地区シリーズ':        '（追記）現在は各リーグ6チームによる新方式。',
    '順位表':             '（追記）タイブレーク制の導入でゲーム構造が変化している。',
    'ウエーバー方式':      '（追記）NPBについて「逆指名制度」とあるが、逆指名制は2007年に廃止され、現在は完全ウェーバー制（育成ドラフトを含む）。',
    'パシフィックリーグ':  '球団名の記述は2006年のものです。',
    'ワンポイントリリーフ': '（追記）2020年の最低3人打者ルール導入により廃止された。辞典では通常の戦術として説明されているが、現在は使用できない。',
    'スピットボール':      '（追記）辞典では「意外な変化をするため現在は禁止されている」と記載されており廃止済みだが、2021年の粘着物質取締り強化により、関連する不正投球への対応がさらに厳格化された。',
}

for e in ej_all:
    note = EJ_ADDENDA.get(e['en'].lower())
    if note:
        e['ja'] = e['ja'].rstrip().rstrip('。') + '。\n' + note if e['ja'] else note

# ph: PHI・Philadelphia Philliesの定義が混入しているため正しい定義に上書き
_ph_entry = next((e for e in ej_all if e['en'] == 'ph'), None)
if _ph_entry:
    _ph_entry['ja'] = '= pinch hitter の略。スコアカードやボックススコア上の略記号。cf. pinch hitter'

# uppercut: 主観的コメントを削除
_uc_entry = next((e for e in ej_all if e['en'].lower() == 'uppercut'), None)
if _uc_entry:
    _uc_entry['ja'] = _uc_entry['ja'].replace('：あまりよい打ち方とはされない。', '').replace('あまりよい打ち方とはされない。', '').strip()

# twi-nighter: カンマ欠落・スペルミスを修正
_twi_entry = next((e for e in ej_all if e['en'] == 'twi-nighter'), None)
if _twi_entry:
    _twi_entry['ja'] = _twi_entry['ja']\
        .replace('day-night double header', 'day-night doubleheader')\
        .replace('nightcap twin killing', 'nightcap, twin killing')

# L: 「LA = Los Angeles Dodgers の略記号。」を除去
_l_entry = next((e for e in ej_all if e['en'] == 'L'), None)
if _l_entry:
    _l_entry['ja'] = _l_entry['ja'].replace('LA = Los Angeles Dodgers の略記号。', '')\
        .replace('losses となる。', 'losses となる。cf. losses/Losses').rstrip()\
        .rstrip('。') + '。'

# L10: 「最近の過去10試合」→「直近10試合」に修正
_l10_entry = next((e for e in ej_all if e['en'] == 'L10'), None)
if _l10_entry:
    _l10_entry['ja'] = _l10_entry['ja'].replace('最近の過去10試合', '直近10試合').replace('成績欄の欄', '成績欄')

# 全EJエントリーの誤字一括修正
for _e in ej_all:
    if '成績蘭' in _e.get('ja', ''):
        _e['ja'] = _e['ja'].replace('成績蘭', '成績欄')

# Wordファイルで省略されていたKの全文を原本から補完
_k_entry = next((e for e in ej_all if e['en'] == 'K'), None)
if _k_entry:
    _k_entry['ja'] = ('strikeout（三振）の略。スコアカードやボックススコア上の略記号。'
        'なぜKなのかについては二説あって、スコア記録の略記をアルファベット1文字で考案した'
        'ヘンリー・チャドウィックがKを三振につけたという説と'
        'ニューヨーク・ヘラルド誌のM. J. ケリーが発明したという説があるが、'
        '三振奪取王のドワイト・グッデンがDr. Kと呼称されたことから定着したといわれている。'
        'cf. K(reversed)')

seen = set(); je_all = []
for e in je:
    k = e['ja']
    if k in JE_EXCLUDE: continue
    if k not in seen and e['en']: seen.add(k); je_all.append(e)

def to_hira(s):
    result = ''
    for ch in s:
        c = ord(ch)
        result += chr(c - 0x60) if 0x30A1 <= c <= 0x30F6 else ch
    return result

def je_sort_key(e):
    yomi = e.get('yomi') or e['ja']
    return to_hira(yomi)
je_all.sort(key=je_sort_key)

for e in je_all:
    note = JE_ADDENDA.get(e['ja'])
    if note:
        if e.get('def'):
            e['def'] = e['def'].rstrip('。') + '。\n' + note
        else:
            e['en'] = e['en'].rstrip('。') + '。\n' + note if e['en'] else note

# アッパースイング: 主観的コメントを削除
for e in je_all:
    if e['ja'] == 'アッパースイング' and e.get('def'):
        e['def'] = e['def'].replace('バットの振りが「アッパー気味になる」のはよくないとされる。', '').strip()

# 個別上書き
for e in je_all:
    if e['ja'] in ('ベーブルース', 'ルース'):
        e['en'] = 'Babe Ruth参照'; e['def'] = ''
    if e.get('def'):
        e['def'] = re.sub(r'^=[^\s　]+[^\s　]+[\s　]+', '', e['def'])
    # 見出し語末尾の「 [」除去
    e['ja'] = e['ja'].rstrip(' [')
    # en中の「英語表現] 日本語説明」を分割
    if ']' in e['en']:
        bracket_pos = e['en'].index(']')
        overflow = e['en'][bracket_pos+1:].lstrip()
        e['en'] = e['en'][:bracket_pos].strip()
        if overflow:
            e['def'] = overflow + (e.get('def') or '')
    # （阿部）を全削除
    e['en']  = re.sub(r'[（(]阿部[）)]', '', e['en']).strip()
    if e.get('def'):
        e['def'] = re.sub(r'[（(]阿部[）)]', '', e['def']).strip()

# 2006年以降の新ルールを和英に追加
JE_MISSING_REFS = [
    # 他のエントリから参照されているが辞書になかった語
    {'ja': '暴投',              'en': 'wild pitch',        'def': '投手の投球がワイルドピッチになること。捕手が捕れないほど大きく外れた投球。野手の送球ミスはワイルドスローという。'},
    {'ja': 'ダブルプレー',      'en': 'double play',       'def': '1つのプレーで2人のアウトを取ること。ゲッツーともいう。cf. ゲッツー'},
    {'ja': 'プレイボール',      'en': 'play ball',         'def': '試合開始・再開を宣言する審判のコール。試合を再開するとき球審は「プレイ」とコールする。'},
    {'ja': 'ホームランボール',  'en': 'home run ball',     'def': '本塁打になる球。打者が長打を打ちやすい甘い投球のこと。'},
    {'ja': 'ベースヒット',      'en': 'base hit',          'def': '安打。打者が野手に処理されることなく塁に出ること。ヒット参照。'},
    {'ja': 'ノーツー',          'en': 'no-two / 0-2',      'def': 'ノーボール・ツーストライクのカウント。投手有利のカウント。'},
    {'ja': 'キーストンコンビ',  'en': 'keystone combination', 'def': '二塁手と遊撃手のコンビのこと。ダブルプレーの要となる。= keystone duo'},
    {'ja': 'ダブルプレーコンビ','en': 'double play combination', 'def': 'キーストンコンビ参照。二塁手と遊撃手のこと。'},
    {'ja': 'イレギュラーバウンド','en': 'irregular bounce / bad hop', 'def': '予想外の方向に跳ねる打球。グラウンドの状態や石などにより球が不規則に変化すること。'},
    {'ja': 'ロードゲーム',      'en': 'road game / away game', 'def': '敵地での試合。ホームゲームの対義語。'},
    {'ja': 'ワインドアップ',    'en': 'windup',            'def': 'ワインドアップポジションからの投球動作。走者なしのとき主に使われる投球フォーム。'},
    {'ja': 'ジャッジ',          'en': 'judgment / judge',  'def': '審判の判定のこと。'},
    {'ja': 'ジーエム',          'en': 'GM / general manager', 'def': 'ゼネラルマネージャーの略称。球団の編成責任者。'},
    {'ja': 'ラウンダーズ',      'en': 'rounders',          'def': '野球の前身ともいわれるイギリスの球技。タウンボールとも関連する。'},
    {'ja': 'ウェイバー',        'en': 'waiver',            'def': '選手の放出手続き。ウェーバー方式参照。'},
    {'ja': 'サードベースマン',  'en': 'third baseman',     'def': '三塁手。サード参照。'},
    {'ja': '送球ミス',          'en': 'throwing error',    'def': 'スローイングエラー。野手の悪送球によるエラー。暴投（wild throw）ともいう。'},
    {'ja': '球場整備員',        'en': 'groundskeeper',     'def': 'グラウンドキーパー。グラウンドの整備を担当するスタッフ。'},
    {'ja': '軸足',              'en': 'pivot foot',        'def': '投手板に接する足。右投手は右足、左投手は左足。投手の軸足参照。'},
    {'ja': 'セットアップポジション', 'en': 'set position', 'def': 'セットポジション。走者がいるときに投手がとる投球前の静止姿勢。ワインドアップポジションの対義語。'},
    # 残り24件対応
    {'ja': 'ヘッドスラディング',    'en': 'headfirst sliding',    'def': 'ヘッドスライディングの別表記。頭から滑り込むスライディング。'},
    {'ja': 'テキサスレンジャーズ',  'en': 'Texas Rangers',         'def': 'テキサス州アーリントンを本拠地とするMLBチーム。'},
    {'ja': 'ウエーティングサークル','en': 'on-deck circle',        'def': '次打者が待機するサークル。ウェイティングサークルともいう。cf. on-deck circle'},
    {'ja': 'ウェイティングサークル','en': 'on-deck circle',        'def': 'ウエーティングサークルの別表記。次打者が準備する円形の場所。cf. on-deck circle'},
    {'ja': 'ウェーバー方式',        'en': 'waiver system',         'def': 'ウェイバー参照。球団が選手を放出する際の優先順位制度。'},
    {'ja': 'ウェイティング',        'en': 'on-deck / waiting',     'def': '次打者が準備すること。ウェイティングサークルで待機する行為。'},
    {'ja': 'アナハイムエンジェルス','en': 'Anaheim Angels',         'def': 'カリフォルニア州アナハイムを本拠地とするMLBチームの旧称。現在はロサンジェルス・エンゼルス。'},
    {'ja': 'カンバス',              'en': 'canvas base',           'def': 'キャンバス（帆布）製のベース。塁を指す。'},
    {'ja': 'クリーンアップトリオ',  'en': 'cleanup trio',          'def': '3・4・5番打者のこと。中軸打者3人の総称。'},
    {'ja': 'ゴールデングローブ賞',  'en': 'Gold Glove Award',      'def': '各ポジションで最も優れた守備の選手に贈られる年間表彰。ゴールドグラブ賞参照。'},
    {'ja': 'ローレイズリリーフ投手賞','en': 'Rolaids Relief Man Award','def': 'シーズン最優秀救援投手に贈られた賞。現在はMLBではトレバー・ホフマン賞として継続。'},
    {'ja': 'スクリュウーボール',    'en': 'screwball',             'def': 'スクリューボールの別表記。逆カーブともいう変化球。'},
    {'ja': '日本シリーズ',          'en': 'Japan Series',          'def': '日本プロ野球の日本一を決めるシリーズ。セ・リーグとパ・リーグの優勝チームが対戦する。'},
    {'ja': 'ダブルプレーシフト',    'en': 'double play shift',     'def': 'ダブルプレーを取りやすいように内野手が通常より浅めに守る守備シフト。'},
    {'ja': 'サークル',              'en': 'on-deck circle',        'def': 'ネクストバッターズサークルのこと。次打者が準備する円形の場所。'},
    {'ja': '投手版',                'en': "pitcher's plate",       'def': '投手板（ピッチャーズプレート）の別表記。投手が投球時に踏む板。'},
    {'ja': 'フィールダーズチョイス','en': "fielder's choice",      'def': 'フィールダーズチョイス。野手選択。打球を処理する際に野手がアウトにする走者を選択すること。cf. fielder\'s choice'},
    {'ja': 'フィルダースチョイス',  'en': "fielder's choice",      'def': 'フィールダーズチョイスの別表記。野手選択。cf. fielder\'s choice'},
    {'ja': 'ウィルソン',            'en': 'Wilson',                'def': '野球の道具・人名に関連するウィルソン。スポーツ用品メーカーまたは関係者の名称。'},
    {'ja': 'タウンボール',          'en': 'town ball',             'def': '野球の前身のひとつ。ラウンダーズに似たアメリカの球技。'},
    {'ja': 'テキサスリーガーズヒット', 'en': 'Texas Leaguer / Texas League single', 'def': 'テキサスヒットの正式名称。内野手と外野手の中間にぽとりと落ちる安打。テキサスヒット参照。'},
    {'ja': 'テキサスリーガー',      'en': 'Texas Leaguer',         'def': 'テキサスリーガーズヒットの略。テキサスヒット参照。'},
]
_je_missing_seen = set()

JE_NEW_RULES = [
    {'ja': 'ABSチャレンジ', 'yomi': 'エービーエスチャレンジ', 'en': 'ABS Challenge System',                  'def': '選手がストライク・ボールの判定に異議を申し立て、電子ストライクゾーンシステムによる再審査を要求できる制度。各チームに1試合あたり一定回数のチャレンジ権が与えられ、成功した場合は権利が返還される。球審の判定を完全に自動化するのではなく、人間の審判を主体としつつ誤審を救済するハイブリッド方式を採用。マイナーリーグでの試験運用を経て、2026年シーズンからMLBで正式に導入された。'},
    {'ja': 'ビデオ判定拡大',       'en': 'Expanded Instant Replay',               'def': '審判の判定に対し、監督がビデオ検証（チャレンジ）を要求できる制度。ニューヨークのリプレイ・オペレーション・センターで専任スタッフが映像を確認し、判定を覆すかどうかを決定する。ホームラン判定、タッグプレー、フォースプレー、フェア・ファウルなど多くのプレーが対象となる。ボール・ストライクの判定は対象外。'},
    {'ja': '本塁衝突禁止',         'en': 'Home Plate Collision Rule',             'def': '走者が捕手に故意に体当たりすることを原則禁止。捕手側もボールを持っていない状態で走路をブロックすることが禁止される。違反した場合、走者はアウト、あるいは捕手の妨害としてセーフの判定が下される。捕手の負傷が相次いだことを受けて2014年に導入された。'},
    {'ja': '危険なスライディング禁止', 'en': 'Slide Rule (Takeout Slide Ban)',     'def': '走者は塁に到達することを目的とした正当なスライディングをしなければならない。ゲッツー崩しを目的とした体当たりや故意の接触は禁止。違反した場合は走者アウトに加え、打者走者も併殺となる。2016年のポストシーズンで物議を醸したタックルプレーが導入の直接的なきっかけとなった。'},
    {'ja': '故意四球（申告制）',    'en': 'Automatic Intentional Walk',           'def': '監督がベンチから審判に申告するだけで故意四球が成立し、投手が4球を実際に投げる必要がなくなった。それまでは4球を投げる間に暴投などのアクシデントが稀に起きることもあった。試合時間短縮策の一環として2017年に導入された。'},
    {'ja': '最低3人打者ルール',     'en': 'Three-Batter Minimum Rule',            'def': 'リリーフ投手は最低3人の打者に対して投げるか、イニングが終了するまで交代できない。左打者一人だけに登板してすぐ交代する「ワンポイントリリーフ」が試合のテンポを損なうとして廃止された。負傷による交代はこのルールの例外とされる。2020年に導入され、ブルペンの起用戦略に大きな変化をもたらした。'},
    {'ja': 'タイブレーク制',        'en': 'Extra-Innings Tiebreaker Rule',        'def': '延長戦の各回開始時に、直前の打順の打者が二塁走者として自動的に置かれる。得点が入りやすくなるため試合の長時間化を防ぎ、投手の消耗も抑えられる。2020年にコロナ禍の特別ルールとして導入され、試合時間短縮の効果が認められ2023年に正式採用となった。'},
    {'ja': '粘着物質取締り',        'en': 'Foreign Substance Crackdown',          'def': '投手がボールにスパイダータックなどの粘着物質を塗り、回転数を人工的に高める行為への取り締まりを強化。審判が試合中に投手の手やグラブを検査でき、違反が発覚した場合は即時退場および10試合の出場停止処分となる。打者との極端な不均衡が問題視されたことを受けて2021年から本格的に強化された。なお松脂（ロジン）の使用は従来通り認められている。'},
    {'ja': 'ユニバーサルDH',        'en': 'Universal Designated Hitter',          'def': '指名打者制が両リーグに適用され、投手が打席に立つことがなくなった。アメリカンリーグは1973年から採用していたが、ナショナルリーグは長年投手が打つ伝統を守り続けていた。2020年にコロナ禍の特別措置として一時導入され、2022年から正式に両リーグ統一採用となった。'},
    {'ja': '大谷ルール',            'en': 'Ohtani Rule (Two-Way Player Exception)', 'def': 'DHを兼任している投手が、マウンドを降りた後も指名打者としてその試合に打席に立ち続けられる特例ルール。通常、投手が降板するとDHも消滅するが、この特例により二刀流選手の継続出場が可能になった。2022年のユニバーサルDH導入と同時に設けられ、事実上大谷翔平のために作られたルールとして「大谷ルール」と呼ばれる。'},
    {'ja': 'ピッチクロック',        'en': 'Pitch Clock',                          'def': '投手は走者なしで15秒以内、走者ありで20秒以内に投球動作を開始しなければならない。打者側にも投球の8秒前までに構えを完了する義務があり、違反するとボールまたはストライクが自動宣告される。2023年の導入初年度に平均試合時間が約30分短縮されるという劇的な効果をもたらした。'},
    {'ja': '守備シフト制限',        'en': 'Shift Restriction',                    'def': '投球時に内野手4人全員がダート（土の部分）に位置し、二塁ベースを挟んで左右に2人ずつ配置しなければならない。極端なシフトが安打数の減少や得点力低下につながるとして批判され、2023年に禁止された。外野手の配置については現時点では制限がない。'},
    {'ja': 'ベース大型化',          'en': 'Larger Bases',                         'def': '一・二・三塁のベースが従来の15インチ四方から18インチ四方に拡大された。接触事故の減少が主な目的だが、ベースが大きくなることで一塁到達距離がわずかに短くなり、内野安打や盗塁の成功率が若干上がる効果もある。ピッチクロックや牽制制限との相乗効果で、2023年以降に盗塁数が大幅に増加した。'},
    {'ja': '牽制回数制限',          'en': 'Disengagement Limit',                  'def': '投手が1打者との対戦中に牽制球を投げる、またはプレートを外す行為は合計2回までに制限される。3回目で走者をアウトにできなかった場合は自動的にボークとなり、全走者が1つ進塁する。ピッチクロックと連動した制度であり、2023年の導入後、盗塁数が大幅に増加した一因とされている。'},
]
_je_seen = {e['ja'] for e in je_all}
# DOCXの「オークランド・アスレティックス」は_OLD_TEAMS_JEの「アスレチックス」と重複するため除外
je_all = [e for e in je_all if e.get('ja') != 'オークランド・アスレティックス']
_je_seen.discard('オークランド・アスレティックス')
for r in JE_MISSING_REFS:
    if r['ja'] not in _je_seen:
        _je_seen.add(r['ja'])
        je_all.append(r)
for r in JE_NEW_RULES:
    if r['ja'] not in _je_seen:
        _je_seen.add(r['ja'])
        je_all.append(r)

# 球場名の和英エントリ
_STADIUMS_JE = [
    {'ja': 'オリオール・パーク・アット・カムデン・ヤーズ', 'en': 'Oriole Park at Camden Yards', 'def': 'ボルティモア・オリオールズの本拠地。メリーランド州ボルティモア。'},
    {'ja': 'フェンウェイ・パーク',         'en': 'Fenway Park',                  'def': 'ボストン・レッドソックスの本拠地。マサチューセッツ州ボストン。1912年開場。グリーン・モンスター（左翼フェンス）で有名。'},
    {'ja': 'ヤンキー・スタジアム',         'en': 'Yankee Stadium',               'def': 'ニューヨーク・ヤンキースの本拠地。ニューヨーク州ブロンクス。2009年開場。'},
    {'ja': 'トロピカーナ・フィールド',     'en': 'Tropicana Field',              'def': 'タンパベイ・レイズの本拠地。フロリダ州セントピーターズバーグ。屋根付きドーム球場。'},
    {'ja': 'ロジャーズ・センター',         'en': 'Rogers Centre',                'def': 'トロント・ブルージェイズの本拠地。カナダ・オンタリオ州トロント。開閉式屋根を持つ。'},
    {'ja': 'ギャランティード・レート・フィールド', 'en': 'Guaranteed Rate Field', 'def': 'シカゴ・ホワイトソックスの本拠地。イリノイ州シカゴ。'},
    {'ja': 'プログレッシブ・フィールド',   'en': 'Progressive Field',            'def': 'クリーブランド・ガーディアンズの本拠地。オハイオ州クリーブランド。'},
    {'ja': 'コメリカ・パーク',             'en': 'Comerica Park',                'def': 'デトロイト・タイガースの本拠地。ミシガン州デトロイト。'},
    {'ja': 'カウフマン・スタジアム',       'en': 'Kauffman Stadium',             'def': 'カンザスシティ・ロイヤルズの本拠地。ミズーリ州カンザスシティ。'},
    {'ja': 'ターゲット・フィールド',       'en': 'Target Field',                 'def': 'ミネソタ・ツインズの本拠地。ミネソタ州ミネアポリス。'},
    {'ja': 'ミニット・メイド・パーク',     'en': 'Minute Maid Park',             'def': 'ヒューストン・アストロズの本拠地。テキサス州ヒューストン。開閉式屋根を持つ。'},
    {'ja': 'エンジェル・スタジアム',       'en': 'Angel Stadium',                'def': 'ロサンゼルス・エンゼルスの本拠地。カリフォルニア州アナハイム。'},
    {'ja': 'サター・ヘルス・パーク',       'en': 'Sutter Health Park',           'def': 'オークランド・アスレティックスが新球場建設中に使用する暫定本拠地。カリフォルニア州サクラメント。'},
    {'ja': 'T-モバイル・パーク',           'en': 'T-Mobile Park',                'def': 'シアトル・マリナーズの本拠地。ワシントン州シアトル。開閉式屋根を持つ。'},
    {'ja': 'グローブ・ライフ・フィールド', 'en': 'Globe Life Field',             'def': 'テキサス・レンジャーズの本拠地。テキサス州アーリントン。2020年開場の開閉式屋根付き球場。'},
    {'ja': 'トゥルーイスト・パーク',       'en': 'Truist Park',                  'def': 'アトランタ・ブレーブスの本拠地。ジョージア州カンバーランド。'},
    {'ja': 'ローンデポ・パーク',           'en': 'loanDepot park',               'def': 'マイアミ・マーリンズの本拠地。フロリダ州マイアミ。開閉式屋根を持つ。'},
    {'ja': 'シティ・フィールド',           'en': 'Citi Field',                   'def': 'ニューヨーク・メッツの本拠地。ニューヨーク州クイーンズ。2009年開場。'},
    {'ja': 'シティズンズ・バンク・パーク', 'en': 'Citizens Bank Park',           'def': 'フィラデルフィア・フィリーズの本拠地。ペンシルベニア州フィラデルフィア。'},
    {'ja': 'ナショナルズ・パーク',         'en': 'Nationals Park',               'def': 'ワシントン・ナショナルズの本拠地。ワシントンD.C.。'},
    {'ja': 'リグレー・フィールド',         'en': 'Wrigley Field',                'def': 'シカゴ・カブスの本拠地。イリノイ州シカゴ。1914年開場の歴史的球場。ツタに覆われた外野フェンスで有名。'},
    {'ja': 'グレート・アメリカン・ボールパーク', 'en': 'Great American Ball Park', 'def': 'シンシナティ・レッズの本拠地。オハイオ州シンシナティ。'},
    {'ja': 'アメリカン・ファミリー・フィールド', 'en': 'American Family Field',   'def': 'ミルウォーキー・ブルワーズの本拠地。ウィスコンシン州ミルウォーキー。開閉式屋根を持つ。'},
    {'ja': 'PNCパーク',                    'en': 'PNC Park',                     'def': 'ピッツバーグ・パイレーツの本拠地。ペンシルベニア州ピッツバーグ。アレゲニー川沿いに立地する美しい球場として名高い。'},
    {'ja': 'ブッシュ・スタジアム',         'en': 'Busch Stadium',                'def': 'セントルイス・カーディナルスの本拠地。ミズーリ州セントルイス。'},
    {'ja': 'チェース・フィールド',         'en': 'Chase Field',                  'def': 'アリゾナ・ダイヤモンドバックスの本拠地。アリゾナ州フェニックス。開閉式屋根と天然芝を持つ。'},
    {'ja': 'クアーズ・フィールド',         'en': 'Coors Field',                  'def': 'コロラド・ロッキーズの本拠地。コロラド州デンバー。標高1600m超の高地にあり打球が飛びやすい球場として有名。'},
    {'ja': 'ドジャー・スタジアム',         'en': 'Dodger Stadium',               'def': 'ロサンゼルス・ドジャースの本拠地。カリフォルニア州ロサンゼルス。1962年開場。MLB最大収容人数を誇る。'},
    {'ja': 'ペトコ・パーク',               'en': 'Petco Park',                   'def': 'サンディエゴ・パドレスの本拠地。カリフォルニア州サンディエゴ。'},
    {'ja': 'オラクル・パーク',             'en': 'Oracle Park',                  'def': 'サンフランシスコ・ジャイアンツの本拠地。カリフォルニア州サンフランシスコ。マッコビー・コーブに面し、本塁打が湾に飛び込むことで有名。'},
]
for r in _STADIUMS_JE:
    if r['ja'] not in _je_seen:
        _je_seen.add(r['ja'])
        je_all.append(r)

# 旧球団名エントリ（JE）
_OLD_TEAMS_JE = [
    {'ja': 'モントリオール・エクスポズ',           'en': 'Montreal Expos',          'def': 'ナショナル・リーグ東地区。1969年創設。カナダ・モントリオール本拠地。（追記）現在は無い。2005年にワシントンへ移転しワシントン・ナショナルズとなった。'},
    {'ja': 'フロリダ・マーリンズ',                 'en': 'Florida Marlins',         'def': 'ナショナル・リーグ東地区。1993年創設。フロリダ州マイアミ本拠地。（追記）現在は無い。2012年にマイアミ・マーリンズに改称。'},
    {'ja': 'タンパベイ・デビルレイズ',             'en': 'Tampa Bay Devil Rays',    'def': 'アメリカン・リーグ東地区。1998年創設。フロリダ州セントピーターズバーグ本拠地。（追記）現在は無い。2008年にタンパベイ・レイズに改称。'},
    {'ja': 'クリーブランド・インディアンス',       'en': 'Cleveland Indians',       'def': 'アメリカン・リーグ中地区。1901年創設。オハイオ州クリーブランド本拠地。（追記）現在は無い。2022年にクリーブランド・ガーディアンズに改称。'},
    {'ja': 'アナハイム・エンゼルス',               'en': 'Anaheim Angels',          'def': 'アメリカン・リーグ西地区。1961年創設。カリフォルニア州アナハイム本拠地。（追記）現在は無い。2005年にロサンゼルス・エンゼルス・オブ・アナハイムに改称、2016年にロサンゼルス・エンゼルスに改称。'},
    {'ja': 'ロサンゼルス・エンゼルス・オブ・アナハイム', 'en': 'Los Angeles Angels of Anaheim', 'def': 'アメリカン・リーグ西地区。アナハイム・エンゼルスが2005年に改称。（追記）現在は無い。2016年にロサンゼルス・エンゼルスに改称。'},
    {'ja': 'オークランド・アスレチックス',         'en': 'Oakland Athletics',       'def': 'アメリカン・リーグ西地区。1901年創設（旧フィラデルフィア・アスレチックス）。（追記）2024年シーズン終了をもってオークランドを撤退。2025年よりアスレチックスに改称、サクラメントを暫定本拠地とする。\ncf. Oakland Athletics (the)'},
    {'ja': 'アスレチックス',                       'en': 'Athletics',               'def': 'アメリカン・リーグ西地区。オークランド・アスレチックスが2025年に改称。サクラメントを暫定本拠地とし、ラスベガスに新球場を建設中（2028年開場予定）。'},
]
for r in _OLD_TEAMS_JE:
    if r['ja'] not in _je_seen:
        _je_seen.add(r['ja'])
        je_all.append(r)

# 漢字・ラテン文字で始まるエントリーにyomi（カタカナ読み）を付与（JE_NEW_RULES追加後）
print("Adding yomi to kanji entries …", file=sys.stderr)
for e in je_all:
    if e.get('yomi'):
        continue  # 手動設定済みはスキップ
    ch = e['ja'][0]
    if not (0x30A1 <= ord(ch) <= 0x30F6) and not (0x3041 <= ord(ch) <= 0x3096):
        e['yomi'] = to_yomi(e['ja'])

je_all.sort(key=je_sort_key)

# テキサスヒット：末尾の「テキサスリーガーズヒット=テキサス・ヒット参照。」を削除
for e in je_all:
    if e['ja'] == 'テキサスヒット' and e.get('def'):
        e['def'] = re.sub(r'テキサスリーガーズヒット\s*=\s*テキサス・?ヒット参照。?', '', e['def']).strip()
        break

# アルプス・スタンド：定義文を修正
for e in je_all:
    if e['ja'] == 'アルプス・スタンド':
        e['def'] = '阪神球団の本拠地、甲子園球場の内野席と外野席の間に設けられた大観覧席のこと。高くそびえているように見えることからこの名がつけられたのであろうが、英語にはこのような表現はない。なお、スタンドの最上段は英語でupper deckと言う。'
        break

# メジャーリーグ：「メジャーリーグ参照」→「マイナーリーグ参照」に修正
for e in je_all:
    if e['ja'] == 'メジャーリーグ' and e.get('def'):
        e['def'] = e['def'].replace('メジャーリーグ参照', 'マイナーリーグ参照')

# スパイクシューズ：ブラケット除去後に「だけでよい」→「だけでもよい」
for e in je_all:
    if e['ja'] == 'スパイクシューズ' and e.get('def'):
        e['def'] = e['def'].replace('だけでよい。', 'だけでもよい。')

# ワンポイントリリーフ：参考文献リストが混入しているので除去（（追記）は保持）
for e in je_all:
    if e['ja'] == 'ワンポイントリリーフ' and e.get('def'):
        idx = e['def'].find('主要参考文献')
        if idx >= 0:
            before = e['def'][:idx].rstrip('。 \t\n') + '。'
            addendum_idx = e['def'].find('（追記）')
            if addendum_idx >= 0:
                e['def'] = before + '\n' + e['def'][addendum_idx:]
            else:
                e['def'] = before

# 「とか」→「または」（全エントリー）
# ただし「とか」の後が「ら」「も」「な」「く」の場合は除外（ことから・とかく等の誤置換防止）
_toka_re = re.compile(r'とか(?![らもなく])')
for e in ej_all:
    e['ja'] = _toka_re.sub('または', e['ja'])
for e in je_all:
    e['en'] = _toka_re.sub('または', e['en'])
    if e.get('def'):
        e['def'] = _toka_re.sub('または', e['def'])

# 「1. 」「２．　」など番号＋点の後ろのスペースを除去
_num_space = re.compile(r'([0-9０-９]+[.．])[ \t　]+')
# 番号項目の前に改行を挿入：「。2．」や「 2．」（スペース後の全角ピリオド付き番号）
_num_newline = re.compile(r'(?<!\n)(?:[。．][ \t　]*|[ \t　]+)([0-9０-９]+[．])')
def fmt_numbered(s):
    if not s: return s
    s = _num_space.sub(r'\1', s)
    s = _num_newline.sub(r'\n\1', s)
    return s
for e in ej_all:
    e['ja'] = fmt_numbered(e['ja'])
for e in je_all:
    e['en'] = fmt_numbered(e['en'])
    if e.get('def'):
        e['def'] = fmt_numbered(e['def'])

# ── 解決不可能な参照を削除 ──────────────────────────────────────────────────────
# JEマップとEJマップを構築して参照先が存在するか確認
_valid_ja = set()
for e in je_all:
    _valid_ja.add(e['ja'])
    _valid_ja.add(e['ja'].replace('・', ''))
_valid_en = set()
for e in ej_all:
    k = e['en'].lower()
    _valid_en.add(k)
    for p in (k.split('/') if '/' in k else k.split(', ') if ', ' in k else []):
        if p.strip(): _valid_en.add(p.strip())

def _ja_ref_ok(raw):
    """正規化後にJEまたはEJに存在するか確認"""
    t = re.sub(r'^詳細は', '', raw)
    t = re.sub(r'^combination', '', t)
    t = re.sub(r'[をはがのへでとにもや]$|から$|まで$|より$', '', t)
    t = re.sub(r'の項目$', '', t)
    t = re.sub(r'の項$', '', t)
    t = re.sub(r'試合$', '', t)
    t = re.sub(r'[をはがのへでとにもや]$|から$|まで$|より$', '', t)
    t = t.replace('・', '').strip()
    return (t in _valid_ja or raw in _valid_ja
            or t.lower() in _valid_en or t.lower().replace('-', ' ') in _valid_en)

def _en_ref_ok(raw):
    """EJに存在するか確認"""
    k = raw.strip().lower()
    alt = k.replace('-', ' ') if '-' in k else k.replace(' ', '-')
    return (k in _valid_en or alt in _valid_en
            or k + ' (the)' in _valid_en or alt + ' (the)' in _valid_en
            or k + ' (a)' in _valid_en)

# JEの参照テキスト削除：「TERM参照」パターン
_re_ja_ref = re.compile(
    r'[「｢]([^」｣]+)[」｣]参照'                   # 「〇〇」参照
    r'|([ぁ-んァ-ヶー一-龯・（）a-zA-Z0-9]+(?:とノーゲーム)?)'  # 〇〇参照
    r'参照'
)
# EJのcf.テキスト削除用パターン（EJの説明文中）
_re_cf = re.compile(r'cf\.\s*([A-Za-z0-9][A-Za-z0-9 \-\/\(\)\']*)')

def _strip_bad_ja_refs(text):
    """解決不可能な参照を削除（「TERM参照」ごと除去）"""
    if not text: return text
    def repl(m):
        raw = (m.group(1) or m.group(2) or '').strip()
        if not raw or _ja_ref_ok(raw):
            return m.group(0)  # 解決可能→そのまま
        # 解決不可能→削除（前後の読点・スペースも整理）
        return ''
    result = _re_ja_ref.sub(repl, text)
    # 削除後に残る「。参照」「、参照」などのゴミを除去
    result = re.sub(r'[、。]\s*$', '。', result.strip())
    result = re.sub(r'\s{2,}', ' ', result)
    return result

_CF_FORCE_REMOVE = {'handy man'}  # 辞書から削除したエントリのcf.参照も除去

def _strip_bad_cf(text):
    """解決不可能なcf.参照を削除"""
    if not text: return text
    def repl(m):
        raw = m.group(1).strip()
        if raw.lower() in _CF_FORCE_REMOVE:
            return ''
        if _en_ref_ok(raw):
            return m.group(0)
        return ''
    result = _re_cf.sub(repl, text)
    result = re.sub(r'cf\.\s*,\s*', 'cf. ', result)  # cf. が残った場合の先頭カンマ除去
    result = re.sub(r',\s*,', ',', result)
    result = re.sub(r'cf\.\s*$', '', result).strip()
    # cf.が消えて先頭にカンマだけ残るケース（例: 「：, role player」→「：cf. role player」）
    result = re.sub(r'(：|:)\s*,\s*', r'\1 cf. ', result)
    # スペース・句読点の整理
    result = re.sub(r'\s{2,}', ' ', result).strip()
    return result

_removed_refs = 0
for e in je_all:
    for field in ('en', 'def'):
        orig = e.get(field, '')
        if orig:
            cleaned = _strip_bad_ja_refs(orig)
            if cleaned != orig:
                e[field] = cleaned
                _removed_refs += 1
for e in ej_all:
    orig = e.get('ja', '')
    if orig:
        cleaned = _strip_bad_cf(orig)
        if cleaned != orig:
            e['ja'] = cleaned
            _removed_refs += 1

print(f"Removed {_removed_refs} unresolvable refs", file=sys.stderr)

# 二重句点の修正（addendum適用時のrstrip漏れを吸収）
for e in ej_all:
    if e.get('ja') and '。。' in e['ja']:
        e['ja'] = e['ja'].replace('。。', '。')
for e in je_all:
    for f in ('def', 'en'):
        if e.get(f) and '。。' in e[f]:
            e[f] = e[f].replace('。。', '。')
# ─────────────────────────────────────────────────────────────────────────────

# ── 拡充データ（見出し/短文/長文/例/関連語）のマージ ──────────────────────────
# dict_enrich.json があれば en(小文字) で EJ にマージ。new:true は新規エントリ追加。
_enrich_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'dict_enrich.json')
if os.path.exists(_enrich_path):
    with open(_enrich_path, encoding='utf-8') as _f:
        _enrich_entries = json.load(_f).get('entries', [])
    _ej_by_en = {}
    for _e in ej_all:
        _k = (_e.get('en') or '').strip().lower()
        if _k and _k not in _ej_by_en:
            _ej_by_en[_k] = _e
    _added = _merged = 0
    for _en in _enrich_entries:
        _k = _en['en'].strip().lower()
        _tgt = _ej_by_en.get(_k)
        if _tgt is None:
            if _en.get('new'):
                _tgt = {'en': _en['en'], 'ja': _en.get('short', '')}
                ej_all.append(_tgt); _ej_by_en[_k] = _tgt; _added += 1
            else:
                print(f"  enrich: '{_en['en']}' が辞書に無く new でもないためスキップ", file=sys.stderr)
                continue
        _tgt['short'] = _en.get('short', '')
        _tgt['long'] = _en.get('long', '')
        if _en.get('example'):
            _tgt['ex'] = _en['example']
        _tgt['rel'] = _en.get('related', [])
        if not _tgt.get('ja'):
            _tgt['ja'] = _en.get('short', '')
        _merged += 1
    # 新規追加でソート順が乱れるので再ソート（英和は sk/en 基準）
    ej_all.sort(key=lambda x: (x.get('sk') or x['en']).lower())
    print(f"Enrich: merged {_merged}, added {_added} new", file=sys.stderr)

    # ── 和英（JE）拡充のマージ（ja で照合。short=英訳, long=説明, related=和英見出し） ──
    with open(_enrich_path, encoding='utf-8') as _f:
        _je_enrich = json.load(_f).get('je_entries', [])
    if _je_enrich:
        _je_by_ja = {}
        for _e in je_all:
            _k = (_e.get('ja') or '').strip()
            if _k and _k not in _je_by_ja:
                _je_by_ja[_k] = _e
        _jm = 0
        for _je in _je_enrich:
            _tgt = _je_by_ja.get(_je['ja'].strip())
            if _tgt is None:
                print(f"  je-enrich: '{_je['ja']}' が和英に無くスキップ", file=sys.stderr)
                continue
            _tgt['short'] = _je.get('short', '')
            _tgt['long'] = _je.get('long', '')
            if _je.get('example'):
                _tgt['ex'] = _je['example']
            _tgt['rel'] = _je.get('related', [])
            _jm += 1
        print(f"JE-Enrich: merged {_jm}", file=sys.stderr)

# スナップショット使用時は、ここで完成データ（JSON）を最終データとして採用（上流の結果を上書き）
if _USE_SNAP:
    with open(_EJ_SNAP, encoding='utf-8') as _f: ej_all = json.load(_f)
    with open(_JE_SNAP, encoding='utf-8') as _f: je_all = json.load(_f)

EJ_COUNT = len(ej_all)
JE_COUNT = len(je_all)
print(f"Final → EJ: {EJ_COUNT},  JE: {JE_COUNT}", file=sys.stderr)

EJ_JSON = json.dumps(ej_all, ensure_ascii=False)
JE_JSON = json.dumps(je_all, ensure_ascii=False)

# 今日の豆知識（普遍的なルール・記録。事実確認しやすいものだけ）
TRIVIA = [
    "1試合は9イニング。9回終了時に同点なら決着がつくまで延長戦を行う。",
    "打者は3ストライクで三振、4ボールで四球（フォアボール）となり一塁へ進む。",
    "ワールドシリーズは7回戦制（4勝を先取したチームが優勝）。",
    "完全試合（パーフェクトゲーム）は、27人の打者を1人も出塁させずに抑える快挙。",
    "ノーヒットノーランは、無安打無得点に抑えること。四球や失策で走者は出てもよい。",
    "サイクルヒットは、1試合で単打・二塁打・三塁打・本塁打をすべて打つこと。",
    "グランドスラム（満塁本塁打）は一度に4点が入る、最も得点効率のよい一打。",
    "インフィールドフライは、無死または1死で走者が一・二塁か満塁のとき、内野への平凡なフライで打者が自動的にアウトになるルール。",
    "ボークは走者がいるときの投手の反則投球動作。宣告されると全走者が1つ進塁する。",
    "タッチアップは、フライが捕球された後に走者が塁に触れ直してから進塁するプレー。",
    "サイ・ヤング賞は、その年の最も優れた投手に贈られる賞（1956年創設）。名は通算511勝の大投手サイ・ヤングにちなむ。",
    "MVPは最優秀選手賞。両リーグそれぞれ1人ずつ選ばれる。",
    "サイ・ヤングの通算511勝はMLB史上最多で、今後破られないと言われる記録。",
    "ジョー・ディマジオの56試合連続安打（1941年）は、不滅の大記録とされる。",
    "バリー・ボンズの通算762本塁打はMLB歴代最多。シーズン73本塁打（2001年）も最多記録。",
    "ピート・ローズの通算4256安打はMLB歴代最多。",
    "リッキー・ヘンダーソンの通算1406盗塁はMLB歴代最多。",
    "2022年から両リーグで指名打者（DH）制を採用（ユニバーサルDH）。投手は打席に立たなくなった。",
    "指名打者（DH）は投手に代わって打撃のみを行い、守備にはつかない。",
    "2023年からピッチクロックを導入。走者なしで15秒、走者ありで20秒以内に投球する。",
    "2023年から塁ベースが一辺15インチから18インチに拡大され、盗塁が増える一因となった。",
    "2023年から極端な守備シフトが制限され、内野手は内野の土の上に2人ずつ配置しなければならない。",
    "延長戦のタイブレーク制では、10回以降は無死・走者二塁から攻撃を始める。",
    "振り逃げは、2ストライク後の第3ストライクを捕手が正規に捕球できなかったとき、打者が一塁へ走れるプレー。",
    "ストライクゾーンは、ホームベース上の空間で、打者のひざ頭の下部から肩とズボン上部の中間まで。",
    "セーブは、リードを守って試合を締めくくった救援投手に記録される。",
    "打率は「安打÷打数」、防御率（ERA）は「自責点×9÷投球回」で計算する。",
    "OPSは出塁率（OBP）と長打率（SLG）を足した打撃指標で、総合力の目安になる。",
    "大谷翔平は2024年、史上初となる「1シーズン50本塁打・50盗塁」を達成した。",
    "ダブルヘッダーは同じ2チームが1日に2試合続けて行うこと。",
]
TRIVIA_JSON = json.dumps(TRIVIA, ensure_ascii=False)

# ── HTML ──────────────────────────────────────────────────────────────────────

HTML = r"""<!DOCTYPE html>
<html lang="ja">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1,viewport-fit=cover">
<meta name="apple-mobile-web-app-capable" content="yes">
<meta name="apple-mobile-web-app-status-bar-style" content="default">
<meta name="mobile-web-app-capable" content="yes">
<title>MLB野球用語辞典</title>
<style>
/* ── Design tokens (pop) ── */
:root{
  --c-primary:#3a3a5c;
  --c-accent:#8ad3f7;
  --c-coral:#ff5a4e;
  --c-orange:#f8bd6b;
  --c-ink:#1e1e2e;--c-muted:#6b6b80;--c-muted2:#a8a8bc;
  --c-canvas:#fdf8ef;--c-parchment:#ffffff;--c-card:#edf7ee;--c-hairline:#e8e0d0;
  --c-border:#1e1e2e;
  --bw:2.5px;
  --shadow-s:2px 2px 0 var(--c-border);
  --shadow-m:3px 3px 0 var(--c-border);
  --shadow-l:4px 4px 0 var(--c-border);
  --c-nav-bg:#ffffff;
  --c-idx-bg:rgba(255,255,255,0.95);
  --c-card-bg:#ffffff;
  --c-score-bg:#ffffff;
  --c-ok-bg:#d8f5e2;--c-ok-border:#1e1e2e;--c-ok-text:#156a35;
  --c-ok-mark:#2fbf6b;
  --c-ng-bg:#ffddda;--c-ng-border:#1e1e2e;--c-ng-text:#a02318;
  --c-ng-mark:#ff5a4e;
  --c-fav:#ff6f9c;
  --r-pill:9999px;--r-lg:20px;--r-md:14px;--r-sm:10px;
  --ff-d:"SF Pro Display",system-ui,-apple-system,sans-serif;
  --ff-t:"SF Pro Text",system-ui,-apple-system,sans-serif;
  --safe-b:env(safe-area-inset-bottom,0px);
  --safe-t:env(safe-area-inset-top,0px);
}
/* ── Reset ── */
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
html,body{height:100%;overflow:hidden;-webkit-text-size-adjust:100%}
body{font-family:var(--ff-t);font-size:17px;line-height:1.47;color:var(--c-ink);
  background:#ffffff;-webkit-font-smoothing:antialiased;
  transition:background .25s,color .25s}

/* ── Shell ── */
#app{display:flex;flex-direction:column;height:100%;max-width:600px;margin:0 auto;
  background:var(--c-canvas);position:relative;overflow:hidden}
.tab-panel{display:none;flex:1;flex-direction:column;overflow:hidden;min-height:0}
.tab-panel.active{display:flex}
/* プルダウン更新のインジケータ */
#ptr-ind{position:fixed;left:50%;top:0;transform:translateX(-50%) translateY(-52px);z-index:2000;
  background:var(--c-card);border:2px solid var(--c-border);border-radius:var(--r-pill);
  box-shadow:var(--shadow-s);padding:6px 16px;font-family:var(--ff-t);font-weight:800;font-size:12px;
  color:var(--c-primary);opacity:0;transition:opacity .15s;pointer-events:none;white-space:nowrap}
#ptr-ind.show{opacity:1}

/* ── バナー広告 ── */
#ad-banner{display:none;flex-shrink:0;align-items:center;justify-content:center;gap:8px;
  height:52px;background:var(--c-score-bg);border-top:1px dashed var(--c-border);z-index:99}
#ad-banner.show{display:flex}
.ad-banner-tag{font-family:var(--ff-t);font-size:10px;font-weight:800;color:#fff;
  background:var(--c-muted2);border-radius:var(--r-pill);padding:1px 8px}
.ad-banner-txt{font-family:var(--ff-t);font-size:12px;color:var(--c-muted2)}
/* ── Bottom nav ── */
#bnav{display:flex;background:var(--c-nav-bg);border-top:1px solid var(--c-hairline);
  flex-shrink:0;z-index:100;padding-bottom:var(--safe-b);
  transition:background .25s,border-color .25s}
.bnav-btn{flex:1;display:flex;flex-direction:column;align-items:center;justify-content:center;
  padding:6px 0 4px;border:none;background:transparent;cursor:pointer;
  color:var(--c-muted);font-family:var(--ff-t);font-size:9px;font-weight:600;
  letter-spacing:-0.06px;line-height:1.3;gap:3px;transition:color .15s;
  -webkit-tap-highlight-color:transparent;min-width:0}
/* アイコンは各自の色（--ic）を薄く色付け。選択中は濃く＋ラベルも同色 */
.bnav-btn.active{color:var(--ic);font-weight:800}
.bnav-btn svg{width:24px;height:24px;color:var(--ic,var(--c-muted));
  stroke:currentColor;fill:currentColor;fill-opacity:.16;
  stroke-width:1.9;stroke-linecap:round;stroke-linejoin:round;flex-shrink:0;
  opacity:.55;transition:opacity .15s}
.bnav-btn.active svg{opacity:1;stroke-width:2.1}

/* ── App title ── */
.app-title{padding:14px 16px 2px;text-align:center;flex-shrink:0;
  background:var(--c-canvas);transition:background .25s}
.app-title h1{
  font-family:'Helvetica Neue','Arial Black','Hiragino Sans','Hiragino Kaku Gothic ProN',
    'Yu Gothic','YuGothic','Noto Sans JP',sans-serif;
  font-size:15px;font-weight:900;font-style:normal;
  white-space:nowrap;overflow:hidden;text-overflow:ellipsis;
  letter-spacing:0.8px;line-height:1.2;margin:0;
  color:var(--c-primary);
  display:flex;align-items:center;justify-content:center;gap:7px}
.app-title h1 img{height:26px;width:auto;flex-shrink:0}

/* ── Search ── */
.search-bar{padding:10px 14px;background:var(--c-canvas);
  border-bottom:1px solid var(--c-hairline);flex-shrink:0;
  transition:background .25s,border-color .25s}
.search-wrap{display:flex;align-items:center;background:var(--c-parchment);
  border:var(--bw) solid var(--c-border);box-shadow:var(--shadow-s);
  border-radius:var(--r-pill);padding:9px 14px;gap:8px;transition:background .25s}
.search-wrap svg{width:15px;height:15px;color:var(--c-muted);flex-shrink:0;
  stroke:currentColor;fill:none;stroke-width:2.2;stroke-linecap:round;stroke-linejoin:round}
.search-wrap input{flex:1;border:none;background:transparent;font-family:var(--ff-t);
  font-size:16px;color:var(--c-ink);outline:none;letter-spacing:-0.3px}
.search-wrap input::placeholder{color:var(--c-muted)}
.s-clear{background:none;border:none;cursor:pointer;padding:2px;color:var(--c-muted);
  display:none;flex-shrink:0;line-height:1;font-size:17px;-webkit-tap-highlight-color:transparent}
.s-clear.show{display:block}

/* ── Sub-tabs ── */
.sub-tabs{display:flex;background:var(--c-canvas);border-bottom:1px solid var(--c-hairline);
  flex-shrink:0;gap:8px;padding:8px 14px;transition:background .25s,border-color .25s}
.sub-tab{flex:1;padding:9px;border:var(--bw) solid var(--c-border);background:var(--c-parchment);
  border-radius:var(--r-pill);box-shadow:var(--shadow-s);
  font-family:var(--ff-t);font-size:15px;font-weight:700;color:var(--c-muted);
  cursor:pointer;transition:all .12s;
  -webkit-tap-highlight-color:transparent}
.sub-tab:active{transform:translate(2px,2px);box-shadow:none}
.sub-tab.active{color:var(--c-ink);background:var(--c-accent)}
/* ♥ お気に入りサブタブ（ピンク・下のSVGと同じ丸みハート） */
.sub-tab-fav{flex:0 0 auto;width:52px;padding:6px;display:flex;align-items:center;justify-content:center}
.sub-tab-fav svg{width:23px;height:23px;fill:#ff6f9c;stroke:none;display:block}
.sub-tab-fav.active{background:#ff9cbb}
.sub-tab-fav.active svg{fill:#fff}
/* お気に入り表示中の「辞書に戻る」バー */
.fav-back-bar{display:flex;align-items:center;gap:10px;padding:8px 14px;
  background:var(--c-canvas);border-bottom:1px solid var(--c-hairline)}
.fav-back-title{font-family:var(--ff-d);font-weight:900;font-size:15px;color:var(--c-ink)}

/* ── Dict body ── */
.dict-body{flex:1;display:flex;overflow:hidden;position:relative;min-height:0}
.entry-list{flex:1;overflow-y:auto;-webkit-overflow-scrolling:touch;
  padding-right:28px}
.entry-list::-webkit-scrollbar{display:none}

/* ── Entry ── */
.entry-item{padding:13px 14px;border-bottom:1px solid var(--c-hairline);
  display:flex;align-items:flex-start;gap:10px;cursor:pointer;
  background:var(--c-canvas);transition:background .1s,border-color .25s;
  -webkit-tap-highlight-color:transparent}
.entry-item:active{background:var(--c-parchment)}
.entry-text{flex:1;min-width:0}
.entry-head{font-family:var(--ff-d);font-size:21px;font-weight:600;
  line-height:1.3;letter-spacing:-0.28px;color:var(--c-ink);
  white-space:nowrap;overflow:hidden;text-overflow:ellipsis;padding-bottom:2px}
.entry-badge{font-size:10px;font-weight:700;padding:1px 5px;border-radius:3px;
  vertical-align:middle;margin-left:6px;letter-spacing:0}
.entry-badge.ej{background:var(--c-accent);color:var(--c-ink);border:1.5px solid var(--c-border)}
.entry-badge.je{background:var(--c-orange);color:var(--c-ink);border:1.5px solid var(--c-border)}
.entry-sub{font-family:var(--ff-t);font-size:13px;color:var(--c-muted);
  line-height:1.38;letter-spacing:-0.16px;margin-top:3px;
  overflow:hidden;display:-webkit-box;-webkit-line-clamp:2;-webkit-box-orient:vertical}
.fav-btn{background:transparent;border:none;cursor:pointer;padding:4px;
  color:var(--c-muted);flex-shrink:0;-webkit-tap-highlight-color:transparent;
  transition:color .15s,transform .15s}
.fav-btn.on{color:var(--c-fav)}
.fav-btn:active{transform:scale(.82)}
.fav-btn svg{width:20px;height:20px;stroke:currentColor;
  stroke-width:1.8;stroke-linecap:round;stroke-linejoin:round}
@keyframes fav-remove{0%{opacity:1;transform:translateX(0)}60%{opacity:.3;transform:translateX(18px)}100%{opacity:0;transform:translateX(24px)}}
.fav-removing{animation:fav-remove .35s ease-in forwards;pointer-events:none}

/* ── Index bar ── */
#idx-bar,#abbr-idx-bar{position:absolute;right:0;top:0;bottom:0;width:26px;
  display:flex;flex-direction:column;justify-content:space-between;
  align-items:center;padding:6px 0;z-index:10;
  background:var(--c-idx-bg);backdrop-filter:blur(4px);
  -webkit-backdrop-filter:blur(4px)}
.idx-l{font-family:var(--ff-t);font-size:9px;font-weight:700;
  color:var(--c-primary);cursor:pointer;padding:1px 2px;border-radius:3px;
  transition:all .12s;-webkit-tap-highlight-color:transparent;
  line-height:1.3;min-width:20px;text-align:center;user-select:none}
.idx-l:hover,.idx-l.hit{background:var(--c-primary);color:#fff}
#idx-bar.ja-full .idx-l{font-size:8px;padding:0 1px;min-width:18px;line-height:1.2}
#idx-bar.ja-full .idx-l.row-head{font-size:11px;font-weight:900;color:var(--c-primary)}

/* ── Section header ── */
.sec-hdr{padding:5px 14px 3px;background:var(--c-accent);
  font-family:var(--ff-d);font-size:20px;font-weight:800;
  letter-spacing:.231px;color:var(--c-ink);position:sticky;top:0;
  z-index:5;border-bottom:2px solid var(--c-border);border-top:2px solid var(--c-border);
  transition:background .25s,border-color .25s,color .25s}
.search-mode-hdr{font-size:13px;font-weight:700;letter-spacing:.5px;
  color:var(--c-muted);padding:8px 14px 4px;text-transform:uppercase}

/* ── Bottom sheet modal ── */
#quiz-flash{position:absolute;inset:0;z-index:800;pointer-events:none;
  display:flex;align-items:center;justify-content:center;opacity:0;transition:opacity .12s}
#quiz-flash.show{opacity:1}
#quiz-flash.hide{opacity:0;transition:opacity .3s}
#quiz-flash-mark{font-size:160px;line-height:1;font-weight:900;
  text-shadow:0 0 40px rgba(0,0,0,.15)}
#quiz-flash.ng #quiz-flash-mark{color:var(--c-ng-mark)}
#quiz-flash.ok #quiz-flash-mark{color:var(--c-ok-mark)}
/* 今日の10問ボーナス「＋N点」大きく表示（濃い色プレートでどの背景でもはっきり読める） */
#quiz-flash-bonus{display:none;position:absolute;top:18%;left:0;right:0;margin:0 auto;
  width:max-content;max-width:88%;white-space:nowrap;text-align:center;
  font-family:var(--ff-d);font-weight:900;font-size:clamp(52px,17vw,92px);line-height:1;
  letter-spacing:-1px;padding:12px 34px;border:3.5px solid var(--c-border);
  border-radius:var(--r-pill);box-shadow:0 8px 0 rgba(0,0,0,.28);
  transform-origin:center;pointer-events:none}
#quiz-flash.bonus #quiz-flash-bonus{display:block;animation:bonus-pop .6s cubic-bezier(.34,1.56,.64,1) both}
@keyframes bonus-pop{0%{transform:scale(.2) rotate(-8deg);opacity:0}55%{transform:scale(1.18) rotate(4deg);opacity:1}100%{transform:scale(1) rotate(0);opacity:1}}
/* 正解時（.char付き）: キャラが下から大きく飛び出す（マークは非表示） */
#quiz-flash-char{display:none;position:absolute;bottom:0;left:0;right:0;margin:0 auto;
  height:52vh;max-height:440px;width:auto;transform-origin:center bottom;
  filter:drop-shadow(0 8px 16px rgba(0,0,0,.28));pointer-events:none}
#quiz-flash.char #quiz-flash-mark{display:none}
#quiz-flash.char #quiz-flash-char{display:block;animation:char-jump-in 1.15s cubic-bezier(.22,.85,.3,1) forwards}
@keyframes char-jump-in{
  0%{transform:translateY(120%) scale(.78)}
  70%{transform:translateY(-5%) scale(1.03)}
  100%{transform:translateY(0) scale(1)}
}
/* レア（ハート/シー）: 2回はっきり飛び出し＋頂点で少し長め */
#quiz-flash.char.twice #quiz-flash-char{animation:char-jump-twice 2.1s cubic-bezier(.34,1.2,.5,1) forwards}
@keyframes char-jump-twice{
  0%{transform:translateY(125%) scale(.72)}     /* 画面下から登場 */
  18%{transform:translateY(-20%) scale(1.06)}   /* 1回目の頂点 */
  40%{transform:translateY(70%) scale(.92)}     /* しっかり下がる */
  60%{transform:translateY(-20%) scale(1.06)}   /* 2回目の頂点 */
  78%{transform:translateY(-20%) scale(1.06)}   /* 頂点で少し停止 */
  100%{transform:translateY(0) scale(1)}        /* 着地 */
}
#modal-backdrop{position:fixed;inset:0;background:rgba(0,0,0,0);z-index:999;
  display:none;transition:background .28s}
#modal-backdrop.open{display:block;background:rgba(0,0,0,.45)}
#modal{position:fixed;bottom:0;left:50%;transform:translate(-50%,100%);
  width:100%;max-width:600px;
  max-height:82vh;z-index:1000;
  background:var(--c-canvas);
  border:var(--bw) solid var(--c-border);border-bottom:none;
  border-radius:24px 24px 0 0;
  display:flex;flex-direction:column;
  transition:transform .3s cubic-bezier(.32,.72,0,1),background .25s;
  box-shadow:0 -4px 32px rgba(0,0,0,.18)}
#modal.open{transform:translate(-50%,0)}
.modal-drag{width:36px;height:4px;background:var(--c-muted2);
  border-radius:2px;margin:10px auto 0;flex-shrink:0;cursor:grab}
.modal-hdr{display:flex;align-items:center;padding:8px 16px 10px;
  border-bottom:1px solid var(--c-hairline);flex-shrink:0;
  transition:border-color .25s}
.modal-close{width:32px;height:32px;border-radius:50%;border:none;
  background:var(--c-parchment);cursor:pointer;display:flex;align-items:center;
  justify-content:center;color:var(--c-ink);font-size:15px;font-weight:600;
  flex-shrink:0;-webkit-tap-highlight-color:transparent;transition:background .15s;
  line-height:1;margin-left:auto}
.modal-close:active{background:var(--c-hairline)}
.modal-fav-btn{background:transparent;border:none;cursor:pointer;
  padding:4px;color:var(--c-muted);
  -webkit-tap-highlight-color:transparent;transition:color .15s}
.modal-fav-btn.on{color:var(--c-fav)}
.modal-fav-btn svg{width:24px;height:24px;stroke:currentColor;
  stroke-width:1.8;stroke-linecap:round;stroke-linejoin:round}
.modal-body{flex:1;min-height:0;overflow-y:auto;padding:18px 20px 32px;
  -webkit-overflow-scrolling:touch}
.modal-body::-webkit-scrollbar{display:none}
.modal-term{font-family:var(--ff-d);font-size:32px;font-weight:600;
  line-height:1.14;letter-spacing:-0.374px;color:var(--c-ink);margin-bottom:14px;
  transition:color .25s}
.modal-def{font-family:var(--ff-t);font-size:18px;font-weight:400;
  line-height:1.72;letter-spacing:-0.3px;color:var(--c-ink);
  transition:color .25s}
.def-xref{color:var(--c-primary);text-decoration:underline;cursor:pointer;
  text-decoration-thickness:1px;text-underline-offset:3px;font-weight:600;
  -webkit-tap-highlight-color:transparent}
/* 拡充表示（見出し/短文/長文/例/関連語） */
.m-short{font-family:var(--ff-t);font-size:18px;font-weight:700;line-height:1.7;margin:0 0 10px}
.m-long{font-family:var(--ff-t);font-size:16px;font-weight:400;line-height:1.75;color:var(--c-ink);margin:0 0 12px}
.m-ex{font-family:var(--ff-t);font-size:15px;line-height:1.7;color:#333;margin:0 0 12px}
.m-ex .lbl,.m-rel .lbl{font-weight:700;color:var(--c-primary)}
.m-ex .en{font-style:normal}
.m-rel{font-family:var(--ff-t);font-size:15px;line-height:1.9;margin:0}
.m-rel .def-xref{margin-right:2px}

/* ── Fav tab ── */
.fav-hdr{padding:16px 14px 8px;background:var(--c-canvas);
  border-bottom:1px solid var(--c-hairline);flex-shrink:0;transition:background .25s}
.fav-hdr h1{font-family:var(--ff-d);font-size:26px;font-weight:900;letter-spacing:-0.28px}
.empty{display:flex;flex-direction:column;align-items:center;justify-content:center;
  flex:1;padding:48px 24px;text-align:center;gap:14px;color:var(--c-muted)}
.empty svg{width:40px;height:40px;stroke:currentColor;stroke-width:1.5;
  fill:none;stroke-linecap:round;stroke-linejoin:round;opacity:.35}
.empty img{height:110px;width:auto;opacity:.9}
.empty p{font-size:15px;line-height:1.55}

/* ── Quiz ── */
#tab-quiz{overflow-y:auto;-webkit-overflow-scrolling:touch}
#tab-quiz::-webkit-scrollbar{display:none}
.quiz-inner{display:flex;flex-direction:column;min-height:100%;padding-bottom:24px}
.quiz-char{height:120px;background:var(--c-accent);
  display:flex;align-items:center;justify-content:center;
  border-bottom:var(--bw) solid var(--c-border);flex-shrink:0;
  transition:background .25s,border-color .25s}
.quiz-char img{height:106px;width:auto;object-fit:contain;transition:opacity .3s}
@keyframes char-pop{0%{transform:scale(1)}45%{transform:scale(1.18) rotate(-4deg)}100%{transform:scale(1)}}
.char-pop{animation:char-pop .5s cubic-bezier(.34,1.56,.64,1)}
/* ダイヤモンド上のキャラ用（中央寄せtranslateを保持したままpop） */
@keyframes char-pop-d{0%{transform:translate(-50%,-50%) scale(1)}45%{transform:translate(-50%,-50%) scale(1.2) rotate(-4deg)}100%{transform:translate(-50%,-50%) scale(1)}}
.char-pop-d{animation:char-pop-d .5s cubic-bezier(.34,1.56,.64,1)}
/* ── クイズ目次 ── */
#quiz-menu{flex:1;overflow-y:auto;padding:14px 16px 20px;display:flex;flex-direction:column;gap:12px}
#quiz-menu::-webkit-scrollbar{display:none}
.qmenu-title-row{display:flex;align-items:center;justify-content:center;flex-wrap:wrap;gap:6px 10px;margin-bottom:0;
  font-family:var(--ff-d);font-size:22px;font-weight:900;color:var(--c-ink)}
.qmenu-title-row img{height:44px;width:auto}
.qmenu-runs{font-family:var(--ff-t);font-size:14px;font-weight:800;color:var(--c-ink);white-space:nowrap}
.qmenu-runs b{color:var(--c-coral);font-family:var(--ff-d);font-size:19px}
.qmenu-card{display:flex;align-items:center;gap:14px;padding:11px 16px;
  border:var(--bw) solid var(--c-border);border-radius:var(--r-lg);box-shadow:var(--shadow-m);
  cursor:pointer;text-align:left;transition:all .12s;-webkit-tap-highlight-color:transparent}
.qmenu-card:active{transform:translate(3px,3px);box-shadow:none}
.qmenu-card img{height:56px;width:56px;object-fit:contain;flex-shrink:0}
.qmenu-quiz{background:var(--c-accent)}
.qmenu-daily{background:var(--c-orange)}
.qmenu-review{background:#ffd9d6}
.qmenu-txt{display:flex;flex-direction:column;gap:3px;min-width:0}
.qmenu-group-label{font-family:var(--ff-d);font-weight:900;font-size:15px;color:var(--c-muted);
  margin:6px 2px -4px;letter-spacing:.5px}
.qmenu-name{font-family:var(--ff-d);font-size:20px;font-weight:900;color:var(--c-ink)}
.qmenu-sub{font-size:13px;font-weight:600;color:var(--c-ink);opacity:.75;line-height:1.4}
.qmenu-score{background:var(--c-parchment);border:var(--bw) solid var(--c-border);
  border-radius:var(--r-lg);box-shadow:var(--shadow-s);padding:8px 0 12px}
.qmenu-score .score-cols{display:flex}
.qmenu-score .score-col{flex:1;text-align:center}
/* 点の入り方 */
.qmenu-howto{background:var(--c-card);border:2px solid var(--c-border);border-radius:var(--r-md);
  box-shadow:var(--shadow-s);padding:11px 14px}
.qmenu-howto-ttl{display:flex;align-items:center;gap:6px;width:100%;background:none;border:none;padding:0;
  font-family:var(--ff-d);font-weight:900;font-size:14px;color:var(--c-ink);cursor:pointer;
  -webkit-tap-highlight-color:transparent}
.howto-chev{font-size:11px;color:var(--c-muted);transition:transform .15s}
.howto-chev.open{transform:rotate(180deg)}
.howto-swingrow{display:flex;align-items:center;justify-content:space-between;
  margin:9px 0 4px;padding:7px 10px;background:var(--c-score-bg);border-radius:var(--r-md)}
.howto-swing-lbl{font-family:var(--ff-t);font-weight:800;font-size:13px;color:var(--c-ink)}
.howto-swing-hint{font-family:var(--ff-t);font-weight:700;font-size:12px;color:var(--c-primary);margin:4px 2px 0}
.qmenu-howto ul{margin:8px 0 0;padding-left:18px}
.howto-sub-ttl{font-family:var(--ff-d);font-weight:900;font-size:13px;color:var(--c-ink);margin:12px 0 2px;
  padding-top:9px;border-top:1px dashed var(--c-hairline)}
.qmenu-howto li{font-family:var(--ff-t);font-size:12.5px;color:var(--c-ink);opacity:.85;line-height:1.6}
.qmenu-howto b{color:var(--c-coral)}
/* ── サブページ共通（戻るバー） ── */
.subpage-bar{display:flex;align-items:center;gap:10px;padding:6px 14px;
  background:var(--c-canvas);border-bottom:1px solid var(--c-hairline);flex-shrink:0}
.subpage-title{font-family:var(--ff-d);font-size:17px;font-weight:900;color:var(--c-ink)}
.daily-runs-lbl{margin-left:auto;font-family:var(--ff-d);font-size:14px;font-weight:900;color:var(--c-primary)}
.daily-runs-lbl #daily-runs{color:var(--c-coral);font-size:17px}
.back-btn{padding:6px 14px;border:2px solid var(--c-border);border-radius:var(--r-pill);
  background:var(--c-parchment);box-shadow:var(--shadow-s);font-family:var(--ff-t);
  font-size:13px;font-weight:800;color:var(--c-ink);cursor:pointer;flex-shrink:0;
  transition:all .12s;-webkit-tap-highlight-color:transparent}
.back-btn:active{transform:translate(2px,2px);box-shadow:none}
.subpage-hdr{display:flex;align-items:center;gap:12px}
/* ── ダイヤモンド（塁）エリア ── */
#quiz-play{flex:1;display:flex;flex-direction:column;overflow-y:auto;min-height:0}
#quiz-play::-webkit-scrollbar{display:none}
#quiz-char{height:158px;background:#c9ecff}
.diamond-wrap{position:relative;height:100%;width:100%}
.quiz-char img.diamond-img{position:absolute;inset:0;width:100%;height:100%;object-fit:fill;mix-blend-mode:multiply}
.quiz-char img.dchar{position:absolute;height:26%;width:auto;transform:translate(-50%,-50%);
  filter:drop-shadow(0 2px 2px rgba(0,0,0,.25));transition:opacity .3s}
.quiz-char img.dchar-home{left:50%;top:80%;height:34%}
/* 塁上ランナー: 足が塁に乗るよう下端寄りにアンカー */
.quiz-char img.dchar-runner{height:24%;transform:translate(-50%,-84%)}
/* 盗塁待ち＝各塁から次の塁方向へ1歩リードした位置に置く */
#dchar-1{left:74%;top:44%}  /* 一塁→二塁方向へリード */
#dchar-2{left:43%;top:20%}  /* 二塁→三塁方向へリード */
#dchar-3{left:26%;top:53%}  /* 三塁→本塁方向へリード */
/* 塁間を実際に走る走者（塁パスに沿って移動） */
.quiz-char img.dchar-mover{position:absolute;height:24%;width:auto;
  transform:translate(-50%,-84%);z-index:8;pointer-events:none;
  filter:drop-shadow(0 3px 5px rgba(0,0,0,.28));
  transition:left .42s linear, top .42s linear}
/* ── スイングタイミング演出 ── */
.swing-ui{display:flex;flex-direction:column;align-items:center;justify-content:center;gap:7px;
  margin:10px 14px 0;padding:12px;border:2px solid var(--c-border);border-radius:var(--r-lg);
  background:#c9ecff;box-shadow:var(--shadow-s)}
.swing-pitch{font-family:var(--ff-d);font-weight:900;font-size:13px;color:#fff;
  background:var(--c-primary);border:2px solid var(--c-border);border-radius:var(--r-pill);
  padding:2px 12px;box-shadow:var(--shadow-s)}
.swing-track{position:relative;width:78%;height:20px;border-radius:var(--r-pill);
  background:rgba(255,255,255,.85);border:2px solid var(--c-border);overflow:hidden;
  box-shadow:var(--shadow-s)}
.swing-zone{position:absolute;top:0;bottom:0}
.swing-zone-good{background:#bfe3ff}
.swing-zone-great{background:#ffe08a}
.swing-zone-perfect{background:#ff8f6b}
.swing-marker{position:absolute;top:-3px;bottom:-3px;width:4px;left:0;margin-left:-2px;
  background:var(--c-ink);border-radius:2px;box-shadow:0 0 0 1px #fff}
.swing-guide{font-family:var(--ff-t);font-weight:700;font-size:11px;color:var(--c-ink);
  opacity:.75;letter-spacing:.2px}
.swing-btn{font-family:var(--ff-d);font-weight:900;font-size:16px;color:#fff;
  background:var(--c-coral);border:3px solid var(--c-border);border-radius:var(--r-pill);
  padding:6px 26px;box-shadow:var(--shadow-m);cursor:pointer;-webkit-tap-highlight-color:transparent}
.swing-btn:active{transform:translate(2px,2px);box-shadow:none}
.swing-btn:disabled{opacity:.55}
.swing-result{font-family:var(--ff-d);font-weight:900;font-size:17px;color:var(--c-ink);min-height:20px}
.swing-result.hit{color:var(--c-coral)}
.swing-result.out{color:var(--c-muted)}
/* 設定: スイング操作トグル */
.swing-toggle{flex-shrink:0;padding:8px 18px;border:var(--bw) solid var(--c-border);
  border-radius:var(--r-pill);background:var(--c-accent);box-shadow:var(--shadow-s);
  font-family:var(--ff-t);font-size:14px;font-weight:800;color:var(--c-ink);cursor:pointer;
  -webkit-tap-highlight-color:transparent;transition:all .12s}
.swing-toggle:active{transform:translate(2px,2px);box-shadow:none}
.swing-toggle.auto{background:var(--c-muted2);color:#fff}
/* ── アウトカウント・得点（フィールド上のバッジ） ── */
.diamond-hud{position:absolute;top:8px;z-index:6;
  font-family:var(--ff-d);font-weight:900;font-size:13px;
  padding:3px 11px;border-radius:var(--r-pill);
  border:2px solid var(--c-border);background:rgba(255,255,255,.92);
  box-shadow:var(--shadow-s)}
.diamond-hud-out{left:8px;color:var(--c-coral);letter-spacing:2px}
.diamond-hud-run{right:8px;color:var(--c-primary);font-size:12px}
.diamond-hud-run #quiz-runs,.diamond-hud-run #quiz-runs-today{display:inline-block;color:var(--c-coral);
  font-size:15px;min-width:1.0em;text-align:center}
@keyframes run-bump{0%{transform:scale(1)}40%{transform:scale(1.5)}100%{transform:scale(1)}}
.run-bump{animation:run-bump .45s cubic-bezier(.34,1.56,.64,1)}
/* 難易度バッジ */
.hit-badge{display:inline-block;font-family:var(--ff-t);font-size:11px;font-weight:800;
  padding:1px 7px;border-radius:var(--r-pill);border:1.5px solid var(--c-border);margin-left:6px;vertical-align:middle}
.hit-1{background:var(--c-accent);color:var(--c-ink)}
.hit-2{background:var(--c-orange);color:var(--c-ink)}
.hit-3{background:var(--c-coral);color:#fff}
/* 目次の累計得点 */
.menu-runs-row{text-align:center;font-family:var(--ff-d);font-weight:900;font-size:15px;
  color:var(--c-primary);padding:6px 0 2px;border-top:1px solid var(--c-hairline);margin-top:4px}
.menu-runs-row #menu-runs,.menu-runs-row #menu-runs-today{color:var(--c-coral);font-size:20px}
.quiz-card{position:relative}
/* ── THREE OUT 演出 ── */
#three-out{position:fixed;inset:0;z-index:900;display:none;align-items:center;justify-content:center;
  background:rgba(30,30,46,.55);pointer-events:none;opacity:1;transition:opacity .7s ease}
#three-out.show{display:flex}
#three-out.fade{opacity:0}
.three-out-inner{display:flex;flex-direction:column;align-items:center;gap:20px}
.three-out-txt{font-family:var(--ff-d);font-size:52px;font-weight:900;color:#fff;
  background:var(--c-coral);border:4px solid var(--c-border);border-radius:var(--r-lg);
  box-shadow:var(--shadow-l);padding:18px 34px;transform:rotate(-6deg);
  animation:three-out-pop .45s cubic-bezier(.34,1.56,.64,1)}
.three-out-img{width:62vw;max-width:340px;height:auto;
  filter:drop-shadow(0 8px 16px rgba(0,0,0,.3));
  animation:three-out-img-pop .5s cubic-bezier(.34,1.56,.64,1)}
@keyframes three-out-pop{0%{transform:rotate(-6deg) scale(.3)}70%{transform:rotate(-6deg) scale(1.12)}100%{transform:rotate(-6deg) scale(1)}}
@keyframes three-out-img-pop{0%{transform:scale(.3)}70%{transform:scale(1.1)}100%{transform:scale(1)}}
/* ── CPU対戦 ── */
.qmenu-versus{background:#ffe0a8}
#versus-bar{padding:7px 14px;background:var(--c-score-bg);border-bottom:1px solid var(--c-hairline);
  font-family:var(--ff-t);font-size:13px;color:var(--c-ink);text-align:center}
#versus-bar b{color:var(--c-coral);font-family:var(--ff-d);font-size:15px}
#versus-inning{padding:6px 14px;text-align:center;font-family:var(--ff-d);font-weight:900;
  font-size:14px;color:var(--c-primary);background:var(--c-canvas)}
#versus-board{position:fixed;inset:0;z-index:800;display:none;align-items:flex-start;justify-content:center;
  background:rgba(20,26,46,.6);padding:18px;overflow-y:auto}
#versus-board.show{display:flex}
#versus-board-inner{width:100%;max-width:400px;margin:auto;background:var(--c-card);border:3px solid var(--c-border);
  border-radius:var(--r-lg);box-shadow:var(--shadow-l);padding:16px;text-align:center}
.vs-board-ttl{font-family:var(--ff-d);font-weight:900;font-size:18px;color:var(--c-ink);margin-bottom:6px}
.vs-names{font-family:var(--ff-t);font-weight:800;font-size:13px;color:var(--c-ink);margin-bottom:8px;word-break:break-word}
.vs-score{width:100%;border-collapse:collapse;font-family:var(--ff-t);font-size:12px;margin-bottom:8px}
.vs-score th,.vs-score td{border:1px solid var(--c-hairline);padding:3px 0;text-align:center;min-width:18px}
.vs-score th:first-child{text-align:left;padding-left:4px;white-space:nowrap;font-size:11px}
.vs-score .vs-hdr th{background:var(--c-score-bg);color:var(--c-muted)}
.vs-score .vs-r{font-family:var(--ff-d);font-weight:900;color:var(--c-coral);background:var(--c-score-bg)}
.vs-msg{font-family:var(--ff-t);font-weight:700;font-size:13px;color:var(--c-ink);margin:6px 0 10px}
.vs-btns{display:flex;flex-direction:column;gap:8px}
.vs-btn{padding:11px 0;border:2px solid var(--c-border);border-radius:var(--r-pill);
  background:var(--c-accent);font-family:var(--ff-d);font-weight:900;font-size:15px;color:var(--c-ink);
  cursor:pointer;-webkit-tap-highlight-color:transparent}
.vs-btn:active{transform:translate(2px,2px)}
.vs-btn-sub{background:none;border:none;color:var(--c-muted);font-size:13px;font-weight:700;text-decoration:underline}
#versus-toast{position:fixed;inset:0;z-index:820;display:none;align-items:center;justify-content:center;
  background:rgba(20,26,46,.6);white-space:pre-line;text-align:center;
  font-family:var(--ff-d);font-weight:900;font-size:20px;color:#fff;padding:24px;line-height:1.6}
#versus-toast.show{display:flex}
#rankup-toast{position:fixed;left:50%;top:18%;transform:translateX(-50%) translateY(-16px);z-index:900;
  display:none;opacity:0;transition:opacity .35s,transform .35s;pointer-events:none;max-width:88%;
  background:linear-gradient(135deg,#ffd34d,#ff9f1c);color:#3a2a00;border:3px solid #fff;
  box-shadow:0 8px 26px rgba(0,0,0,.32);border-radius:16px;padding:14px 20px;text-align:center;
  font-family:var(--ff-d);font-weight:900;line-height:1.4}
#rankup-toast.show{display:block;opacity:1;transform:translateX(-50%) translateY(0)}
#rankup-toast .ru-sub{font-size:13px;color:#7a5200}
#rankup-toast .ru-name{font-size:24px;color:#c23b00}
#versus-result{position:fixed;inset:0;z-index:830;display:none;align-items:flex-start;justify-content:center;
  background:rgba(20,26,46,.7);padding:24px;overflow-y:auto}
#versus-result.show{display:flex}
.vs-res-inner{width:100%;max-width:340px;margin:auto;background:var(--c-card);border:3px solid var(--c-border);
  border-radius:var(--r-lg);box-shadow:var(--shadow-l);padding:22px 18px;text-align:center}
.vs-res-char{width:120px;height:auto;margin-bottom:8px}
.vs-res-ttl{font-family:var(--ff-d);font-weight:900;font-size:26px;color:var(--c-ink)}
.vs-res-score{font-family:var(--ff-d);font-weight:900;font-size:18px;color:var(--c-coral);margin:6px 0}
.vs-res-bonus{font-family:var(--ff-t);font-weight:800;font-size:14px;color:var(--c-primary);min-height:18px;margin-bottom:12px}
/* ── 広告オーバーレイ（プレースホルダ） ── */
#ad-overlay{position:fixed;inset:0;z-index:1000;display:none;align-items:flex-start;justify-content:center;
  background:rgba(20,26,46,.75);padding:24px;overflow-y:auto}
#ad-overlay.show{display:flex}
.ad-box{width:100%;max-width:340px;margin:auto;background:var(--c-card);border:3px solid var(--c-border);
  border-radius:var(--r-lg);box-shadow:var(--shadow-l);padding:14px;text-align:center}
.ad-tag{display:inline-block;font-family:var(--ff-t);font-size:11px;font-weight:800;color:#fff;
  background:var(--c-muted2);border-radius:var(--r-pill);padding:2px 12px;margin-bottom:10px}
.ad-body{background:var(--c-score-bg);border:2px dashed var(--c-border);border-radius:var(--r-md);
  padding:26px 12px;margin-bottom:12px}
.ad-body img{width:84px;height:auto;margin-bottom:8px}
.ad-msg{font-family:var(--ff-d);font-weight:900;font-size:18px;color:var(--c-ink)}
.ad-note{font-size:11.5px;color:var(--c-muted2);margin-top:4px}
.ad-close{display:block;width:100%;padding:11px 0;border:2px solid var(--c-border);border-radius:var(--r-pill);
  background:var(--c-accent);font-family:var(--ff-d);font-weight:900;font-size:15px;color:var(--c-ink);
  cursor:pointer;-webkit-tap-highlight-color:transparent}
.ad-close:disabled{opacity:.5}
.ad-premium{display:block;width:100%;margin-top:8px;padding:8px 0;border:none;background:none;
  font-family:var(--ff-t);font-size:12.5px;font-weight:700;color:var(--c-primary);cursor:pointer;text-decoration:underline}
/* ── 辞書検索バー横の略語ボタン ── */
.search-row{display:flex;align-items:center;gap:8px}
.search-row .search-wrap{flex:1;min-width:0}
.abbr-jump-btn{flex-shrink:0;padding:9px 16px;border:var(--bw) solid var(--c-border);
  border-radius:var(--r-pill);background:var(--c-orange);box-shadow:var(--shadow-s);
  font-family:var(--ff-t);font-size:14px;font-weight:800;color:var(--c-ink);cursor:pointer;
  transition:all .12s;-webkit-tap-highlight-color:transparent}
.abbr-jump-btn:active{transform:translate(2px,2px);box-shadow:none}
.quiz-score-bar{display:flex;flex-direction:column;background:var(--c-score-bg);
  border-bottom:1px solid var(--c-hairline);flex-shrink:0;
  transition:background .25s,border-color .25s}
.score-correct-hdr{text-align:center;font-family:var(--ff-t);font-size:11px;
  color:var(--c-muted2);padding:7px 0 0;letter-spacing:.3px}
.score-cols{display:flex}
.score-col{flex:1;text-align:center;padding:2px 8px 9px}
.score-lbl{font-family:var(--ff-t);font-size:13px;color:var(--c-muted);
  font-weight:600;letter-spacing:.2px}
.score-val-row{display:flex;align-items:baseline;justify-content:center;gap:6px;margin-top:1px}
.score-val{font-family:var(--ff-d);font-size:20px;font-weight:600;
  color:var(--c-ink);letter-spacing:-0.28px;transition:color .25s}
.score-rate{font-family:var(--ff-d);font-size:13px;font-weight:500;
  color:var(--c-muted);letter-spacing:0;transition:color .25s}
.quiz-mode-bar{display:flex;padding:10px 14px;gap:8px;background:var(--c-canvas);
  border-bottom:1px solid var(--c-hairline);flex-shrink:0;transition:background .25s}
.qmode-btn{flex:1;padding:8px 14px;border-radius:var(--r-pill);
  border:var(--bw) solid var(--c-border);background:var(--c-parchment);
  box-shadow:var(--shadow-s);
  font-family:var(--ff-t);font-size:14px;font-weight:700;color:var(--c-muted);
  cursor:pointer;transition:all .12s;-webkit-tap-highlight-color:transparent}
.qmode-btn:active{transform:translate(2px,2px);box-shadow:none}
.qmode-btn.active{background:var(--c-primary);color:#fff}
.quiz-card{padding:10px 14px 12px;flex:1}
.pitch-clock{height:10px;background:var(--c-card);border:2px solid var(--c-border);
  border-radius:6px;overflow:hidden;margin-bottom:8px}
.pitch-clock-fill{height:100%;width:100%;background:var(--c-accent)}
.pitch-clock-fill.warn{background:var(--c-coral);animation:pc-blink .5s steps(2,start) infinite}
@keyframes pc-blink{50%{opacity:.4}}
.quiz-q{background:var(--c-parchment);border-radius:var(--r-lg);
  border:var(--bw) solid var(--c-border);box-shadow:var(--shadow-m);
  padding:10px 14px;text-align:center;margin-bottom:8px;transition:background .25s}
.quiz-q-lbl{font-family:var(--ff-t);font-size:10.5px;color:var(--c-muted);
  font-weight:600;text-transform:uppercase;letter-spacing:.5px;margin-bottom:4px}
.quiz-q-term{font-family:var(--ff-d);font-size:21px;font-weight:600;
  line-height:1.18;letter-spacing:-0.28px;color:var(--c-ink);
  word-break:break-word;transition:color .25s}
.choices{display:flex;flex-direction:column;gap:8px;margin-bottom:8px}
.choice{width:100%;padding:12px 14px;border-radius:var(--r-md);
  border:var(--bw) solid var(--c-border);background:var(--c-parchment);
  box-shadow:var(--shadow-s);
  font-family:var(--ff-t);font-size:14.5px;font-weight:600;color:var(--c-ink);
  cursor:pointer;text-align:left;transition:all .12s;
  -webkit-tap-highlight-color:transparent;line-height:1.4}
.choice:active:not(:disabled){transform:translate(2px,2px);box-shadow:none}
.choice.ok{background:var(--c-ok-bg);color:var(--c-ok-text);font-weight:800}
.choice.ng{background:var(--c-ng-bg);color:var(--c-ng-text)}
.choice:disabled{cursor:default}
.next-btn{display:block;width:100%;padding:13px;border-radius:var(--r-pill);
  border:var(--bw) solid var(--c-border);background:var(--c-coral);color:#fff;
  box-shadow:var(--shadow-m);
  font-family:var(--ff-t);font-size:17px;font-weight:800;
  cursor:pointer;-webkit-tap-highlight-color:transparent;transition:all .12s}
.next-btn:active{transform:translate(3px,3px);box-shadow:none}
.next-btn:disabled{opacity:.3;cursor:default;transform:none;box-shadow:var(--shadow-m)}

/* ── Daily quiz tab ── */
#tab-daily{overflow-y:auto;-webkit-overflow-scrolling:touch}
#tab-daily::-webkit-scrollbar{display:none}
.daily-inner{display:flex;flex-direction:column;min-height:100%;padding-bottom:24px}
.daily-progress-bar{display:flex;align-items:center;gap:10px;
  padding:10px 14px;background:var(--c-canvas);
  border-bottom:1px solid var(--c-hairline);flex-shrink:0}
.daily-progress-track{flex:1;height:12px;background:var(--c-parchment);border:2px solid var(--c-border);border-radius:6px;overflow:hidden}
.daily-progress-fill{height:100%;background:var(--c-accent);transition:width .3s}
.daily-progress-lbl{font-size:13px;color:var(--c-muted);white-space:nowrap;font-weight:600}
.daily-result{padding:0 16px 16px;text-align:center;display:flex;flex-direction:column;gap:16px;align-items:stretch;width:100%;box-sizing:border-box}
.daily-result-hero{position:relative;width:100%;min-height:120px;display:flex;flex-direction:column;align-items:center;justify-content:center;gap:8px;padding:18px 0 16px;overflow:hidden;
  background:var(--c-accent);border:var(--bw) solid var(--c-border);box-shadow:var(--shadow-m);border-radius:var(--r-lg)}
.daily-result-char{height:110px;width:auto;animation:char-pop .6s cubic-bezier(.34,1.56,.64,1)}
.daily-result-score{font-family:var(--ff-d);font-size:48px;font-weight:900;color:var(--c-ink)}
.daily-result-sub{font-size:16px;font-weight:700;color:var(--c-ink)}
.daily-bonus{font-family:var(--ff-d);font-size:15px;font-weight:900;color:var(--c-coral);
  background:#fff;border:2px solid var(--c-border);border-radius:var(--r-pill);
  padding:4px 14px;margin-top:2px}
.daily-done-note{font-size:13px;color:var(--c-sub);margin-top:-8px}
#quiz-def-card,#daily-def-card{margin:12px 0 4px;padding:16px;background:var(--c-parchment);
  border:var(--bw) solid var(--c-border);box-shadow:var(--shadow-s);border-radius:var(--r-md);text-align:left}
#daily-def-card .ddc-en,#quiz-def-card .ddc-en{font-size:13px;color:var(--c-primary);margin-bottom:4px}
#daily-def-card .ddc-head,#quiz-def-card .ddc-head{font-size:20px;font-weight:700;color:var(--c-ink);margin-bottom:6px;font-family:var(--ff-d)}
#daily-def-card .ddc-sub,#quiz-def-card .ddc-sub{font-size:13px;color:var(--c-muted);margin-bottom:8px}
#daily-def-card .ddc-def,#quiz-def-card .ddc-def{font-size:14px;color:var(--c-ink);line-height:1.7}
/* Start screen */
.daily-start-inner{display:flex;flex-direction:column;align-items:center;justify-content:center;
  gap:16px;padding:60px 24px 40px;text-align:center}
.daily-start-icon{font-size:56px;line-height:1}
.daily-start-icon img{height:120px;width:auto}
.daily-start-title{font-family:var(--ff-d);font-size:28px;font-weight:900;color:var(--c-ink)}
.daily-start-sub{font-size:15px;color:var(--c-muted)}
.daily-start-btn{margin-top:8px;padding:16px 48px;
  background:var(--c-coral);color:#fff;border:var(--bw) solid var(--c-border);border-radius:50px;
  font-size:18px;font-weight:800;cursor:pointer;letter-spacing:.05em;
  box-shadow:var(--shadow-l);transition:all .12s;-webkit-tap-highlight-color:transparent}
.daily-start-btn:active{transform:translate(4px,4px);box-shadow:none}
/* Calendar */
.daily-cal{padding:0 0 20px;width:100%;box-sizing:border-box}
.cal-inner{width:100%;box-sizing:border-box;margin:0 auto;border:var(--bw) solid var(--c-border);box-shadow:var(--shadow-m);border-radius:var(--r-md);overflow:hidden;background:var(--c-parchment)}
.cal-nav{display:flex;align-items:center;justify-content:space-between;padding:8px 12px 6px}
.cal-title{font-family:var(--ff-d);font-size:17px;font-weight:600;color:var(--c-ink)}
.cal-nav-btn{background:none;border:none;cursor:pointer;font-size:15px;
  color:var(--c-primary);padding:4px 8px;-webkit-tap-highlight-color:transparent;white-space:nowrap}
.cal-grid{display:grid;grid-template-columns:14.2857% 14.2857% 14.2857% 14.2857% 14.2857% 14.2857% 14.2858%;grid-auto-rows:40px;gap:0;width:100%;margin:0 auto}
.cal-dow{text-align:center;font-size:11px;font-weight:700;color:var(--c-muted2);
  padding:4px 0;letter-spacing:.03em;min-width:0;overflow:hidden}
.cal-day{position:relative;text-align:center;
  display:flex;align-items:center;justify-content:center;
  font-size:13px;color:var(--c-ink);min-width:0;overflow:hidden}
.cal-day.empty{color:transparent}
.cal-day.today{font-weight:700;color:var(--c-primary)}
.cal-heart{position:absolute;inset:0;display:flex;align-items:center;justify-content:center;
  pointer-events:none}

/* ── Review tab ── */
#tab-review{overflow-y:auto;-webkit-overflow-scrolling:touch}
#tab-review::-webkit-scrollbar{display:none}
.review-hdr{padding:16px 14px 8px;background:var(--c-canvas);
  border-bottom:1px solid var(--c-hairline);flex-shrink:0;transition:background .25s}
.review-hdr h1{font-family:var(--ff-d);font-size:26px;font-weight:900;letter-spacing:-0.28px}
.review-clear-btn{margin-left:auto;flex-shrink:0;padding:7px 14px;border:2px solid var(--c-border);
  border-radius:var(--r-pill);background:var(--c-coral);box-shadow:var(--shadow-s);
  font-family:var(--ff-t);font-size:13px;font-weight:800;color:#fff;cursor:pointer;
  -webkit-tap-highlight-color:transparent;transition:all .12s}
.review-clear-btn:active{transform:translate(2px,2px);box-shadow:none}
.review-item{padding:12px 14px;border-bottom:1px solid var(--c-hairline);
  display:flex;align-items:center;gap:10px;cursor:pointer;
  background:var(--c-canvas);transition:background .1s;
  -webkit-tap-highlight-color:transparent}
.review-item:active{background:var(--c-parchment)}
.review-del{background:none;border:none;font-size:18px;color:var(--c-muted2);
  cursor:pointer;padding:4px 8px;-webkit-tap-highlight-color:transparent;
  flex-shrink:0;transition:color .15s}
.review-del:active{color:var(--c-ng-border)}

/* ── Settings tab ── */
#tab-settings{overflow-y:auto;-webkit-overflow-scrolling:touch}
#tab-settings::-webkit-scrollbar{display:none}
.settings-inner{padding:20px 16px;display:flex;flex-direction:column;gap:20px}
.settings-hdr{font-family:var(--ff-d);font-size:26px;font-weight:900;
  letter-spacing:-0.28px;color:var(--c-ink);margin-bottom:4px;transition:color .25s}
.settings-section{background:var(--c-card);border-radius:var(--r-lg);
  overflow:hidden;border:var(--bw) solid var(--c-border);box-shadow:var(--shadow-s);
  transition:background .25s,border-color .25s}
.settings-row{display:flex;align-items:center;padding:14px 16px;
  border-bottom:1px solid var(--c-hairline);gap:12px;transition:border-color .25s}
.settings-row:last-child{border-bottom:none}
.settings-row-label{flex:1;font-family:var(--ff-t);font-size:16px;color:var(--c-ink);
  transition:color .25s}
.settings-row-sub{font-size:12px;color:var(--c-muted);margin-top:2px;transition:color .25s}
.reset-btn{flex-shrink:0;padding:7px 14px;border-radius:var(--r-sm);border:2px solid var(--c-border);
  background:var(--c-primary);color:#fff;font-size:13px;font-weight:700;font-family:var(--ff-t);
  box-shadow:var(--shadow-s);cursor:pointer;transition:all .12s}
.reset-btn:active{transform:translate(2px,2px);box-shadow:none}
.reset-confirm{display:none;align-items:center;justify-content:space-between;
  padding:10px 16px 12px;gap:10px;flex-wrap:wrap}
.reset-confirm.show{display:flex}
.reset-confirm-msg{font-size:14px;color:var(--c-muted);flex:1}
.reset-confirm-btns{display:flex;gap:8px;flex-shrink:0}
.reset-confirm-cancel{padding:7px 14px;border-radius:var(--r-sm);border:2px solid var(--c-border);
  background:var(--c-parchment);color:var(--c-ink);font-size:13px;font-weight:700;font-family:var(--ff-t);cursor:pointer}
.reset-confirm-ok{padding:7px 14px;border-radius:var(--r-sm);border:2px solid var(--c-border);
  background:var(--c-coral);color:#fff;font-size:13px;font-weight:700;font-family:var(--ff-t);cursor:pointer}
.reset-confirm-ok:active,.reset-confirm-cancel:active{opacity:.7}
/* Toggle switch */
.toggle{position:relative;width:50px;height:28px;flex-shrink:0}
.toggle input{opacity:0;width:0;height:0;position:absolute}
.toggle-track{position:absolute;inset:0;background:var(--c-muted2);
  border-radius:14px;cursor:pointer;transition:background .22s}
.toggle input:checked + .toggle-track{background:var(--c-ok-mark)}
.toggle-thumb{position:absolute;left:3px;top:3px;width:22px;height:22px;
  background:#fff;border-radius:50%;transition:transform .22s;
  box-shadow:0 1px 3px rgba(0,0,0,.3)}
.toggle input:checked ~ .toggle-thumb{transform:translateX(22px)}
/* About card */
.about-card{background:var(--c-card);border-radius:var(--r-lg);padding:18px 16px;
  border:var(--bw) solid var(--c-border);box-shadow:var(--shadow-s);transition:background .25s,border-color .25s}
.about-title{font-family:var(--ff-d);font-size:17px;font-weight:800;
  color:var(--c-ink);margin-bottom:10px;transition:color .25s}
.about-body{font-family:var(--ff-t);font-size:14px;line-height:1.6;
  color:var(--c-muted);transition:color .25s}
.stats-grid{display:grid;grid-template-columns:1fr 1fr;gap:10px}
.stat-card{background:var(--c-card);border-radius:var(--r-md);
  border:var(--bw) solid var(--c-border);box-shadow:var(--shadow-s);
  padding:14px 14px;text-align:center;transition:background .25s}
.stat-n{font-family:var(--ff-d);font-size:26px;font-weight:900;
  color:var(--c-primary);line-height:1.1}
.stat-l{font-family:var(--ff-t);font-size:12px;color:var(--c-muted);margin-top:3px}
/* ── 初回オンボーディング ── */
/* align-items:flex-start＋カードmargin:auto で、収まる時は中央／大きい文字で溢れる時は上から読めてスクロール可（中央寄せの上下切れ対策） */
#onboard{position:fixed;inset:0;z-index:3000;background:var(--c-canvas);
  display:flex;align-items:flex-start;justify-content:center;padding:20px;overflow-y:auto}
.ob-card{width:100%;max-width:420px;margin:auto;background:#fff;border:var(--bw) solid var(--c-border);
  box-shadow:var(--shadow-l);border-radius:var(--r-lg);padding:22px 20px 24px;text-align:center}
.ob-hero{height:96px;width:auto;margin-bottom:6px}
.ob-title{font-family:var(--ff-d);font-size:24px;font-weight:900;color:var(--c-ink)}
.ob-sub{font-size:13px;color:var(--c-muted);margin-bottom:14px}
.ob-q{font-family:var(--ff-d);font-size:14px;font-weight:900;color:var(--c-primary);
  text-align:left;margin:14px 2px 8px}
.ob-opts{display:flex;gap:8px}
.ob-opts-age{flex-wrap:wrap}
.ob-opt{flex:1;min-width:72px;padding:11px 6px;border:var(--bw) solid var(--c-border);
  border-radius:var(--r-md);background:var(--c-parchment);box-shadow:var(--shadow-s);
  font-family:var(--ff-t);font-size:14px;font-weight:800;color:var(--c-ink);cursor:pointer;
  transition:all .12s;-webkit-tap-highlight-color:transparent}
.ob-opt:active{transform:translate(2px,2px);box-shadow:none}
.ob-opt.sel{background:var(--c-accent)}
.ob-next,.ob-ok{width:100%;margin-top:20px;padding:14px;border:var(--bw) solid var(--c-border);
  border-radius:var(--r-pill);background:var(--c-coral);color:#fff;box-shadow:var(--shadow-m);
  font-family:var(--ff-t);font-size:17px;font-weight:800;cursor:pointer;transition:all .12s;
  -webkit-tap-highlight-color:transparent}
.ob-next:active,.ob-ok:active{transform:translate(3px,3px);box-shadow:none}
.ob-next:disabled{opacity:.4;cursor:default;transform:none;box-shadow:var(--shadow-m)}
/* 低い画面（小型スマホ・文字拡大時）はコンパクト表示で1画面に収める */
@media (max-height:760px){
  #onboard{padding:10px}
  .ob-card{padding:14px 16px 16px}
  .ob-hero{height:56px;margin-bottom:2px}
  .ob-title{font-size:20px}
  .ob-sub{margin-bottom:8px}
  .ob-opt{padding:8px 4px}
  .ob-next,.ob-ok{margin-top:12px;padding:11px}
}
.ob-nick{font-family:var(--ff-d);font-size:26px;font-weight:900;color:var(--c-ink);
  background:var(--c-card);border:var(--bw) solid var(--c-border);border-radius:var(--r-md);
  padding:18px 12px;margin:8px 0 12px;min-height:1.4em;overflow-wrap:anywhere;
  animation:char-pop .4s cubic-bezier(.34,1.56,.64,1)}
.ob-reroll{width:100%;padding:11px;border:2px solid var(--c-border);border-radius:var(--r-pill);
  background:var(--c-parchment);box-shadow:var(--shadow-s);font-family:var(--ff-t);
  font-size:15px;font-weight:800;color:var(--c-ink);cursor:pointer;transition:all .12s;
  -webkit-tap-highlight-color:transparent}
.ob-reroll:active{transform:translate(2px,2px);box-shadow:none}
/* ── 設定: サポート等 ── */
.settings-group-label{font-family:var(--ff-t);font-size:12px;font-weight:800;
  color:var(--c-muted);letter-spacing:.06em;padding:6px 6px 0}
.settings-row-val{font-family:var(--ff-t);font-size:15px;color:var(--c-muted);font-variant-numeric:tabular-nums}
.settings-link{display:flex;align-items:center;gap:12px;width:100%;
  padding:14px 16px;background:transparent;border:none;
  border-bottom:1px solid var(--c-hairline);cursor:pointer;text-align:left;
  -webkit-tap-highlight-color:transparent;transition:background .12s}
.settings-link:last-child{border-bottom:none}
.settings-link:active{background:var(--c-canvas)}
.settings-link .sl-ic{width:22px;height:22px;flex-shrink:0;fill:none;
  stroke:currentColor;stroke-width:2;stroke-linecap:round;stroke-linejoin:round}
.settings-link .sl-txt{flex:1;font-family:var(--ff-t);font-size:16px;font-weight:700;color:var(--c-ink)}
.settings-link .sl-chev{color:var(--c-muted2);font-size:22px;font-weight:700;line-height:1}
/* ── ホーム ── */
#tab-home{overflow-y:auto;-webkit-overflow-scrolling:touch}
#tab-home::-webkit-scrollbar{display:none}
.home-inner{padding:16px 16px 24px;display:flex;flex-direction:column;gap:14px}
.home-hero{display:flex;align-items:center;gap:14px;background:var(--c-accent);
  border:var(--bw) solid var(--c-border);box-shadow:var(--shadow-m);
  border-radius:var(--r-lg);padding:14px 16px}
.home-hero img{height:70px;width:auto;flex-shrink:0}
.home-title{font-family:var(--ff-d);font-size:17px;font-weight:900;color:var(--c-ink);line-height:1.25;letter-spacing:.02em}
.home-date{font-size:16px;font-weight:800;color:var(--c-ink);opacity:.72;margin-top:6px}
.home-player{font-size:13px;font-weight:800;color:var(--c-primary);margin-top:4px}
.home-row2{display:grid;grid-template-columns:1fr 1.4fr;gap:12px}
.home-mini{background:var(--c-card);border:var(--bw) solid var(--c-border);
  box-shadow:var(--shadow-s);border-radius:var(--r-md);padding:12px 14px}
.home-mini-n{font-family:var(--ff-d);font-weight:900;font-size:30px;color:var(--c-coral);line-height:1}
.home-mini-u{font-size:14px;margin-left:2px}
.home-mini-l{font-size:11px;font-weight:700;color:var(--c-muted);margin-top:4px}
.home-rank-name{font-family:var(--ff-d);font-weight:900;font-size:22px;color:var(--c-primary);line-height:1.1}
.home-rank-bar{height:8px;background:#fff;border:2px solid var(--c-border);border-radius:5px;overflow:hidden;margin-top:6px}
.home-rank-fill{height:100%;background:var(--c-orange);transition:width .4s}
.home-card{display:block;width:100%;text-align:left;background:var(--c-card);
  border:var(--bw) solid var(--c-border);box-shadow:var(--shadow-s);
  border-radius:var(--r-lg);padding:14px 16px}
.home-word{cursor:pointer;transition:all .12s;-webkit-tap-highlight-color:transparent}
.home-word:active{transform:translate(2px,2px);box-shadow:none}
.home-card-hd{font-family:var(--ff-t);font-size:13px;font-weight:800;color:var(--c-muted);margin-bottom:6px}
.home-word-term{font-family:var(--ff-d);font-size:24px;font-weight:900;color:var(--c-ink);line-height:1.2;overflow-wrap:anywhere}
.home-word-sub{font-size:14px;color:var(--c-ink);opacity:.8;line-height:1.5;margin-top:4px;
  display:-webkit-box;-webkit-line-clamp:3;-webkit-box-orient:vertical;overflow:hidden}
.home-trivia-body{font-size:15px;line-height:1.6;color:var(--c-ink)}
.home-stats-card{display:flex;align-items:center;gap:16px;padding:16px 18px;
  border:var(--bw) solid var(--c-border);border-radius:var(--r-lg);box-shadow:var(--shadow-m);
  cursor:pointer;text-align:left;background:#ffe9b0;transition:all .12s;-webkit-tap-highlight-color:transparent}
.home-stats-card:active{transform:translate(3px,3px);box-shadow:none}
.home-stats-card img{height:64px;width:64px;object-fit:contain;flex-shrink:0}
/* 成績・ランキング画面 */
#tab-stats{overflow-y:auto}
.stats-hdr2{padding:16px 14px 8px;background:var(--c-canvas);border-bottom:1px solid var(--c-hairline)}
.stats-hdr2 h1{font-family:var(--ff-d);font-size:26px;font-weight:900}
.stats-scroll{padding:14px;display:flex;flex-direction:column;gap:14px}
.stats-card2{background:var(--c-card);border:2px solid var(--c-border);border-radius:var(--r-lg);
  box-shadow:var(--shadow-s);padding:14px}
.stats-card-ttl{font-family:var(--ff-d);font-weight:900;font-size:16px;color:var(--c-ink);margin-bottom:10px}
.rank-you{font-family:var(--ff-d);font-weight:900;font-size:18px;color:var(--c-coral);text-align:center;margin-bottom:10px}
.rank-board{display:flex;flex-direction:column;gap:2px}
.rank-row{display:flex;align-items:center;gap:10px;padding:7px 10px;border-radius:var(--r-md);
  font-family:var(--ff-t);font-size:14px}
.rank-row:nth-child(odd){background:var(--c-score-bg)}
.rank-row.me{background:var(--c-accent);font-weight:800}
.rank-pos{width:2.2em;text-align:center;font-family:var(--ff-d);font-weight:900;color:var(--c-muted)}
.rank-row.me .rank-pos{color:var(--c-coral)}
.rank-nm{flex:1;min-width:0;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
.rank-sc{font-family:var(--ff-d);font-weight:900;color:var(--c-ink)}
.rank-sep{text-align:center;color:var(--c-muted2);font-weight:900;line-height:1;padding:2px 0}
.rank-note{font-size:11px;color:var(--c-muted2);text-align:center;margin-top:8px}
.rank-hist-btn{width:100%;margin-top:10px;padding:9px 10px;border:2px dashed var(--c-border);
  border-radius:var(--r-md);background:var(--c-score-bg);color:var(--c-primary);
  font-family:var(--ff-d);font-weight:900;font-size:13px;cursor:pointer}
#rank-history{margin-top:8px;display:flex;flex-direction:column;gap:10px}
.rh-day{border:1.5px solid var(--c-border);border-radius:var(--r-md);padding:9px 11px;background:var(--c-card)}
.rh-date{font-family:var(--ff-d);font-weight:900;font-size:13px;color:var(--c-ink);margin-bottom:5px}
.rh-line{font-family:var(--ff-t);font-size:13px;display:flex;gap:8px;padding:2px 0}
.rh-line .rh-pos{width:2.2em;text-align:center;font-family:var(--ff-d);font-weight:900;color:var(--c-muted)}
.rh-line .rh-nm{flex:1;min-width:0;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
.rh-line .rh-sc{font-family:var(--ff-d);font-weight:900}
.rh-line.me{color:var(--c-coral);font-weight:800}
.rh-me{margin-top:5px;font-family:var(--ff-t);font-size:12px;color:var(--c-primary);font-weight:800}
.rh-empty{font-size:12px;color:var(--c-muted2);text-align:center;padding:6px 0}
.rec-grid{display:grid;grid-template-columns:1fr 1fr;gap:10px}
.rec-item{background:var(--c-score-bg);border-radius:var(--r-md);padding:12px 6px;text-align:center}
.rec-n{font-family:var(--ff-d);font-weight:900;font-size:22px;color:var(--c-coral)}
.rec-l{font-family:var(--ff-t);font-size:12px;color:var(--c-muted);margin-top:2px}
.rec-rank-row{margin-top:10px;font-family:var(--ff-t);font-size:13px;color:var(--c-ink);text-align:center}
.rec-rank-row b{color:var(--c-primary);font-family:var(--ff-d)}
.home-actions{display:grid;grid-template-columns:1fr 1fr;gap:12px}
.home-act{display:flex;flex-direction:column;align-items:center;gap:4px;padding:12px;
  border:var(--bw) solid var(--c-border);box-shadow:var(--shadow-m);border-radius:var(--r-lg);
  font-family:var(--ff-d);font-weight:900;font-size:16px;color:var(--c-ink);cursor:pointer;
  transition:all .12s;-webkit-tap-highlight-color:transparent}
.home-act:active{transform:translate(3px,3px);box-shadow:none}
.home-act img{height:56px;width:auto}
.home-act-quiz{background:var(--c-accent)}
.home-act-daily{background:var(--c-orange)}
/* ── 日本人選手・この日のMLB史 ── */
.home-jp-cover{display:flex;flex-direction:column;align-items:center;gap:4px;width:100%;
  padding:15px 12px;border:2px solid var(--c-border);border-radius:var(--r-md);
  background:var(--c-accent);box-shadow:var(--shadow-s);cursor:pointer;
  -webkit-tap-highlight-color:transparent;transition:all .12s}
.home-jp-cover:active{transform:translate(2px,2px);box-shadow:none}
.jp-cover-lock{font-family:var(--ff-d);font-weight:900;font-size:15px;color:var(--c-ink)}
.jp-cover-note{font-size:11.5px;color:var(--c-ink);opacity:.6}
.home-jp-body{font-size:14px;color:var(--c-muted)}
.jp-row{display:flex;align-items:baseline;gap:8px;padding:6px 0;border-bottom:1px dashed var(--c-hairline)}
.jp-row:last-child{border-bottom:none}
.jp-name{flex-shrink:0;font-family:var(--ff-t);font-weight:800;color:var(--c-ink);font-size:14px;min-width:5.2em}
.jp-line{flex:1;color:var(--c-ink);opacity:.9;font-size:13.5px;line-height:1.45;overflow-wrap:anywhere}
.jp-line .jp-date{color:var(--c-muted2);font-size:12px;margin-right:5px}
.jp-line .jp-good{color:var(--c-coral);font-weight:800}
.jp-note{font-size:12px;color:var(--c-muted2);padding:2px 0}
.home-otd-body{font-size:14px;color:var(--c-ink);line-height:1.55;opacity:.92}
.home-otd-body .otd-year{font-family:var(--ff-d);font-weight:900;color:var(--c-primary);margin-right:6px}
/* ── 今日の試合・順位表 ── */
.home-games-hd{display:flex;align-items:center;justify-content:space-between;gap:8px}
.games-refresh{font-family:var(--ff-t);font-size:10.5px;font-weight:600;color:var(--c-muted2);white-space:nowrap}
.games-refresh-btn{margin-left:4px;padding:3px 9px;border:1.5px solid var(--c-border);border-radius:var(--r-pill);
  background:#fff;font-family:var(--ff-t);font-size:11px;font-weight:800;color:var(--c-primary);
  cursor:pointer;white-space:nowrap;-webkit-tap-highlight-color:transparent;transition:all .12s}
.games-refresh-btn:active:not(:disabled){transform:translate(1px,1px)}
.games-refresh-btn:disabled{opacity:.6;cursor:default}
.standings-btn{padding:4px 12px;border:2px solid var(--c-border);border-radius:var(--r-pill);
  background:var(--c-accent);box-shadow:var(--shadow-s);font-family:var(--ff-t);
  font-size:12px;font-weight:800;color:var(--c-ink);cursor:pointer;
  transition:all .12s;-webkit-tap-highlight-color:transparent}
.standings-btn:active{transform:translate(2px,2px);box-shadow:none}
.home-games-body{font-size:14px;color:var(--c-muted)}
.home-games-note{font-size:10.5px;color:var(--c-muted2);margin-top:8px;text-align:center;letter-spacing:.3px}
/* 昨日の結果を枠で囲って「今日の試合」と視覚的に分離 */
#yday-games:not(:empty){border:2px solid var(--c-border);border-radius:var(--r-md);background:var(--c-score-bg);padding:8px 10px;margin:6px 0 10px}
.yday-sep{font-family:var(--ff-d);font-weight:900;font-size:13px;color:var(--c-ink);margin:2px 0 6px}
.game-row{display:flex;align-items:center;gap:8px;padding:7px 0;border-bottom:1px dashed var(--c-hairline)}
.game-row:last-child{border-bottom:none}
.game-reveal{padding:5px 12px;border:2px dashed var(--c-border);border-radius:var(--r-pill);
  background:var(--c-score-bg);font-family:var(--ff-t);font-size:12px;font-weight:800;
  color:var(--c-muted);cursor:pointer;-webkit-tap-highlight-color:transparent;transition:all .12s;white-space:nowrap}
.game-reveal:active{transform:translate(2px,2px)}
.games-reveal-all{display:block;width:100%;margin-top:10px;padding:9px 0;
  border:2px solid var(--c-border);border-radius:var(--r-pill);background:var(--c-accent);
  box-shadow:var(--shadow-s);font-family:var(--ff-t);font-size:13px;font-weight:800;
  color:var(--c-ink);cursor:pointer;-webkit-tap-highlight-color:transparent;transition:all .12s}
.games-reveal-all:active{transform:translate(2px,2px);box-shadow:none}
.games-reveal-all-top{margin-top:0;margin-bottom:8px}
.game-team{flex:1;display:flex;align-items:center;gap:6px;min-width:0;
  font-weight:700;color:var(--c-ink);font-size:13.5px}
.game-team.away{justify-content:flex-end}
/* 文字拡大時もチーム名が「…」で切れないよう折り返す（高齢者の大きい文字対応） */
.game-team-name{white-space:normal;overflow:visible;word-break:break-word;line-height:1.2;min-width:0}
.game-team.away .game-team-name{text-align:right}
.team-logo{width:22px;height:22px;flex-shrink:0;object-fit:contain}
.game-mid{flex-shrink:0;text-align:center;min-width:74px}
.game-dh{display:block;font-size:9.5px;font-weight:800;color:var(--c-primary);
  background:var(--c-card);border:1.5px solid var(--c-border);border-radius:var(--r-pill);
  padding:0 6px;margin:0 auto 2px;width:fit-content}
.game-score{font-family:var(--ff-d);font-weight:900;font-size:17px;color:var(--c-ink)}
.game-state{font-size:10px;font-weight:800;border-radius:var(--r-pill);padding:1px 8px;display:inline-block}
.game-state.live{background:var(--c-coral);color:#fff}
.game-state.final{background:var(--c-primary);color:#fff}
.game-state.pre{background:var(--c-parchment);border:1.5px solid var(--c-border);color:var(--c-ink)}
#tab-standings{overflow-y:auto;-webkit-overflow-scrolling:touch}
#tab-standings::-webkit-scrollbar{display:none}
.standings-body{padding:14px 16px 28px;display:flex;flex-direction:column;gap:16px;color:var(--c-muted)}
.stand-div{background:var(--c-parchment);border:var(--bw) solid var(--c-border);
  box-shadow:var(--shadow-s);border-radius:var(--r-lg);overflow:hidden}
.stand-div-hd{background:var(--c-accent);border-bottom:2px solid var(--c-border);
  font-family:var(--ff-d);font-weight:900;font-size:15px;color:var(--c-ink);padding:7px 14px}
.stand-table{width:100%;border-collapse:collapse;font-size:13px}
.stand-table th{font-size:10.5px;font-weight:800;color:var(--c-muted);text-align:right;padding:6px 8px 2px;letter-spacing:.5px}
.stand-table th:first-child{text-align:left;padding-left:14px}
.stand-table td{padding:6px 8px;text-align:right;color:var(--c-ink);border-top:1px dashed var(--c-hairline);
  font-variant-numeric:tabular-nums}
.stand-table td:first-child{text-align:left;padding-left:14px;font-weight:700;white-space:nowrap}
.stand-team{display:flex;align-items:center;gap:6px}
.stand-team .team-logo{width:20px;height:20px}
</style>
</head>
<body>
<div id="app">

  <!-- ── ホーム ── -->
  <div id="tab-home" class="tab-panel active">
    <div class="home-inner">
      <div class="home-hero">
        <img src="images/よろこび.png" alt="">
        <div class="home-hero-txt">
          <div class="home-title">メジャーリーグ野球用語<br>英和・和英辞典 ＆ MLB情報・学習</div>
          <div class="home-date" id="home-date"></div>
          <div class="home-player" id="home-player"></div>
        </div>
      </div>

      <div class="home-row2">
        <div class="home-mini home-streak">
          <div class="home-mini-n"><span id="home-streak-n">0</span><span class="home-mini-u">日</span></div>
          <div class="home-mini-l">連続ログイン</div>
        </div>
        <div class="home-mini home-rank">
          <div class="home-rank-name" id="home-rank">ルーキー</div>
          <div class="home-mini-l">称号（得点 <span id="home-runs">0</span>）</div>
          <div class="home-rank-bar"><div class="home-rank-fill" id="home-rank-fill"></div></div>
          <div class="home-mini-l" id="home-rank-next" style="margin-top:4px"></div>
        </div>
      </div>

      <div class="home-card home-jp-card">
        <div class="home-card-hd">🇯🇵 今日の日本人選手</div>
        <button type="button" class="home-jp-cover" id="home-jp-cover" onclick="revealJpPlayers()">
          <span class="jp-cover-lock">👀 タップして成績を表示</span>
          <span class="jp-cover-note">ネタバレ注意（各選手の個人成績が表示されます）</span>
        </button>
        <div id="home-jp" class="home-jp-body" style="display:none">読み込み中…</div>
      </div>

      <button class="home-card home-word" id="home-word" onclick="openWordOfDay()">
        <div class="home-card-hd">📖 今日の一語</div>
        <div class="home-word-term" id="home-word-term">—</div>
        <div class="home-word-sub" id="home-word-sub"></div>
      </button>

      <div class="home-card home-trivia">
        <div class="home-card-hd">💡 今日の豆知識</div>
        <div class="home-trivia-body" id="home-trivia">—</div>
      </div>

      <div class="home-card home-otd-card">
        <div class="home-card-hd">📅 この日のMLB史</div>
        <div id="home-otd" class="home-otd-body">—</div>
      </div>

      <div class="home-card home-games-card">
        <div class="home-card-hd home-games-hd"><span>⚾ 今日の試合<span class="games-refresh" id="games-usdate"></span> <button class="games-refresh-btn" id="games-refresh-btn" onclick="refreshGames()">🔄 最新に更新</button></span>
          <button class="standings-btn" onclick="goTab('standings')">順位表 →</button>
        </div>
        <button type="button" class="games-reveal-all games-reveal-all-top" id="yday-btn" onclick="toggleYdayGames()">📅 昨日の結果を見る ▼</button>
        <div id="yday-games" style="display:none"></div>
        <div id="games-revealall-top"></div>
        <div id="home-games" class="home-games-body">読み込み中…</div>
        <div class="home-games-note">米国日付の本日の試合／開始時刻は日本時間／結果は「結果」をタップで表示（ネタバレ防止）</div>
      </div>

    </div>
  </div>

  <!-- ── 順位表 ── -->
  <div id="tab-standings" class="tab-panel">
    <div class="fav-hdr subpage-hdr"><button class="back-btn" onclick="goTab('home')">← ホーム</button><h1>順位表</h1></div>
    <div id="standings-body" class="standings-body">読み込み中…</div>
  </div>

  <!-- ── 辞書 ── -->
  <div id="tab-dict" class="tab-panel">
    <div class="app-title">
      <h1><img src="images/デフォ.png" alt="">メジャーリーグ野球用語英和・和英辞典</h1>
    </div>
    <div class="search-bar">
      <div class="search-row">
        <div class="search-wrap">
          <svg viewBox="0 0 20 20"><path fill-rule="evenodd" d="M8 4a4 4 0 100 8 4 4 0 000-8zM2 8a6 6 0 1110.89 3.476l4.817 4.817a1 1 0 01-1.414 1.414l-4.816-4.816A6 6 0 012 8z" clip-rule="evenodd" stroke="none" fill="currentColor"/></svg>
          <input id="s-input" type="text" placeholder="検索…" autocomplete="off" autocorrect="off" autocapitalize="none" spellcheck="false">
          <button class="s-clear" id="s-clear" onclick="clearSearch()">✕</button>
        </div>
        <button class="abbr-jump-btn" onclick="goTab('abbr')">略語</button>
      </div>
    </div>
    <div class="sub-tabs">
      <button class="sub-tab active" id="dict-tab-ej" onclick="setDictMode('ej')">英和 (EN→JA)</button>
      <button class="sub-tab" id="dict-tab-je" onclick="setDictMode('je')">和英 (JA→EN)</button>
      <button class="sub-tab sub-tab-fav" id="dict-tab-fav" onclick="toggleDictFav()" aria-label="お気に入り"><svg viewBox="0 0 24 24"><path d="M20.84 4.61a5.5 5.5 0 00-7.78 0L12 5.67l-1.06-1.06a5.5 5.5 0 00-7.78 7.78l1.06 1.06L12 21.23l7.78-7.78 1.06-1.06a5.5 5.5 0 000-7.78z"/></svg></button>
    </div>
    <div id="fav-back-bar" class="fav-back-bar" style="display:none">
      <button class="back-btn" onclick="exitDictFav()">← 辞書に戻る</button>
      <span class="fav-back-title">お気に入り</span>
    </div>
    <div class="dict-body">
      <div class="entry-list" id="dict-list"></div>
      <div id="idx-bar"></div>
    </div>
  </div>

  <!-- ── 略語 ── -->
  <div id="tab-abbr" class="tab-panel">
    <div class="fav-hdr subpage-hdr"><button class="back-btn" onclick="goTab('dict')">← 辞書</button><h1>統計略語辞典</h1></div>
    <div class="search-bar">
      <div class="search-wrap">
        <svg viewBox="0 0 20 20"><path fill-rule="evenodd" d="M8 4a4 4 0 100 8 4 4 0 000-8zM2 8a6 6 0 1110.89 3.476l4.817 4.817a1 1 0 01-1.414 1.414l-4.816-4.816A6 6 0 012 8z" clip-rule="evenodd" stroke="none" fill="currentColor"/></svg>
        <input id="abbr-search" type="text" placeholder="略語を検索…" autocomplete="off" autocorrect="off" autocapitalize="none" spellcheck="false">
        <button class="s-clear" id="abbr-clear" onclick="clearAbbrSearch()">✕</button>
      </div>
    </div>
    <div class="dict-body">
      <div class="entry-list" id="abbr-list"></div>
      <div id="abbr-idx-bar"></div>
    </div>
  </div>

  <!-- ── お気に入りは辞書内の♥サブタブへ統合 ── -->
  <div id="tab-fav-unused" class="tab-panel" style="display:none">
  </div>

  <!-- ── クイズ ── -->
  <div id="tab-quiz" class="tab-panel">

    <!-- 目次 -->
    <div id="quiz-menu">
      <div class="qmenu-title-row"><img src="images/よろこび.png" alt=""><span>あそぶ・まなぶ</span>
        <span class="qmenu-runs">本日 <b id="menu-runs-today">0</b> 点 ／ 累計 <b id="menu-runs">0</b> 点</span></div>
      <div class="qmenu-group-label">📚 学ぶ</div>
      <button class="qmenu-card qmenu-quiz" onclick="startQuizPlay()">
        <img src="images/バッター.png" alt="">
        <span class="qmenu-txt"><span class="qmenu-name">フリーバッティング</span><span class="qmenu-sub">英和・和英ミックス！自己ベストを伸ばそう</span></span>
      </button>
      <button class="qmenu-card qmenu-daily" onclick="goTab('daily')">
        <img src="images/ピッチャー.png" alt="">
        <span class="qmenu-txt"><span class="qmenu-name">今日の10問</span><span class="qmenu-sub">毎日コツコツ10問チャレンジ</span></span>
      </button>
      <button class="qmenu-card qmenu-review" onclick="goTab('review')">
        <img src="images/デフォ.png" alt="">
        <span class="qmenu-txt"><span class="qmenu-name">復習</span><span class="qmenu-sub">間違えた問題をおさらい</span></span>
      </button>
      <div class="qmenu-group-label">🔥 勝負</div>
      <button class="qmenu-card qmenu-versus" onclick="startVersus()">
        <img src="images/ガッツ.png" alt="">
        <span class="qmenu-txt"><span class="qmenu-name">試合</span><span class="qmenu-sub">9回勝負！勝てばボーナス</span></span>
      </button>
      <div class="qmenu-howto">
        <button class="qmenu-howto-ttl" onclick="toggleStats()">🏆 今日の成績ランキング（夜12時リセット） <span class="howto-chev" id="stats-chev">▼</span></button>
        <div id="stats-body" style="display:none">
          <div class="rank-you" id="rank-you">—</div>
          <div class="rank-board" id="rank-board"></div>
          <button class="rank-hist-btn" onclick="toggleRankHistory()">📅 過去のランキングを見る（1週間） <span id="rank-hist-chev">▼</span></button>
          <div id="rank-history" style="display:none"></div>
          <div class="rec-grid">
            <div class="rec-item"><div class="rec-n" id="rec-today">0</div><div class="rec-l">本日の得点</div></div>
            <div class="rec-item"><div class="rec-n" id="rec-total">0</div><div class="rec-l">累計得点</div></div>
            <div class="rec-item"><div class="rec-n" id="rec-vs">0-0</div><div class="rec-l">対戦成績</div></div>
            <div class="rec-item"><div class="rec-n" id="rec-streak">0</div><div class="rec-l">連続ログイン</div></div>
          </div>
          <div class="rec-rank-row">称号：<b id="rec-rank-name">ルーキー</b> <span id="rec-rank-next"></span></div>
          <div class="home-rank-bar" style="margin:6px 4px 0"><div class="home-rank-fill" id="rec-rank-fill"></div></div>
          <div class="rank-note">本日の得点でランキング（毎日リセット）／累計得点で称号が上がります</div>
        </div>
      </div>
      <div class="qmenu-howto">
        <button class="qmenu-howto-ttl" id="howto-toggle" onclick="toggleHowto()">⚾ 点の入り方 <span class="howto-chev" id="howto-chev">▼</span></button>
        <div id="howto-body" style="display:none">
          <div class="howto-swingrow">
            <span class="howto-swing-lbl">スイング操作</span>
            <button class="swing-toggle" id="swing-toggle" onclick="toggleSwingAuto()">手動</button>
          </div>
          <div class="howto-swing-hint" id="howto-swing-hint"></div>
          <ul id="howto-manual">
            <li>正解すると打席へ。バーの<b>赤</b>で止めると長打！（<b>黄</b>＝単打）</li>
            <li><b>往路</b>＝ホームラン／<b>復路</b>＝二塁打（難問は三塁打）／それ以降＝単打</li>
            <li>走者が<b>ホームインで得点</b>。不正解や凡打はアウト、<b>3アウトで終了</b>。</li>
          </ul>
          <ul id="howto-auto" style="display:none">
            <li>正解すると<b>自動で単打</b>（タイミング操作なし）。</li>
            <li><b>長打・本塁打はねらえません</b>（狙うなら手動）。堅実だが天井は低め。</li>
            <li>走者が<b>ホームインで得点</b>。不正解はアウト、<b>3アウトで終了</b>。</li>
          </ul>
          <div class="howto-sub-ttl">📅 今日の10問</div>
          <ul>
            <li>毎日<b>10問</b>に挑戦（1日1回）。スイングはなく<b>正解／不正解のみ</b>。</li>
            <li><b>10問全問正解＝＋5点</b>／<b>6〜9問正解＝＋3点</b>（本日・累計に加算）。</li>
          </ul>
        </div>
      </div>
      <div class="qmenu-score">
        <div class="score-correct-hdr">正解数</div>
        <div class="score-cols">
          <div class="score-col">
            <div class="score-lbl">英和</div>
            <div class="score-val-row">
              <div class="score-val" id="sc-ej">0 / __EJ_COUNT__</div>
              <div class="score-rate" id="sc-ej-rate">--%</div>
            </div>
          </div>
          <div class="score-col">
            <div class="score-lbl">和英</div>
            <div class="score-val-row">
              <div class="score-val" id="sc-je">0 / __JE_COUNT__</div>
              <div class="score-rate" id="sc-je-rate">--%</div>
            </div>
          </div>
        </div>
      </div>
    </div>

    <!-- プレイ画面 -->
    <div id="quiz-play" style="display:none">
    <div class="quiz-inner">
      <div class="subpage-bar"><button class="back-btn" onclick="quizBack()">← 目次</button></div>
      <div id="versus-bar" style="display:none"></div>
      <div id="versus-inning" style="display:none"></div>
      <div class="quiz-char" id="quiz-char">
        <div class="diamond-wrap">
          <img class="diamond-img" src="images/ベース.png" alt="">
          <img class="dchar dchar-home" id="quiz-char-img" src="images/デフォ.png" alt="">
          <img class="dchar dchar-runner" id="dchar-1" src="images/盗塁待ち.png" alt="" style="display:none">
          <img class="dchar dchar-runner" id="dchar-2" src="images/盗塁待ち.png" alt="" style="display:none">
          <img class="dchar dchar-runner" id="dchar-3" src="images/盗塁待ち.png" alt="" style="display:none">
          <div class="diamond-hud diamond-hud-out"><span class="quiz-outs" id="quiz-outs">OUT ○○○</span></div>
          <div class="diamond-hud diamond-hud-run" id="hud-runs">本日 <span id="quiz-runs-today">0</span>／通算 <span id="quiz-runs">0</span></div>
        </div>
      </div>
      <div class="swing-ui" id="swing-ui" style="display:none">
        <div class="swing-pitch" id="swing-pitch">Lv1 直球</div>
        <div class="swing-track">
          <div class="swing-zone swing-zone-good" id="sz-good"></div>
          <div class="swing-zone swing-zone-great" id="sz-great"></div>
          <div class="swing-zone swing-zone-perfect" id="sz-perfect"></div>
          <div class="swing-marker" id="swing-marker"></div>
        </div>
        <div class="swing-guide" id="swing-guide"></div>
        <button class="swing-btn" id="swing-btn" onclick="doSwing()">打て！</button>
        <div class="swing-result" id="swing-result"></div>
      </div>
      <div class="quiz-card">
        <div class="pitch-clock"><div class="pitch-clock-fill" id="pitch-clock-fill"></div></div>
        <div class="quiz-q">
          <div class="quiz-q-lbl" id="qq-lbl">問題</div>
          <div class="quiz-q-term" id="qq-term">—</div>
        </div>
        <div class="choices" id="choices"></div>
        <button class="next-btn" id="next-btn" disabled onclick="quizNext()">次の問題 →</button>
        <div id="quiz-def-card" style="display:none"></div>
      </div>
    </div>
    </div>

    <!-- スリーアウト演出 -->
    <div id="three-out"><div class="three-out-inner"><span class="three-out-txt">THREE OUT!</span><img class="three-out-img" src="images/うなだれ.png" alt=""></div></div>

    <!-- 対戦: イニング結果ボード -->
    <div id="versus-board"><div id="versus-board-inner"></div></div>
    <!-- 対戦: 開始トースト -->
    <div id="versus-toast"></div>
    <div id="rankup-toast"></div>
    <!-- 対戦: 試合結果 -->
    <div id="versus-result"><div class="vs-res-inner">
      <img class="vs-res-char" src="images/よろこび.png" alt="">
      <div class="vs-res-ttl">勝利！</div>
      <div class="vs-res-score">あなた 0 - 0 相手</div>
      <div class="vs-res-bonus"></div>
      <button class="vs-btn" onclick="closeVersusResult()">目次へ戻る</button>
    </div></div>
  </div>

  <!-- ── 広告（プレースホルダ／ネイティブで実広告に差し替え） ── -->
  <div id="ad-overlay">
    <div class="ad-box">
      <div class="ad-tag">広告</div>
      <div class="ad-body" id="ad-body">
        <img src="images/デフォ.png" alt="">
        <div class="ad-msg">広告スペース</div>
        <div class="ad-note">実際の広告はアプリ版で表示されます</div>
      </div>
      <button class="ad-close" id="ad-close" onclick="closeAd()" disabled>とじる（<span id="ad-count">5</span>）</button>
      <button class="ad-premium" onclick="goTab('settings');closeAd()">広告を消す（プレミアム）</button>
    </div>
  </div>

  <!-- ── 今日の10問 ── -->
  <div id="tab-daily" class="tab-panel">
    <div class="daily-inner">
      <div class="subpage-bar"><button class="back-btn" onclick="quizMenu()">← 目次</button><span class="subpage-title">今日の10問</span><span class="daily-runs-lbl">⚾ 本日 <span id="daily-runs-today">0</span> ／ 累計 <span id="daily-runs">0</span></span></div>

      <!-- スタート画面 -->
      <div id="daily-start" style="display:none">
        <div class="daily-start-inner">
          <div class="daily-start-icon"><img src="images/バッター.png" alt=""></div>
          <div class="daily-start-title">今日の10問</div>
          <div class="daily-start-sub">毎日10問でコツコツ学ぼう</div>
          <button class="daily-start-btn" onclick="startDailyQuiz()">スタート ▶</button>
        </div>
      </div>

      <div class="quiz-char" id="daily-char" style="display:none"><img id="daily-char-img" src="images/デフォ.png" alt="character"></div>
      <div class="daily-progress-bar" id="daily-prog-bar">
        <div class="daily-progress-track"><div class="daily-progress-fill" id="daily-prog-fill" style="width:0%"></div></div>
        <div class="daily-progress-lbl" id="daily-prog-lbl">0 / 10</div>
      </div>
      <div class="quiz-card" id="daily-quiz-area">
        <div class="quiz-q">
          <div class="quiz-q-lbl" id="daily-qq-lbl">問題</div>
          <div class="quiz-q-term" id="daily-qq-term">—</div>
        </div>
        <div class="choices" id="daily-choices"></div>
        <button class="next-btn" id="daily-next-btn" disabled onclick="dailyNextQ()">次の問題 →</button>
        <div id="daily-def-card" style="display:none"></div>
      </div>
      <div class="daily-result" id="daily-result" style="display:none">
        <div class="daily-result-hero">
          <img src="images/よろこび.png" alt="" class="daily-result-char">
          <div style="position:relative;display:inline-block">
            <div class="daily-result-score" id="daily-result-score">0 / 10</div>
            <span style="position:absolute;right:-1.2em;bottom:0.15em;font-size:16px;font-weight:700;color:var(--c-ink)">問</span>
          </div>
          <div class="daily-result-sub">今日の10問 完了！</div>
          <div class="daily-bonus" id="daily-bonus" style="display:none"></div>
          <div class="daily-done-note">明日またチャレンジしよう</div>
        </div>
        <div class="daily-cal">
          <div class="cal-inner">
            <div class="cal-nav">
              <button class="cal-nav-btn" onclick="calPrev()">← 前の月</button>
              <div class="cal-title" id="cal-title"></div>
              <button class="cal-nav-btn" onclick="calNext()">次の月 →</button>
            </div>
            <div class="cal-grid" id="cal-grid"></div>
          </div>
        </div>
      </div>
      <div class="daily-cal" id="daily-cal-area" style="display:none">
        <div class="cal-inner">
          <div class="cal-nav">
            <button class="cal-nav-btn" onclick="calPrev()">←</button>
            <div class="cal-title" id="cal-title2"></div>
            <button class="cal-nav-btn" onclick="calNext()">→</button>
          </div>
          <div class="cal-grid" id="cal-grid2"></div>
        </div>
      </div>
    </div>
  </div>

  <!-- ── 復習リスト ── -->
  <div id="tab-review" class="tab-panel">
    <div class="review-hdr subpage-hdr"><button class="back-btn" onclick="quizMenu()">← 目次</button><h1>復習リスト</h1><button class="review-clear-btn" id="review-clear-btn" onclick="askClearReview()">全削除</button></div>
    <div class="reset-confirm" id="review-clear-confirm">
      <span class="reset-confirm-msg">復習リストをすべて削除しますか？</span>
      <div class="reset-confirm-btns">
        <button class="reset-confirm-cancel" onclick="cancelClearReview()">いいえ</button>
        <button class="reset-confirm-ok" onclick="clearReview()">はい</button>
      </div>
    </div>
    <div id="review-list"></div>
  </div>

  <!-- ── 設定 ── -->
  <div id="tab-settings" class="tab-panel">
    <div class="settings-inner">
      <div class="settings-hdr">設定</div>

      <div class="stats-grid">
        <div class="stat-card">
          <div class="stat-n" id="stat-ej">__EJ_COUNT__</div>
          <div class="stat-l">英和エントリ数</div>
        </div>
        <div class="stat-card">
          <div class="stat-n" id="stat-je">__JE_COUNT__</div>
          <div class="stat-l">和英エントリ数</div>
        </div>
      </div>

      <div class="settings-section">
        <div class="settings-row">
          <div class="entry-text">
            <div class="settings-row-label">問題スコアをリセット</div>
            <div class="settings-row-sub" id="quiz-score-sub">英和: 0正解 / 0問　和英: 0正解 / 0問</div>
          </div>
          <button class="reset-btn" id="quiz-reset-btn" onclick="askResetQuiz()">リセット</button>
        </div>
        <div class="reset-confirm" id="quiz-reset-confirm">
          <span class="reset-confirm-msg">スコアをリセットしますか？</span>
          <div class="reset-confirm-btns">
            <button class="reset-confirm-cancel" onclick="cancelResetQuiz()">キャンセル</button>
            <button class="reset-confirm-ok" onclick="resetQuizScore()">リセットする</button>
          </div>
        </div>
      </div>

      <div class="settings-group-label">問題</div>
      <div class="settings-section">
        <div class="settings-row">
          <div class="entry-text">
            <div class="settings-row-label">効果音</div>
            <div class="settings-row-sub">ヒットの「カキーン」などの野球効果音</div>
          </div>
          <button class="swing-toggle" id="sfx-toggle" onclick="toggleSfx()">オン</button>
        </div>
      </div>

      <div class="settings-group-label">プレミアム</div>
      <div class="settings-section">
        <div class="settings-row">
          <div class="entry-text">
            <div class="settings-row-label">プレミアム（月額200円）</div>
            <div class="settings-row-sub">広告なし・回数無制限。<span id="premium-status">未登録（無料）</span></div>
          </div>
          <span id="premium-actions"></span>
        </div>
        <button class="settings-link" onclick="restorePremium()">
          <svg class="sl-ic" viewBox="0 0 24 24" style="color:var(--c-primary)"><path d="M21 12a9 9 0 1 1-3-6.7"/><path d="M21 3v6h-6"/></svg>
          <span class="sl-txt">購入を復元</span><span class="sl-chev">›</span>
        </button>
      </div>

      <div class="settings-group-label">通知</div>
      <div class="settings-section">
        <div class="settings-row">
          <div class="entry-text">
            <div class="settings-row-label">デイリー通知</div>
            <div class="settings-row-sub">毎朝8時ごろ「今日の10問」をお知らせ（アプリ版のみ）</div>
          </div>
          <button class="swing-toggle" id="notif-toggle" onclick="toggleDailyNotif()">オフ</button>
        </div>
      </div>

      <div class="settings-group-label">サポート</div>
      <div class="settings-section">
        <button class="settings-link" onclick="rateApp()">
          <svg class="sl-ic" viewBox="0 0 24 24" style="color:#f5a13c"><polygon points="12 2 15 9 22 9.3 16.5 14 18.5 21 12 17 5.5 21 7.5 14 2 9.3 9 9"/></svg>
          <span class="sl-txt">アプリを評価する</span><span class="sl-chev">›</span>
        </button>
        <button class="settings-link" onclick="reportBug()">
          <svg class="sl-ic" viewBox="0 0 24 24" style="color:var(--c-coral)"><path d="M4 3v18"/><path d="M4 4h13l-2 4 2 4H4"/></svg>
          <span class="sl-txt">不具合を報告</span><span class="sl-chev">›</span>
        </button>
        <button class="settings-link" onclick="openExt('privacy.html')">
          <svg class="sl-ic" viewBox="0 0 24 24" style="color:#2fbf6b"><path d="M12 2l8 3v6c0 5-3.5 8.5-8 11-4.5-2.5-8-6-8-11V5z"/></svg>
          <span class="sl-txt">プライバシーポリシー</span><span class="sl-chev">›</span>
        </button>
        <button class="settings-link" onclick="openExt('terms.html')">
          <svg class="sl-ic" viewBox="0 0 24 24" style="color:var(--c-accent)"><path d="M12 2l8 3v6c0 5-3.5 8.5-8 11-4.5-2.5-8-6-8-11V5z"/></svg>
          <span class="sl-txt">利用規約</span><span class="sl-chev">›</span>
        </button>
      </div>

      <div class="settings-group-label">このアプリについて</div>
      <div class="settings-section">
        <div class="settings-row"><div class="settings-row-label">バージョン</div><div class="settings-row-val">1.0.0</div></div>
      </div>

      <div class="about-card">
        <div class="about-title">辞書データについて</div>
        <div class="about-body" id="about-body">本アプリの辞書データは、著者の承諾を得て、以下の書籍に基づいて作成しています。<br><br>
          <strong>大リーグ早わかり野球用語<br>英和・和英小辞典</strong><br>
          阿部 達 編著／現代図書<br>
          2006年3月3日発行<br><br>
          原典は2006年発行のため、例文や固有名詞など一部に古い情報が含まれます。新ルール等は改訂しております。</div>
      </div>
    </div>
  </div>

  <!-- ── バナー広告（辞書・略語・お気に入り・今日の10問・復習に表示。プレミアムは非表示） ── -->
  <div id="ad-banner"><span class="ad-banner-tag">広告</span><span class="ad-banner-txt">バナー広告スペース（アプリ版で表示）</span></div>

  <!-- ── ボトムナビ ── -->
  <nav id="bnav">
    <button class="bnav-btn active" id="bnav-home" onclick="goTab('home')" style="--ic:#ff5a4e">
      <svg viewBox="0 0 24 24"><path d="M4 4H20V13.5L12 21L4 13.5Z"/></svg>
      ホーム
    </button>
    <button class="bnav-btn" id="bnav-dict" onclick="goTab('dict')" style="--ic:#3aa0e0">
      <svg viewBox="0 0 24 24"><path d="M4 19.5A2.5 2.5 0 016.5 17H20"/><path d="M6.5 2H20v20H6.5A2.5 2.5 0 014 19.5v-15A2.5 2.5 0 016.5 2z"/></svg>
      辞書
    </button>
    <button class="bnav-btn" id="bnav-quiz" onclick="goQuizNav()" style="--ic:#f5a13c">
      <svg viewBox="0 0 24 24"><circle cx="12" cy="12" r="9"/><path fill="none" d="M6.2 5.6Q9.6 12 6.2 18.4"/><path fill="none" d="M17.8 5.6Q14.4 12 17.8 18.4"/></svg>
      問題
    </button>
    <button class="bnav-btn" id="bnav-settings" onclick="goTab('settings')" style="--ic:#8a8fb0">
      <svg viewBox="0 0 24 24" style="fill:none"><circle cx="12" cy="12" r="3"/><path d="M19.4 15a1.65 1.65 0 00.33 1.82l.06.06a2 2 0 010 2.83 2 2 0 01-2.83 0l-.06-.06a1.65 1.65 0 00-1.82-.33 1.65 1.65 0 00-1 1.51V21a2 2 0 01-2 2 2 2 0 01-2-2v-.09A1.65 1.65 0 009 19.4a1.65 1.65 0 00-1.82.33l-.06.06a2 2 0 01-2.83 0 2 2 0 010-2.83l.06-.06A1.65 1.65 0 004.68 15a1.65 1.65 0 00-1.51-1H3a2 2 0 01-2-2 2 2 0 012-2h.09A1.65 1.65 0 004.6 9a1.65 1.65 0 00-.33-1.82l-.06-.06a2 2 0 010-2.83 2 2 0 012.83 0l.06.06A1.65 1.65 0 009 4.68a1.65 1.65 0 001-1.51V3a2 2 0 012-2 2 2 0 012 2v.09a1.65 1.65 0 001 1.51 1.65 1.65 0 001.82-.33l.06-.06a2 2 0 012.83 0 2 2 0 010 2.83l-.06.06A1.65 1.65 0 0019.4 9a1.65 1.65 0 001.51 1H21a2 2 0 012 2 2 2 0 01-2 2h-.09a1.65 1.65 0 00-1.51 1z"/></svg>
      設定
    </button>
  </nav>

  <!-- ── モーダル背景 ── -->
  <!-- ── 正解/不正解フラッシュ ── -->
  <div id="quiz-flash"><img id="quiz-flash-char" src="images/よろこび.png" alt=""><span id="quiz-flash-mark"></span><span id="quiz-flash-bonus"></span></div>

  <div id="modal-backdrop" onclick="closeModal()"></div>

  <!-- ── ボトムシートモーダル ── -->
  <div id="modal">
    <div class="modal-drag"></div>
    <div class="modal-hdr">
      <button class="modal-fav-btn" id="mfav-btn" onclick="toggleMFav()">
        <svg id="mfav-svg" viewBox="0 0 24 24"><path d="M20.84 4.61a5.5 5.5 0 00-7.78 0L12 5.67l-1.06-1.06a5.5 5.5 0 00-7.78 7.78l1.06 1.06L12 21.23l7.78-7.78 1.06-1.06a5.5 5.5 0 000-7.78z"/></svg>
      </button>
      <button class="modal-close" onclick="closeModal()">✕</button>
    </div>
    <div class="modal-body">
      <div class="modal-term" id="m-term"></div>
      <div class="modal-def" id="m-def"></div>
    </div>
  </div>

  <!-- ── 初回オンボーディング ── -->
  <div id="onboard" style="display:none">
    <div class="ob-card">
      <div id="ob-step1">
        <img class="ob-hero" src="images/よろこび.png" alt="">
        <div class="ob-title">ようこそ！</div>
        <div class="ob-sub">はじめに教えてね</div>
        <div class="ob-q">性別</div>
        <div class="ob-opts" id="ob-gender">
          <button type="button" class="ob-opt" data-v="m" onclick="obSel('gender',this)">男</button>
          <button type="button" class="ob-opt" data-v="f" onclick="obSel('gender',this)">女</button>
          <button type="button" class="ob-opt" data-v="n" onclick="obSel('gender',this)">無回答</button>
        </div>
        <div class="ob-q">年代</div>
        <div class="ob-opts ob-opts-age" id="ob-age">
          <button type="button" class="ob-opt" data-v="10" onclick="obSel('age',this)">10代以下</button>
          <button type="button" class="ob-opt" data-v="20" onclick="obSel('age',this)">20代</button>
          <button type="button" class="ob-opt" data-v="30" onclick="obSel('age',this)">30代</button>
          <button type="button" class="ob-opt" data-v="40" onclick="obSel('age',this)">40代</button>
          <button type="button" class="ob-opt" data-v="50" onclick="obSel('age',this)">50代</button>
          <button type="button" class="ob-opt" data-v="60" onclick="obSel('age',this)">60代以上</button>
        </div>
        <button class="ob-next" id="ob-next" disabled onclick="obToStep2()">つぎへ →</button>
      </div>
      <div id="ob-step2" style="display:none">
        <img class="ob-hero" src="images/バッター.png" alt="">
        <div class="ob-title">ニックネームを選ぼう</div>
        <div class="ob-sub">気に入らなければ引き直してね</div>
        <div class="ob-nick" id="ob-nick">—</div>
        <button class="ob-reroll" onclick="obReroll()">🔄 引き直す</button>
        <button class="ob-ok" onclick="obConfirm()">これにする！</button>
      </div>
    </div>
  </div>
</div>

<script>
// ── Data ─────────────────────────────────────────────────────────────────────
const EJ = __EJ_DATA__;
const JE = __JE_DATA__;
const TRIVIA = __TRIVIA_DATA__;
const ABBR = EJ.filter(e => e.cat === 'abbr');
const EJ_TOTAL = EJ.length;
const JE_TOTAL = JE.length;

// ── State ─────────────────────────────────────────────────────────────────────
let tab='dict', dictMode='ej', favMode='ej', qMode='ej';
let favs=new Set(JSON.parse(localStorage.getItem('mlb_favs')||'[]'));
let modalEntry=null, modalMode='ej';
let sq='';
const qs={ej:{ok:0,n:0},je:{ok:0,n:0}};
let cq=null, ans=false;
// 不正解キュー: {entry, retryAt} の配列（retryAt = 正解するまで出し続ける出題番号）
const retryQueue={ej:[],je:[]};
// 旧テーマ設定の後始末（テーマ機能は廃止）
localStorage.removeItem('mlb_theme');
localStorage.removeItem('mlb_dark');

// ── Util ──────────────────────────────────────────────────────────────────────
const $=id=>document.getElementById(id);
function esc(s){return(s||'').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;').replace(/'/g,'&#39;')}
function favKey(e){return(e.en||'')+'||'+(e.ja||'')+(e.def||'')}
function saveFavs(){localStorage.setItem('mlb_favs',JSON.stringify([...favs]))}
function isFav(e){return favs.has(favKey(e))}
function toggleFav(e){const k=favKey(e);favs.has(k)?favs.delete(k):favs.add(k);saveFavs()}
// ── サポート系リンク ──
const REVIEW_URL='';   // App Store公開後にレビューURLを設定
function openExt(url){ try{ window.open(url,'_blank'); }catch(e){ location.href=url; } }
function rateApp(){ if(REVIEW_URL) openExt(REVIEW_URL); else alert('公開後にApp Storeで評価いただけます。応援ありがとうございます！'); }
function reportBug(){ openExt('https://fantamstick.com/contact'); }
// ── 初回オンボーディング（性別・年代・ニックネーム） ──
const NICK={
  coolPre:['豪快な','剛腕','鉄壁の','快速','電光石火の','不動の','百戦錬磨の','孤高の','精密機械','逆転の','無敵の','伝説の','絶対的','変幻自在の','一撃必殺の','沈黙の','冷静沈着な','熱血','天才','韋駄天','豪速球','超人','神速の','爆裂','無双の','大器','鉄腕','快足の','一匹狼','ミスター'],
  coolSuf:['エース','4番','キャプテン','クローザー','スラッガー','守護神','大砲','遊撃手','二刀流','先発','リードオフマン','主砲','番長','ストッパー','助っ人','強打者','豪腕投手','切り込み隊長','扇の要','鉄人','代打の切り札','ムードメーカー','幻の左腕','いぶし銀','広角打法','韋駄天ランナー','火消し','決め球王'],
  cutePre:['きらめく','ふんわり','ちいさな','にこにこ','ぴょんぴょん','もちもち','ときめき','ゆるふわ','キラキラ','まんまる','ぽかぽか','ふわもこ','ぷにぷに','るんるん','わくわく','ほんわか','すやすや','ふわり','こつこつ','ぱたぱた'],
  cuteSuf:['ベースちゃん','ボールちゃん','グローブちゃん','ホームランガール','ルーキーちゃん','マネージャー','応援団長','かっとばし娘','一年生','ベンチ女子','ミットちゃん','バットちゃん','スパイクちゃん','主将ちゃん','応援女子','球場の妖精'],
  netaM:['三振王','ベンチ温め係','帰宅部エース','二日酔いの4番','遅刻魔クローザー','満塁の申し子','サヨナラ男','敬遠された男','代走のプロ','ヤジ将軍','カレー好き遊撃手','補欠の星','エラーの帝王','肩だけは一流','声だけデカい4番','三日坊主エース','ネクストの主','塁上の哲学者','空振りアーティスト','ビール売りの星','牽制で果てた男','雨天中止を願う男','素振りだけ天才','ヘルメット忘れの常習犯','自打球の被害者','ファウルで粘るだけ','ダグアウトの語り部','伝令だけ全力','円陣を組みたがる男','守備位置わからない','ベンチで一番声出す','グラブにボール入らない','背番号だけ一流','ヒーローインタビュー志望','ガム噛みすぎ4番','ネクストで素振りしすぎ','初球必ず見送る男','走塁ミスの申し子','ベースカバー忘れ番長','サインを見落とす天才','監督より監督な男','日程を勘違いする男','延長戦に消える男','タイムかけすぎ投手','ヒマワリの種係','ベンチプレスだけ4番','背筋だけ守護神','ルール知らない助っ人','応援歌だけ完璧','グラウンド整備が生きがい'],
  netaF:['三振ガール','おやつ番長','昼寝の女王','サボり上手','ベンチアイドル','かき氷担当','応援だけは一流','スコア係','まかない女子','送りバント職人','声出し隊長','補欠アイドル','ベンチの主','現地観戦の達人','グッズ全部買う人','ルーキーのお姉さん','塁間のお散歩','日焼け対策は万全','タオル振るだけ達人','推しの分まで応援','ハイタッチ職人','ヒロイン狙いの女子','うちわ全力製作係','ビジター遠征の鬼','ユニ全種類制覇','7回に必ず泣く人','ジェット風船担当','神宮の常連','売り子に詳しすぎる','スタメン全部言える','双眼鏡の女王','勝ちメシ研究家','雨でも現地','記録より記憶の人','応援ボードの匠','始球式に憧れる','ビール1杯で酔う人','三塁側の主','ハッピー投げ職人','負けても笑顔の人']
};
function _p(a){return a[Math.floor(Math.random()*a.length)];}
function genNickname(g){
  if(g==='m') return Math.random()<0.55 ? _p(NICK.coolPre)+_p(NICK.coolSuf) : _p(NICK.netaM);
  if(g==='f') return Math.random()<0.55 ? _p(NICK.cutePre)+_p(NICK.cuteSuf) : _p(NICK.netaF);
  const r=Math.random();
  if(r<0.3) return _p(NICK.coolPre)+_p(NICK.coolSuf);
  if(r<0.6) return _p(NICK.cutePre)+_p(NICK.cuteSuf);
  if(r<0.8) return _p(NICK.netaM);
  return _p(NICK.netaF);
}
let _ob={gender:null,age:null,nick:''};
function getProfile(){ try{return JSON.parse(localStorage.getItem('mlb_profile')||'null');}catch(e){return null;} }
function obSel(kind,btn){
  _ob[kind]=btn.getAttribute('data-v');
  btn.parentElement.querySelectorAll('.ob-opt').forEach(b=>b.classList.toggle('sel',b===btn));
  $('ob-next').disabled=!(_ob.gender&&_ob.age);
}
function obToStep2(){
  $('ob-step1').style.display='none';
  $('ob-step2').style.display='';
  obReroll();
}
function obReroll(){
  // 常連・伏兵の名前はユーザーが選べないよう除外
  const roster=rosterPool(); const banned={}; roster.regs.forEach(n=>banned[n]=1); roster.occs.forEach(n=>banned[n]=1);
  let nk=genNickname(_ob.gender);
  for(let k=0;k<40 && banned[nk];k++) nk=genNickname(_ob.gender);
  _ob.nick=nk;
  const el=$('ob-nick'); el.textContent=_ob.nick;
  el.style.animation='none';void el.offsetWidth;el.style.animation='';
}
function obConfirm(){
  const prof={gender:_ob.gender,age:_ob.age,nick:_ob.nick,ts:todayStr()};
  localStorage.setItem('mlb_profile',JSON.stringify(prof));
  track('onboarding_complete',{gender:_ob.gender,age_band:_ob.age}); // 匿名の属性集計
  gaSetUserProps(); // 以降の全イベントに性別・年代を紐づける
  $('onboard').style.display='none';
  goTab('home');
}
function maybeShowOnboard(){
  if(getProfile()) return false;
  _ob={gender:null,age:null,nick:''};
  $('ob-step2').style.display='none';
  $('ob-step1').style.display='';
  $('ob-next').disabled=true;
  document.querySelectorAll('#onboard .ob-opt').forEach(b=>b.classList.remove('sel'));
  $('onboard').style.display='flex';
  return true;
}

// 行代表文字セット（ア行ならア）
const _KANA_ROW_HEADS='アカサタナハマヤラワ';
const _KANA_DAKU={'ガ':'カ','ギ':'キ','グ':'ク','ゲ':'ケ','ゴ':'コ','ザ':'サ','ジ':'シ','ズ':'ス','ゼ':'セ','ゾ':'ソ','ダ':'タ','ヂ':'チ','ヅ':'ツ','デ':'テ','ド':'ト','バ':'ハ','ビ':'ヒ','ブ':'フ','ベ':'ヘ','ボ':'ホ','パ':'ハ','ピ':'ヒ','プ':'フ','ペ':'ヘ','ポ':'ホ','ヴ':'ウ'};
const _KANA_SMALL={'ァ':'ア','ィ':'イ','ゥ':'ウ','ェ':'エ','ォ':'オ','ッ':'ツ','ャ':'ヤ','ュ':'ユ','ョ':'ヨ'};
// 個別かな正規化：ひらがな→カタカナ、濁点→清音、小文字→大文字
function kanaSection(ch){
  let c=ch;
  const code=c.charCodeAt(0);
  if(code>=0x3041&&code<=0x3096) c=String.fromCharCode(code+0x60); // hiragana→katakana
  if(_KANA_DAKU[c]) c=_KANA_DAKU[c];
  if(_KANA_SMALL[c]) c=_KANA_SMALL[c];
  return c;
}
// 後方互換：行代表文字を返す旧kanaRow
function kanaRow(ch){return kanaSection(ch);}
const JA_KANA_INDEX=['ア','イ','ウ','エ','オ','カ','キ','ク','ケ','コ','サ','シ','ス','セ','ソ','タ','チ','ツ','テ','ト','ナ','ニ','ヌ','ネ','ノ','ハ','ヒ','フ','ヘ','ホ','マ','ミ','ム','メ','モ','ヤ','ユ','ヨ','ラ','リ','ル','レ','ロ','ワ'];

// ── Tab ───────────────────────────────────────────────────────────────────────
function goTab(t){
  document.querySelectorAll('.tab-panel').forEach(p=>p.classList.remove('active'));
  document.querySelectorAll('.bnav-btn').forEach(b=>b.classList.remove('active'));
  $('tab-'+t).classList.add('active');
  // ナビに無いタブは親タブをハイライト（略語→辞書、今日/復習→クイズ）
  const _navMap={abbr:'dict',daily:'quiz',review:'quiz',standings:'home'};
  const _nb=$('bnav-'+(_navMap[t]||t));
  if(_nb) _nb.classList.add('active');
  tab=t;
  // カンニング防止: 未回答のまま辞書・略語を開いたら、そのクイズには戻れない
  if((t==='dict'||t==='abbr')&&(lastQuizSection==='play'||lastQuizSection==='daily')){
    const unanswered=lastQuizSection==='play'?!ans:!dailyAns;
    if(unanswered){
      // フリーバッティングは辞書を開いた時点で打ち切り（裏で時間切れの×・不正解音が鳴らないように）
      if(lastQuizSection==='play' && !versusActive){
        stopPitchClock(); cq=null; ans=true;           // onTimeUp のガードで時間切れ処理を無効化
        if(typeof sw==='object'&&sw&&sw.active){ sw.active=false; if(sw.raf) cancelAnimationFrame(sw.raf); }
      }
      lastQuizSection='menu';
    }
  }
  // 辞書ボタン: お気に入り(♥)表示中なら通常の辞書に戻す（位置を復元）
  if(t==='dict' && dictFav){ exitDictFav(); }
  else if(t==='dict' && sq){
    // 辞書タブに戻るとき: 検索をクリアして最後に見ていた位置へ
    sq=''; $('s-input').value=''; $('s-clear').classList.remove('show');
    renderDictList();
    const _pos=localStorage.getItem('mlb_pos');
    const _parts=_pos?_pos.split('|'):null;
    if(_parts && _parts[0]===dictMode && _parts[1])
      requestAnimationFrame(()=>setTimeout(()=>scrollToLetter(_parts[1]),80));
  }
  else if(t==='dict'){
    // 通常表示で辞書を開いたとき: 最後に見ていた位置へ復元（先頭に戻さない）
    const _pos=localStorage.getItem('mlb_pos');
    const _parts=_pos?_pos.split('|'):null;
    if(_parts && _parts[0]===dictMode && _parts[1])
      requestAnimationFrame(()=>setTimeout(()=>scrollToLetter(_parts[1]),80));
  }
  if(t==='home') renderHome();
  if(t==='standings') renderStandings();
  if(t==='abbr') renderAbbrList();
  if(t==='quiz'){ if(lastQuizSection==='play') quizShowPlay(); else quizToMenu(); }
  if(t==='daily'){ lastQuizSection='daily'; initDaily(); }
  if(t==='review'){ lastQuizSection='review'; renderReviewList(); }
  if(t==='settings') updateScores();
  updateBanner();
}
// バナー広告の表示制御（辞書・略語・お気に入り・今日の10問・復習のみ／プレミアムは非表示）
function updateBanner(){
  const b=$('ad-banner'); if(!b) return;
  const show=(typeof isPremium==='function'&&!isPremium())&&['dict','abbr','daily','review'].indexOf(tab)>=0;
  b.classList.toggle('show',show);
}
// ── ホーム ──────────────────────────────────────────────────────────────────
const RANKS=[
  {min:0,name:'ルーキー'},{min:20,name:'ベンチ入り'},{min:50,name:'レギュラー'},
  {min:100,name:'クリーンナップ'},{min:175,name:'オールスター'},{min:275,name:'ゴールドグラブ'},
  {min:400,name:'シルバースラッガー'},{min:600,name:'MVP'},{min:850,name:'リーグの顔'},
  {min:1200,name:'レジェンド'},{min:1700,name:'殿堂入り'},{min:2100,name:'球界の重鎮'},
  {min:2500,name:'生ける伝説'},{min:3000,name:'世界の主砲'},{min:3500,name:'ミスター・ベースボール'},
  {min:4000,name:'野球の神様'}
];
function currentRank(runs){
  let r=RANKS[0],next=null;
  for(let i=0;i<RANKS.length;i++){ if(runs>=RANKS[i].min){ r=RANKS[i]; next=RANKS[i+1]||null; } }
  return {r,next};
}
// 連続ログイン更新（1日1回）
function updateStreak(){
  const today=todayStr();
  let s={last:'',count:0};
  try{ s=JSON.parse(localStorage.getItem('mlb_streak'))||s; }catch(e){}
  if(s.last===today) return s.count;
  // 前日かどうか
  const d=new Date(); d.setDate(d.getDate()-1);
  const y=d.getFullYear()+'-'+String(d.getMonth()+1).padStart(2,'0')+'-'+String(d.getDate()).padStart(2,'0');
  s.count=(s.last===y)?s.count+1:1;
  s.last=today;
  localStorage.setItem('mlb_streak',JSON.stringify(s));
  return s.count;
}
// 日付から決定的にインデックスを作る（毎日変わる）
function daySeed(){
  const d=new Date();
  return d.getFullYear()*10000+(d.getMonth()+1)*100+d.getDate();
}
let _wordOfDay=null;
function pickWordOfDay(){
  // 人名・略語・日本語なしを除いたEJから、日付で決定的に選ぶ
  const pool=EJ.filter(e=>e.cat!=='abbr' && e.ja && /[ぁ-んァ-ヶー一-龯]/.test(e.ja) && !isPersonName2(e,'ej'));
  if(!pool.length) return null;
  return pool[daySeed()%pool.length];
}
function openWordOfDay(){ if(_wordOfDay) openModal(_wordOfDay,'ej'); }
function renderHome(){
  const d=new Date();
  try{ todayField(); }catch(e){} // その日のフィールドを確定させ、過去ランキングに残るようにする
  if($('home-date')) $('home-date').textContent=(d.getMonth()+1)+'月'+d.getDate()+'日（'+'日月火水木金土'[d.getDay()]+'）';
  const _pf=getProfile();
  if($('home-player')) $('home-player').textContent=_pf&&_pf.nick?('⚾ '+_pf.nick):'';
  // 連続ログイン
  const streak=updateStreak();
  if($('home-streak-n')) $('home-streak-n').textContent=streak;
  // 称号
  const runs=(typeof runsTotal==='number')?runsTotal:0;
  const {r,next}=currentRank(runs);
  if($('home-rank')) $('home-rank').textContent=r.name;
  if($('home-runs')) $('home-runs').textContent=runs;
  if($('home-rank-fill')){
    const pct=next?Math.min(100,Math.round((runs-r.min)/(next.min-r.min)*100)):100;
    $('home-rank-fill').style.width=pct+'%';
  }
  if($('home-rank-next')) $('home-rank-next').textContent=next?('次の称号「'+next.name+'」まであと'+Math.max(0,next.min-runs)+'点'):'最高称号に到達！';
  // 今日の一語
  _wordOfDay=pickWordOfDay();
  if(_wordOfDay){
    if($('home-word-term')) $('home-word-term').textContent=_wordOfDay.en;
    if($('home-word-sub')){
      // 拡充済みは short（簡潔な訳）を表示。無ければ生データを整形
      let sub=(_wordOfDay.short||'').trim();
      if(!sub) sub=(_wordOfDay.ja||'').replace(/\s*cf\..*$/i,'').replace(/\s*e\.g\..*$/i,'').replace(/\n/g,' ');
      sub=sub.replace(/[。、\s]+$/,'').trim();
      if(sub.length>60) sub=sub.substring(0,60)+'…';
      $('home-word-sub').textContent=sub;
    }
  }
  // 今日の豆知識
  if($('home-trivia')&&TRIVIA.length) $('home-trivia').textContent=TRIVIA[daySeed()%TRIVIA.length];
  // この日のMLB史
  renderOnThisDay();
  // 今日の日本人選手（ネタバレ防止でタップ表示。一度開けばセッション中は継続）
  if(jpRevealed){ revealJpPlayers(); } else { hideJpPlayers(); }
  // 今日の試合（ネット必要）
  loadGames();
}
// ── 成績・本日のランキング（ローカル。大母数の中で本日得点を競う） ────────────────
// 上位は僅差、下は長い裾野（10点前後で約100位）。母数は日替わりで約140〜180人。
// ── 生活パターン（1日の点数の伸び方）─────────────────────────────
// w = 各時間帯(0〜23時)に稼ぐ相対ポイント。cumFrac で「今の時刻までの達成率(0〜1)」に変換。
const PERSONAS={
  steady:  {w:[13,11,6,1,1,1,1,3,3,2,2,2,6,4,3,6,4,3,7,5,4,7,5,6]}, // 深夜まで起きてて一気に、以降まんべんなく
  nightowl:{w:[16,14,9,3,1,1,1,1,1,1,1,1,1,1,1,2,2,3,4,5,6,9,13,16]},// 深夜〜未明と夜更けにドカッと
  latenite:{w:[9,20,10,2,1,1,1,1,1,1,1,1,1,1,1,1,1,2,2,3,3,5,7,9]},  // 1〜2時に爆発するタイプ
  commuter:{w:[2,2,1,1,1,1,2,8,9,5,2,2,3,2,2,2,3,6,9,8,5,3,2,1]}, // 通勤時間帯(朝夕)にドカッと
  meal:    {w:[2,2,1,1,1,1,2,6,5,2,2,3,8,7,3,2,2,3,5,9,8,4,2,1]}, // ご飯の時間(朝昼晩)だけ
  evening: {w:[3,2,1,1,1,1,1,2,2,2,2,2,3,3,3,4,5,6,8,9,8,7,5,3]}, // 夜型・帰宅後に伸びる
  morning: {w:[3,2,1,1,1,2,4,8,7,6,5,4,4,3,3,3,3,3,3,3,2,2,2,2]}, // 朝活タイプ
  weekend: {w:[7,6,4,1,1,1,2,4,5,5,6,6,7,6,5,6,6,6,7,7,6,6,5,5], wkScale:0.28}, // 週末だけすさまじい(平日は控えめ)
  closer:  {w:[1,1,0,0,0,0,0,0,0,0,0,0,0,0,0,1,1,2,3,5,9,15,22,26]}, // 日中は沈み、夜12時間際に一気に追い込む
  dawn:    {w:[26,20,11,3,1,1,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,0,1]}, // 日付が変わった瞬間から一気に稼ぐ早朝リーダー
};
const PERSONA_ORDER=['steady','nightowl','latenite','commuter','weekend','meal','evening','nightowl','morning','latenite','commuter','meal'];
function cumFrac(w,h){ let tot=0; for(const x of w) tot+=x; if(tot<=0) return 0; let acc=0; const full=Math.floor(h); for(let i=0;i<full&&i<24;i++) acc+=w[i]; if(full<24) acc+=(w[full]||0)*(h-full); return Math.min(1,acc/tot); }
function nowHour(){ const d=new Date(); return d.getHours()+d.getMinutes()/60; }
// 日をまたいで固定の登場人物：常連(いつも上位)＋伏兵(たまに急上昇)。名前は永続。
function rosterPool(){
  let r; try{ r=JSON.parse(localStorage.getItem('mlb_roster')); }catch(e){}
  if(r&&r.regs&&r.regs.length>=4&&r.occs&&r.occs.length>=5) return r;
  const my=(getProfile()&&getProfile().nick)||''; const used={}; used[my]=1;
  const pick=()=>{ for(let k=0;k<40;k++){const c=genNickname(''); if(c&&!used[c]){used[c]=1;return c;}} return 'ライバル'; };
  const regs=[]; for(let i=0;i<4;i++) regs.push(pick());  // 常連4名（毎日2〜4名が上位に）
  const occs=[]; for(let i=0;i<5;i++) occs.push(pick());  // 伏兵5名（たまに一気に上位へ）
  const out={regs,occs};
  try{ localStorage.setItem('mlb_roster',JSON.stringify(out)); }catch(e){}
  return out;
}
function todayField(){
  const key='mlb_field_'+todayStr();
  try{ const c=JSON.parse(localStorage.getItem(key)); if(c&&c.names&&c.names.length>=40&&c.finals&&c.finals.length>=250&&c.personas&&c.personas.length>=250&&c.N>=250) return c; }catch(e){}
  const wknd=(function(){const g=new Date().getDay(); return g===0||g===6;})();
  const N=290+Math.floor(Math.random()*111);     // 全体 290〜400人（この時間帯でも大勢が参加している設定）
  const p=1.7+Math.random()*0.8;                 // 落ち方の急さ（日替わり 1.7〜2.5）
  // 1位の「1日の最終」を日替わりで大きく変動。週末は高め
  let T=90+Math.floor(Math.random()*71);         // 平日 90〜160点
  if(wknd) T+=25+Math.floor(Math.random()*61);   // 週末は +25〜85点（盛り上がる日）
  // その日の型
  const roll=Math.random();
  let dtype='flat';
  if(roll<0.16) dtype='runaway';                 // 独走の1位
  else if(roll<0.38) dtype='duel';               // 1位と2位だけが高得点で競る
  else if(roll<0.56) dtype='tight';              // 全体が接戦
  // 全N人分の最終スコア（べき乗カーブ：上位はなだらか、下位まで自然に分布）。tightは平坦め
  const shape=(dtype==='tight')?(1.0+Math.random()*0.35):(1.4+Math.random()*0.5);
  const finals=[];
  for(let r=0;r<N;r++) finals.push(Math.max(1,Math.round(T*Math.pow(Math.max(0,1-r/N), shape))));
  if(dtype==='runaway'){ T=Math.round(T*(1.28+Math.random()*0.3)); finals[0]=T; }        // 独走
  else if(dtype==='duel'){ finals[0]=T; finals[1]=Math.round(T*(0.93+Math.random()*0.05)); for(let r=2;r<N;r++) finals[r]=Math.round(finals[r]*0.7); } // 1・2位だけ突出
  const top=finals[0];
  const roster=rosterPool();
  const my=(getProfile()&&getProfile().nick)||'';
  const used={}; used[my]=1;
  const shuf=a=>{a=a.slice();for(let i=a.length-1;i>0;i--){const j=Math.floor(Math.random()*(i+1));[a[i],a[j]]=[a[j],a[i]];}return a;};
  // 今日出勤する常連（2〜4名。各72%で出勤＝たまにお休みの日も）
  let regs=shuf(roster.regs).filter(()=>Math.random()<0.72);
  if(regs.length<2) regs=shuf(roster.regs).slice(0,2+Math.floor(Math.random()*2));
  if(regs.length>4) regs=regs.slice(0,4);
  // たまに一気に上位へ来る伏兵（各18%で今日出現、最大2名）
  let surg=shuf(roster.occs).filter(()=>Math.random()<0.18);
  if(surg.length>2) surg=surg.slice(0,2);
  // 上位ランクは常連＋伏兵で埋め、残りは日替わりのランダム名。伏兵は出た日はトップ寄りに
  const topNames=shuf(surg).concat(shuf(regs));
  const names=[];  // 表示は上位20人なので、名前は上位40人分だけ用意（ドラマで沈む人がいても表示は名前付きに収まる）
  topNames.forEach(n=>{ if(names.length<40 && !used[n]){ used[n]=1; names.push(n); } });
  while(names.length<40){ let nm='ライバル'; for(let k=0;k<30;k++){const c=genNickname(''); if(c&&!used[c]){nm=c;used[c]=1;break;}} names.push(nm); }
  // 生活パターンを全N人に割り当て（各人が時刻ごとに違う伸び方＝毎分の順位が自然に動く）
  const personas=[]; while(personas.length<N){ const b=PERSONA_ORDER.slice(); for(let i=b.length-1;i>0;i--){const j=Math.floor(Math.random()*(i+1));[b[i],b[j]]=[b[j],b[i]];} personas.push(...b); } personas.length=N;
  // 日替わりの順位ドラマ（closer=夜の追い込み / dawn=0時からの早朝リーダー / late系=競り合い）
  const froll=Math.random(); let finish='normal';
  if(froll<0.18)      finish='comeback';   // 日中は下位、夜12時間際に一気に上位へ逆転
  else if(froll<0.34) finish='deadheat';   // 1・2位が最後まで競り合い、終盤に得点を伸ばす
  else if(froll<0.48) finish='midnight';   // 0時回った瞬間から一気に稼ぐ早朝リーダーが出る
  else if(froll<0.58) finish='triple';     // 上位3人が夜に三つ巴の追い込み
  if(finish==='deadheat'){ personas[0]='closer'; personas[1]='nightowl'; }
  else if(finish==='comeback'){ const slot=2+Math.floor(Math.random()*5); personas[slot]='closer'; } // 3〜7位枠→日中沈み夜に急浮上
  else if(finish==='midnight'){ const slot=Math.floor(Math.random()*3); personas[slot]='dawn'; }      // 上位1〜3枠→0時から爆走
  else if(finish==='triple'){ personas[0]='closer'; personas[1]='latenite'; personas[2]='nightowl'; }
  const f={top,N,p:shape,names,personas,wknd,finals,dtype,finish};
  try{ localStorage.setItem(key,JSON.stringify(f)); }catch(e){}
  pruneOldFields();
  return f;
}
// 古い日のフィールドを整理（過去ランキング用に直近8日分だけ保持）
function pruneOldFields(){
  try{
    const keep=8; const cut=new Date(); cut.setDate(cut.getDate()-keep);
    const cutS=cut.getFullYear()+'-'+String(cut.getMonth()+1).padStart(2,'0')+'-'+String(cut.getDate()).padStart(2,'0');
    for(let i=localStorage.length-1;i>=0;i--){
      const k=localStorage.key(i);
      if(k&&k.indexOf('mlb_field_')===0){ const ds=k.slice('mlb_field_'.length); if(ds<cutS) localStorage.removeItem(k); }
    }
  }catch(e){}
}
// ランクr(1位=最強)の「1日の最終スコア」
function finalForRank(rank,f){ if(f.finals&&f.finals[rank-1]!=null) return f.finals[rank-1]; return Math.max(1, Math.round(f.top*Math.pow(Math.max(0,1-(rank-1)/f.N), f.p))); }
// 今この瞬間のスコア = 最終 × その人の生活パターンの達成率
const DAY_END=24-1/3600; // 23:59:59 で「最終値」に到達させる
function liveScoreForRank(rank,f,h){
  const base=finalForRank(rank,f);
  const P=PERSONAS[f.personas[rank-1]]||PERSONAS.steady;
  const hh=Math.min(h,DAY_END)*(24/DAY_END); // 23:59:59 で達成率100%になるよう調整
  let sc=base*cumFrac(P.w,hh);
  if(P.wkScale && !f.wknd) sc*=P.wkScale; // 週末型は平日は伸びない
  return Math.max(0,Math.round(sc));
}
function scoreForRank(rank,f){ return finalForRank(rank,f); }
function rankForScore(s,f){
  if(s>=f.top) return 1;
  if(s<=0) return f.N;
  const r=1 + f.N*(1 - Math.pow(s/f.top, 1/f.p));
  return Math.min(f.N, Math.max(1, Math.round(r)));
}
function renderStats(){
  const f=todayField();
  const myNick=(getProfile()&&getProfile().nick)||'あなた';
  const us=runsToday;
  const h=nowHour();
  // 今の時刻の全体の進み具合（早朝は低い→少点でも上位に入れる）
  const avgP=(cumFrac(PERSONAS.steady.w,h)+cumFrac(PERSONAS.commuter.w,h)+cumFrac(PERSONAS.evening.w,h)+cumFrac(PERSONAS.meal.w,h))/4;
  // 全ライバル（N人）の現時点スコアを算出。上位20人を表示し、あなたの順位も同じ分布から出す（矛盾なし）
  const NF=f.finals.length;
  const rivals=[]; for(let r=1;r<=NF;r++) rivals.push({name:(f.names[r-1]||('選手'+r)),score:liveScoreForRank(r,f,h),me:false});
  // 終盤（夜〜終了間際）はユーザーを最高2位までに。duelは上位2人をキープ（最高3位）
  if(avgP>=0.8 && us>0){
    const keep=(f.dtype==='duel')?2:1;
    const bylive=rivals.slice().sort((a,b)=>b.score-a.score);
    for(let n=0;n<keep && n<bylive.length;n++){ if(bylive[n].score<=us) bylive[n].score=us+(keep-n)+Math.floor(Math.random()*3); }
  }
  const merged=rivals.concat([{name:myNick,score:us,me:true}]).sort((a,b)=> b.score-a.score || (a.me?1:-1));
  const idx=merged.findIndex(x=>x.me);
  const inTop20=idx<20;
  const total=f.N;
  const myRank=Math.min(idx+1, total);
  const you=$('rank-you'); if(you) you.textContent='あなたは '+myRank+' 位 / 全'+total+'人中';
  const bd=$('rank-board');
  if(bd){
    let rows=merged.slice(0,20).map((x,i)=>
      '<div class="rank-row'+(x.me?' me':'')+'"><span class="rank-pos">'+(i+1)+'</span>'+
      '<span class="rank-nm">'+esc(x.name)+(x.me?'（あなた）':'')+'</span>'+
      '<span class="rank-sc">'+x.score+' 点</span></div>').join('');
    if(!inTop20){
      rows+='<div class="rank-sep">⋮</div>'+
        '<div class="rank-row me"><span class="rank-pos">'+myRank+'</span>'+
        '<span class="rank-nm">'+esc(myNick)+'（あなた）</span>'+
        '<span class="rank-sc">'+us+' 点</span></div>';
    }
    bd.innerHTML=rows;
  }
  // マイレコード
  if($('rec-today')) $('rec-today').textContent=runsToday;
  if($('rec-total')) $('rec-total').textContent=runsTotal;
  let rec={w:0,l:0}; try{ rec=JSON.parse(localStorage.getItem('mlb_vs_rec'))||rec; }catch(e){}
  if($('rec-vs')) $('rec-vs').textContent=rec.w+'-'+rec.l;
  let streak=0; try{ streak=(JSON.parse(localStorage.getItem('mlb_streak'))||{}).count||0; }catch(e){}
  if($('rec-streak')) $('rec-streak').textContent=streak;
  const {r,next}=currentRank(runsTotal);
  if($('rec-rank-name')) $('rec-rank-name').textContent=r.name;
  if($('rec-rank-next')) $('rec-rank-next').textContent=next?('（次の称号「'+next.name+'」まであと'+Math.max(0,next.min-runsTotal)+'点）'):'（最高称号！）';
  if($('rec-rank-fill')){ const pct=next?Math.min(100,Math.round((runsTotal-r.min)/(next.min-r.min)*100)):100; $('rec-rank-fill').style.width=pct+'%'; }
}
// ── MLBライブ（statsapi.mlb.com 公開API） ─────────────────────────────────────
const TEAM_JA={108:'エンゼルス',109:'Dバックス',110:'オリオールズ',111:'レッドソックス',112:'カブス',113:'レッズ',114:'ガーディアンズ',115:'ロッキーズ',116:'タイガース',117:'アストロズ',118:'ロイヤルズ',119:'ドジャース',120:'ナショナルズ',121:'メッツ',133:'アスレチックス',134:'パイレーツ',135:'パドレス',136:'マリナーズ',137:'ジャイアンツ',138:'カージナルス',139:'レイズ',140:'レンジャーズ',141:'ブルージェイズ',142:'ツインズ',143:'フィリーズ',144:'ブレーブス',145:'ホワイトソックス',146:'マーリンズ',147:'ヤンキース',158:'ブルワーズ'};
const DIV_JA={200:'ア・リーグ西地区',201:'ア・リーグ東地区',202:'ア・リーグ中地区',203:'ナ・リーグ西地区',204:'ナ・リーグ東地区',205:'ナ・リーグ中地区'};
function teamJa(t){return TEAM_JA[t.id]||t.teamName||t.name||'';}
// 米国東部時間での「今日」の日付（MLBの試合日はET基準）
function etDateStr(){
  const p=new Intl.DateTimeFormat('en-CA',{timeZone:'America/New_York',year:'numeric',month:'2-digit',day:'2-digit'}).format(new Date());
  return p; // YYYY-MM-DD
}
// 米国東部時間で days 日ずらした日付（-1=昨日）
function etDateStrOffset(days){
  const d=new Date(Date.now()+days*86400000);
  return new Intl.DateTimeFormat('en-CA',{timeZone:'America/New_York',year:'numeric',month:'2-digit',day:'2-digit'}).format(d);
}
// キャッシュつきfetch（分単位TTL）
async function cachedJson(url,key,ttlMin){
  try{
    const c=JSON.parse(localStorage.getItem(key)||'null');
    if(c&&Date.now()-c.t<ttlMin*60000) return c.d;
  }catch(e){}
  const r=await fetch(url);
  if(!r.ok) throw new Error('http '+r.status);
  const d=await r.json();
  try{localStorage.setItem(key,JSON.stringify({t:Date.now(),d}));}catch(e){}
  return d;
}
function gameStateJa(g){
  const s=g.status&&g.status.abstractGameState;
  if(s==='Live'){
    // 試合中は「◯回表/裏」を表示（linescoreがあれば）
    const ls=g.linescore;
    if(ls&&ls.currentInning){
      const half=(ls.isTopInning!==undefined)?(ls.isTopInning?'表':'裏')
                 :((ls.inningState==='Top'||ls.inningState==='Middle')?'表':'裏');
      return {cls:'live',txt:ls.currentInning+'回'+half};
    }
    return {cls:'live',txt:'試合中'};
  }
  if(s==='Final') return {cls:'final',txt:'終了'};
  // 開始時刻（日本時間で表示。米国の夜＝日本の翌朝になるため日付も付ける）
  try{
    const t=new Intl.DateTimeFormat('ja-JP',{timeZone:'Asia/Tokyo',month:'numeric',day:'numeric',hour:'2-digit',minute:'2-digit'}).format(new Date(g.gameDate));
    return {cls:'pre',txt:'日本時間 '+t+' 開始'};
  }catch(e){ return {cls:'pre',txt:'予定'}; }
}
function teamLogo(t){
  return `<img class="team-logo" src="https://www.mlbstatic.com/team-logos/${t.id}.svg" alt="" onerror="this.style.display='none'">`;
}
// ネタバレ防止: 結果を表示した試合（セッション中のみ保持。アプリ再起動で再び非表示）
const gameRevealed=new Set();
let _lastGames=[];
function revealGame(pk){ gameRevealed.add(pk); loadGames(); }
function revealAllGames(){ _lastGames.forEach(g=>gameRevealed.add(g.gamePk)); loadGames(); }
async function loadGames(){
  const el=$('home-games');
  if(!el) return;
  // 見出しに米国日付を表示（時差で日本の日付とズレるため）
  try{ const p=etDateStr().split('-'); const ud=$('games-usdate'); if(ud) ud.textContent='（米国 '+parseInt(p[1],10)+'/'+parseInt(p[2],10)+'）'; }catch(e){}
  try{
    const d=await cachedJson('https://statsapi.mlb.com/api/v1/schedule?sportId=1&hydrate=linescore&date='+etDateStr(),'mlb_games2',5);
    const games=(d.dates&&d.dates[0]&&d.dates[0].games)||[];
    _lastGames=games;
    if(!games.length){ el.innerHTML='<div style="padding:6px 0">今日は試合がありません</div>'; return; }
    const rowsHtml=games.map(g=>{
      const st=gameStateJa(g);
      const a=g.teams.away, h=g.teams.home;
      const score=(st.cls==='pre')?'vs':(a.score??'-')+' - '+(h.score??'-');
      // ダブルヘッダー（1日2試合）は第1/第2試合を表示
      const dh=(g.doubleHeader&&g.doubleHeader!=='N'&&g.gameNumber)?`<span class="game-dh">第${g.gameNumber}試合</span>`:'';
      // 結果（試合中・終了）はネタバレ防止で1試合ずつタップ表示。未開始は結果が無いので通常表示
      const isResult=(st.cls==='live'||st.cls==='final');
      let mid;
      if(isResult && !gameRevealed.has(g.gamePk)){
        mid=`${dh}<button type="button" class="game-reveal" onclick="revealGame(${g.gamePk})">結果</button>`;
      }else{
        mid=`${dh}<span class="game-score">${esc(String(score))}</span><br><span class="game-state ${st.cls}">${esc(st.txt)}</span>`;
      }
      return `<div class="game-row">
        <span class="game-team away">${teamLogo(a.team)}<span class="game-team-name">${esc(teamJa(a.team))}</span></span>
        <span class="game-mid">${mid}</span>
        <span class="game-team"><span class="game-team-name">${esc(teamJa(h.team))}</span>${teamLogo(h.team)}</span>
      </div>`;
    }).join('');
    // 未表示の結果があれば「すべて表示」ボタン（上部＋末尾）
    const anyHidden=games.some(g=>{const s=gameStateJa(g).cls; return (s==='live'||s==='final')&&!gameRevealed.has(g.gamePk);});
    const revealBtn=anyHidden?`<button type="button" class="games-reveal-all" onclick="revealAllGames()">👀 すべての結果を表示</button>`:'';
    el.innerHTML=rowsHtml+revealBtn;
    const topWrap=$('games-revealall-top'); if(topWrap) topWrap.innerHTML=anyHidden?`<button type="button" class="games-reveal-all games-reveal-all-top" onclick="revealAllGames()">👀 すべての結果を表示</button>`:'';
  }catch(e){
    el.innerHTML='<div style="padding:6px 0">オフラインのため表示できません</div>';
  }
}
// 手動更新：キャッシュを無視して今日の試合を再取得（開いていれば昨日の結果も）
async function refreshGames(){
  const btn=$('games-refresh-btn');
  if(btn){ btn.disabled=true; btn.textContent='更新中…'; }
  try{ localStorage.removeItem('mlb_games2'); }catch(e){}
  try{ await loadGames(); }catch(e){}
  const yb=$('yday-games');
  if(yb && yb.style.display!=='none'){
    try{ localStorage.removeItem('mlb_games_recent'); }catch(e){}
    _ydayLoaded=false;
    try{ await loadYesterdayGames(); }catch(e){}
  }
  if(btn){ btn.disabled=false; btn.textContent='🔄 最新に更新'; }
}
// 昨日（米国日付）の結果を開閉。日本では米国の昨日＝直近の完了試合なので結果を直接表示
let _ydayLoaded=false;
function toggleYdayGames(){
  const box=$('yday-games'), btn=$('yday-btn'); if(!box) return;
  const open=box.style.display==='none';
  box.style.display=open?'':'none';
  if(btn) btn.textContent=open?'📅 昨日の結果を隠す ▲':'📅 昨日の結果を見る ▼';
  if(open && !_ydayLoaded) loadYesterdayGames();
}
async function loadYesterdayGames(){
  const el=$('yday-games'); if(!el) return;
  el.innerHTML='<div style="padding:6px 0">読み込み中…</div>';
  try{
    // 直近7日〜本日(米国日付)を取得し、「終了した試合がある直近の日」を採用（予定・進行中は除外）
    const start=etDateStrOffset(-7), end=etDateStr();
    const d=await cachedJson('https://statsapi.mlb.com/api/v1/schedule?sportId=1&hydrate=linescore&startDate='+start+'&endDate='+end,'mlb_games_recent',30);
    // 終了（スコアあり）＝完了試合。中止/延期/サスペンドは「中止」として表示（予定・進行中は除外）
    const isFinal=g=>g.status&&g.status.abstractGameState==='Final'&&g.teams&&g.teams.away&&g.teams.home&&g.teams.away.score!=null&&g.teams.home.score!=null;
    const isCancelled=g=>{ const ds=(g.status&&g.status.detailedState)||''; if(/Postponed|Cancel|Suspended/i.test(ds)) return true; return g.status&&g.status.abstractGameState==='Final'&&!isFinal(g); };
    const isShown=g=>isFinal(g)||isCancelled(g);
    // 終了(スコアあり)試合がある直近の日を採用。その日の中止試合も併せて表示
    let latest=null;
    (d.dates||[]).forEach(x=>{ const fin=(x.games||[]).filter(isFinal); if(fin.length) latest={date:x.date,games:(x.games||[]).filter(isShown)}; });
    if(!latest){ el.innerHTML='<div style="padding:6px 0">直近に終了した試合がありませんでした</div>'; _ydayLoaded=true; return; }
    const games=latest.games;
    // 日付見出し（例：7/20（日）の結果）
    let head='';
    try{ const dt=new Date(latest.date+'T12:00:00'); head='<div class="yday-sep">📅 '+(dt.getMonth()+1)+'/'+dt.getDate()+'（'+'日月火水木金土'[dt.getDay()]+'）の結果（米国日付）</div>'; }catch(e){}
    el.innerHTML=head+games.map(g=>{
      const a=g.teams.away, h=g.teams.home;
      const cancelled=!isFinal(g);
      const score=cancelled?'ー':((a.score??'-')+' - '+(h.score??'-'));
      const dh=(g.doubleHeader&&g.doubleHeader!=='N'&&g.gameNumber)?`<span class="game-dh">第${g.gameNumber}試合</span>`:'';
      const stTxt=cancelled?'中止':'終了';
      const cls=cancelled?'pre':'final';
      const mid=`${dh}<span class="game-score">${esc(String(score))}</span><br><span class="game-state ${cls}">${esc(stTxt)}</span>`;
      return `<div class="game-row">
        <span class="game-team away">${teamLogo(a.team)}<span class="game-team-name">${esc(teamJa(a.team))}</span></span>
        <span class="game-mid">${mid}</span>
        <span class="game-team"><span class="game-team-name">${esc(teamJa(h.team))}</span>${teamLogo(h.team)}</span>
      </div>`;
    }).join('');
    _ydayLoaded=true;
  }catch(e){ el.innerHTML='<div style="padding:6px 0">オフラインのため表示できません</div>'; }
}
// ── 今日の日本人選手 ─────────────────────────────────────────────────────────
const JP_PLAYERS=[
  {id:660271,ja:'大谷翔平'},{id:808967,ja:'山本由伸'},{id:808963,ja:'佐々木朗希'},
  {id:673548,ja:'鈴木誠也'},{id:684007,ja:'今永昇太'},{id:807799,ja:'吉田正尚'},
  {id:673540,ja:'千賀滉大'},{id:579328,ja:'菊池雄星'},{id:673513,ja:'松井裕樹'},
];
function fmtHitLine(st){
  const ab=st.atBats||0,h=st.hits||0,hr=st.homeRuns||0,rbi=st.rbi||0,sb=st.stolenBases||0,bb=st.baseOnBalls||0;
  let s=ab+'打数'+h+'安打';
  const ex=[];
  if(hr) ex.push('<span class="jp-good">'+hr+'本塁打</span>');
  if(rbi) ex.push(rbi+'打点');
  if(sb) ex.push('<span class="jp-good">'+sb+'盗塁</span>');
  if(!hr&&!rbi&&bb) ex.push(bb+'四球');
  return s+(ex.length?' '+ex.join(' '):'');
}
function fmtPitchLine(st){
  const ip=st.inningsPitched||'0',er=(st.earnedRuns!=null?st.earnedRuns:'-'),so=st.strikeOuts||0;
  let s=ip+'回 '+er+'自責 '+so+'奪三振';
  if(st.wins) s+=' <span class="jp-good">勝利</span>';
  else if(st.saves) s+=' <span class="jp-good">セーブ</span>';
  else if(st.losses) s+=' 敗戦';
  return s;
}
function mdOf(dstr){ const m=/^\d+-(\d+)-(\d+)$/.exec(dstr||''); return m?(parseInt(m[1],10)+'/'+parseInt(m[2],10)):''; }
// ネタバレ防止: タップするまで結果を隠す（セッション中は一度開けば継続）
let jpRevealed=false;
function hideJpPlayers(){
  const cov=$('home-jp-cover'), body=$('home-jp');
  if(cov) cov.style.display='';
  if(body) body.style.display='none';
}
function revealJpPlayers(){
  jpRevealed=true;
  const cov=$('home-jp-cover'), body=$('home-jp');
  if(cov) cov.style.display='none';
  if(body){ body.style.display=''; }
  loadJpPlayers();
}
async function loadJpPlayers(){
  const el=$('home-jp'); if(!el) return;
  try{
    const ids=JP_PLAYERS.map(p=>p.id).join(',');
    const yr=etDateStr().slice(0,4);
    const url='https://statsapi.mlb.com/api/v1/people?personIds='+ids+
      '&hydrate=stats(group=%5Bhitting,pitching%5D,type=%5BgameLog%5D,season='+yr+')';
    const d=await cachedJson(url,'mlb_jp_'+yr,20);
    const today=etDateStr();
    const byId={}; (d.people||[]).forEach(p=>byId[p.id]=p);
    const rows=[];
    JP_PLAYERS.forEach(pl=>{
      const p=byId[pl.id]; if(!p) return;
      // 打撃・投球それぞれの最新スプリットを集め、最も新しい試合を選ぶ
      let best=null;
      (p.stats||[]).forEach(s=>{
        const grp=(s.group||{}).displayName;
        const sp=(s.splits||[]);
        if(!sp.length) return;
        const last=sp[sp.length-1];
        if(!best || (last.date||'')>(best.date||'')) best={grp,date:last.date,stat:last.stat||{}};
      });
      if(!best) return;
      const line=(best.grp==='pitching')?fmtPitchLine(best.stat):fmtHitLine(best.stat);
      rows.push({ja:pl.ja,date:best.date,line,today:best.date===today});
    });
    if(!rows.length){ el.innerHTML='<div class="jp-note">データを取得できませんでした</div>'; return; }
    const played=rows.filter(r=>r.today);
    let show, note='';
    if(played.length){ show=played; }
    else{
      // 本日出場なし → 直近の試合を新しい順に最大3人
      show=rows.slice().sort((a,b)=>(b.date||'').localeCompare(a.date||'')).slice(0,3);
      note='<div class="jp-note">本日は試合がありません（直近の成績）</div>';
    }
    el.innerHTML=note+show.map(r=>
      '<div class="jp-row"><span class="jp-name">'+esc(r.ja)+'</span>'+
      '<span class="jp-line"><span class="jp-date">'+esc(mdOf(r.date))+'</span>'+r.line+'</span></div>'
    ).join('');
  }catch(e){
    el.innerHTML='<div class="jp-note">オフラインのため表示できません</div>';
  }
}
// ── この日のMLB史（On This Day） ─────────────────────────────────────────────
const ON_THIS_DAY={
  '04-08':[[1974,'ハンク・アーロンが通算715号本塁打を放ち、ベーブ・ルースの記録を抜いた。']],
  '04-15':[[1947,'ジャッキー・ロビンソンがMLBデビュー。人種の壁を破り、背番号42は全球団の永久欠番となっている。']],
  '05-01':[[1991,'ノーラン・ライアンが史上初の7度目のノーヒットノーランを達成。同じ日にリッキー・ヘンダーソンが通算盗塁の歴代最多記録を更新した。']],
  '05-02':[[1995,'野茂英雄がドジャースでMLBデビュー。「トルネード投法」で大旋風を巻き起こした。']],
  '05-06':[[1998,'カブスの新人カーリー・ウッドが1試合20奪三振の大記録を達成した。']],
  '08-07':[[2007,'バリー・ボンズが通算756号本塁打を放ち、ハンク・アーロンを抜いて歴代最多となった。']],
  '09-06':[[1995,'カル・リプケンJr.が2131試合連続出場を達成し、ルー・ゲーリッグの記録を更新した。']],
  '09-08':[[1998,'マーク・マグワイアがシーズン62号本塁打を放ち、ロジャー・マリスの61本を抜いた。']],
  '09-19':[[2024,'大谷翔平が史上初の「1シーズン50本塁打・50盗塁」を達成した。']],
  '09-28':[[1941,'テッド・ウィリアムズがシーズン打率.406を記録。以来80年以上、4割打者は現れていない。']],
  '10-01':[[1961,'ロジャー・マリスがシーズン61号本塁打でベーブ・ルースの60本を抜いた。'],[2004,'イチローがシーズン262安打の新記録を樹立し、ジョージ・シスラーの257本を84年ぶりに更新した。']],
};
const OTD_FALLBACK=[
  'サイ・ヤングの通算511勝は、今後破られないとされる不滅の記録。',
  'ベーブ・ルースは投手として通算94勝を挙げた、元祖二刀流だった。',
  'カル・リプケンJr.の連続試合出場2632は世界記録。',
  'ジャッキー・ロビンソンの背番号42は全30球団で永久欠番。',
  'ボストンのフェンウェイ・パークは1912年開場、現存最古のMLB球場。',
  'シカゴのリグレー・フィールドは1914年開場、ツタのフェンスで有名。',
  'リッキー・ヘンダーソンの通算1406盗塁は歴代最多の大記録。',
  'ピート・ローズの通算4256安打はMLB歴代最多。',
  'ノーラン・ライアンは通算7回のノーヒットノーランと5714奪三振の記録を持つ。',
  'サッチェル・ペイジはニグロリーグの伝説で、42歳でMLBデビューした。',
];
function renderOnThisDay(){
  const el=$('home-otd'); if(!el) return;
  const d=new Date();
  const key=String(d.getMonth()+1).padStart(2,'0')+'-'+String(d.getDate()).padStart(2,'0');
  const ev=ON_THIS_DAY[key];
  if(ev&&ev.length){
    el.innerHTML=ev.map(e=>'<div><span class="otd-year">'+e[0]+'年</span>'+esc(e[1])+'</div>').join('');
  }else{
    el.innerHTML=esc(OTD_FALLBACK[daySeed()%OTD_FALLBACK.length]);
  }
}
// ── デイリー通知（Capacitor Local Notifications） ────────────────────────────
let dailyNotif=localStorage.getItem('mlb_daily_notif')==='1';
function capLN(){
  try{ return window.Capacitor && window.Capacitor.Plugins && window.Capacitor.Plugins.LocalNotifications; }
  catch(e){ return null; }
}
const NOTIF_ID=4210;
async function scheduleDailyNotif(){
  const ln=capLN(); if(!ln) return false;
  try{
    const perm=await ln.requestPermissions();
    if(perm && perm.display && perm.display!=='granted') return false;
    await ln.cancel({notifications:[{id:NOTIF_ID}]});
    await ln.schedule({notifications:[{
      id:NOTIF_ID,
      title:'⚾ 今日の10問',
      body:'今日のMLB用語の問題にチャレンジしよう！',
      schedule:{ on:{ hour:8, minute:0 }, repeats:true, allowWhileIdle:true }
    }]});
    return true;
  }catch(e){ return false; }
}
async function cancelDailyNotif(){
  const ln=capLN(); if(!ln) return;
  try{ await ln.cancel({notifications:[{id:NOTIF_ID}]}); }catch(e){}
}
function syncNotifToggle(){
  const b=$('notif-toggle'); if(!b) return;
  b.textContent=dailyNotif?'オン':'オフ';
  b.classList.toggle('auto',!dailyNotif); // オフ時にグレー表示
}
async function toggleDailyNotif(){
  dailyNotif=!dailyNotif;
  localStorage.setItem('mlb_daily_notif',dailyNotif?'1':'0');
  syncNotifToggle();
  if(dailyNotif){
    const ok=await scheduleDailyNotif();
    if(!ok && !capLN()){
      alert('デイリー通知はアプリ版でご利用いただけます。（Web版では通知を予約できません）');
    }else if(!ok){
      alert('通知の許可が必要です。端末の設定から通知を許可してください。');
      dailyNotif=false; localStorage.setItem('mlb_daily_notif','0'); syncNotifToggle();
    }
  }else{
    cancelDailyNotif();
  }
}
async function renderStandings(){
  const el=$('standings-body');
  if(!el) return;
  el.textContent='読み込み中…';
  try{
    const y=new Date().getFullYear();
    const d=await cachedJson('https://statsapi.mlb.com/api/v1/standings?leagueId=103,104&season='+y+'&standingsTypes=regularSeason','mlb_standings',30);
    const recs=d.records||[];
    if(!recs.length){ el.textContent='順位表を取得できませんでした'; return; }
    // 地区ID順（ア東・ア中・ア西・ナ東・ナ中・ナ西）
    const order=[201,202,200,204,205,203];
    recs.sort((a,b)=>order.indexOf(a.division.id)-order.indexOf(b.division.id));
    el.innerHTML=recs.map(rec=>{
      const rows=rec.teamRecords.map(tr=>
        `<tr><td><span class="stand-team">${teamLogo(tr.team)}<span class="game-team-name">${esc(teamJa(tr.team))}</span></span></td><td>${tr.wins}</td><td>${tr.losses}</td><td>${esc(tr.winningPercentage)}</td><td>${esc(tr.gamesBack)}</td></tr>`
      ).join('');
      return `<div class="stand-div"><div class="stand-div-hd">${esc(DIV_JA[rec.division.id]||'')}</div>
        <table class="stand-table"><tr><th>チーム</th><th>勝</th><th>敗</th><th>率</th><th>差</th></tr>${rows}</table></div>`;
    }).join('');
  }catch(e){
    el.textContent='オフラインのため順位表を表示できません';
  }
}
// クイズ系のどのサブ画面にいたか（menu / play / daily / review）
let lastQuizSection='menu';
// 下部ナビ「クイズ」: 直前のクイズ系サブ画面へ戻る
function goQuizNav(){
  // すでに復習／今日の10問を開いている時に「クイズ」を押したら、クイズ本体（目次）へ移動
  if(tab==='review'||tab==='daily'){ lastQuizSection='menu'; goTab('quiz'); return; }
  if(lastQuizSection==='daily') goTab('daily');
  else if(lastQuizSection==='review') goTab('review');
  else goTab('quiz');
}
// 目次へ戻る（今日の10問・復習の「← 目次」用）
function quizMenu(){ lastQuizSection='menu'; goTab('quiz'); }
// ホームの「クイズ」から目次を飛ばして直接プレイ開始
function homeStartQuiz(){ goTab('quiz'); startQuizPlay(); }
// 「成績・ランキング」開閉
function toggleStats(){
  const b=$('stats-body'), c=$('stats-chev'); if(!b) return;
  const open=b.style.display==='none';
  b.style.display=open?'':'none';
  if(c) c.classList.toggle('open',open);
  if(open) renderStats();
}
// 過去ランキング開閉
function toggleRankHistory(){
  const b=$('rank-history'), c=$('rank-hist-chev'); if(!b) return;
  const open=b.style.display==='none';
  b.style.display=open?'':'none';
  if(c) c.textContent=open?'▲':'▼';
  if(open) renderRankHistory();
}
// 直近1週間（今日を除く）の最終順位を表示。フィールドが保存されている日だけ出す
function renderRankHistory(){
  const box=$('rank-history'); if(!box) return;
  const myNick=(getProfile()&&getProfile().nick)||'あなた';
  let hist={}; try{ hist=JSON.parse(localStorage.getItem('mlb_rank_hist')||'{}'); }catch(e){}
  const today=todayStr();
  const days=[];
  for(let i=1;i<=7;i++){
    const d=new Date(); d.setDate(d.getDate()-i);
    const ds=d.getFullYear()+'-'+String(d.getMonth()+1).padStart(2,'0')+'-'+String(d.getDate()).padStart(2,'0');
    let f=null; try{ f=JSON.parse(localStorage.getItem('mlb_field_'+ds)); }catch(e){}
    if(f&&f.finals&&f.names) days.push({ds,f,d});
  }
  if(!days.length){ box.innerHTML='<div class="rh-empty">まだ過去の記録がありません（明日以降にたまっていきます）</div>'; return; }
  const wd='日月火水木金土';
  box.innerHTML=days.map(({ds,f,d})=>{
    // 最終順位＝finals（23:59時点）。名前は rank 順に対応
    const rows=f.finals.map((sc,i)=>({name:f.names[i]||'—',score:sc,me:false}));
    const us=(hist[ds]!=null)?hist[ds]:null;
    let merged=rows.slice();
    if(us!=null) merged=merged.concat([{name:myNick,score:us,me:true}]);
    merged.sort((a,b)=>b.score-a.score||(a.me?1:-1));
    const myRank=(us!=null)?(merged.findIndex(x=>x.me)+1):null;
    const top3=merged.slice(0,3).map((x,i)=>
      '<div class="rh-line'+(x.me?' me':'')+'"><span class="rh-pos">'+(i+1)+'</span>'+
      '<span class="rh-nm">'+esc(x.name)+(x.me?'（あなた）':'')+'</span>'+
      '<span class="rh-sc">'+x.score+'点</span></div>').join('');
    const label=(d.getMonth()+1)+'/'+d.getDate()+'（'+wd[d.getDay()]+'）';
    const meLine=(us!=null)
      ? '<div class="rh-me">あなた：'+myRank+'位（'+us+'点）</div>'
      : '<div class="rh-me">あなた：この日は参加なし</div>';
    return '<div class="rh-day"><div class="rh-date">'+label+'</div>'+top3+meLine+'</div>';
  }).join('');
}
// 「点の入り方」開閉
function toggleHowto(){
  const b=$('howto-body'), c=$('howto-chev'); if(!b) return;
  const open=b.style.display==='none';
  b.style.display=open?'':'none';
  if(c) c.classList.toggle('open',open);
}
// 進行中のクイズ画面を再表示（状態はリセットしない）
function quizShowPlay(){
  const play=$('quiz-play'),menu=$('quiz-menu');
  if(menu) menu.style.display='none';
  if(play) play.style.display='';
}

// ── Dict mode ─────────────────────────────────────────────────────────────────
let dictFav=false;   // 辞書内♥（お気に入り）表示中か
function updateDictTabs(){
  $('dict-tab-ej').classList.toggle('active',!dictFav&&dictMode==='ej');
  $('dict-tab-je').classList.toggle('active',!dictFav&&dictMode==='je');
  const f=$('dict-tab-fav'); if(f) f.classList.toggle('active',dictFav);
  const bb=$('fav-back-bar'); if(bb) bb.style.display=dictFav?'':'none';
}
function setDictMode(m){
  dictFav=false;
  dictMode=m;
  updateDictTabs();
  $('idx-bar').style.display='';
  $('s-input').value=''; sq='';
  $('dict-list').scrollTop=0;
  renderDictList(); renderIdxBar();
}
let _dictSavedTop=0; // お気に入りに入る直前の辞書スクロール位置
// ♥をもう一度押したら元の辞書（英和/和英）に戻る
function toggleDictFav(){
  if(dictFav) exitDictFav();
  else showDictFav();
}
function showDictFav(){
  if(!dictFav) _dictSavedTop=$('dict-list').scrollTop; // 位置を記憶
  dictFav=true;
  updateDictTabs();
  $('idx-bar').style.display='none';
  $('s-input').value=''; sq=''; $('s-clear').classList.remove('show');
  $('dict-list').scrollTop=0;
  renderDictFav();
}
// お気に入りから元の辞書へ戻る（直前のスクロール位置を復元）
function exitDictFav(){
  dictFav=false;
  updateDictTabs();
  $('idx-bar').style.display='';
  renderDictList(); renderIdxBar();
  $('dict-list').scrollTop=_dictSavedTop;
}
// favKey → {エントリ, モード} の逆引き（初回のみ構築）
let _favLookup=null;
function getFavLookup(){
  if(!_favLookup){
    _favLookup=new Map();
    EJ.forEach(e=>_favLookup.set(favKey(e),{e,m:'ej'}));
    JE.forEach(e=>{const k=favKey(e);if(!_favLookup.has(k))_favLookup.set(k,{e,m:'je'});});
  }
  return _favLookup;
}
// 辞書内お気に入り一覧（追加した順・新しいものが上）
function renderDictFav(){
  const lk=getFavLookup();
  // favsは追加順のSet。逆順で新しい順に
  const all=[...favs].reverse().map(k=>lk.get(k)).filter(Boolean);
  const list=$('dict-list');
  if(!all.length){
    list.innerHTML=`<div class="empty"><img src="images/デフォ.png" alt=""><p>お気に入りに追加した単語が<br>ここに表示されます</p></div>`;
    return;
  }
  const frag=document.createDocumentFragment();
  all.forEach(({e,m})=>{
    const head=m==='ej'?e.en:e.ja;
    const sub =m==='ej'?e.ja:(e.en+(e.def?' — '+e.def:''));
    const item=document.createElement('div');
    item.className='entry-item';
    item.innerHTML=`
      <div class="entry-text">
        <div class="entry-head">${esc(head)}<span class="entry-badge ${m}">${m==='ej'?'英和':'和英'}</span></div>
        <div class="entry-sub">${esc(formatDef(sub))}</div>
      </div>
      <button class="fav-btn on" onclick="event.stopPropagation();toggleFavBtn(this,${JSON.stringify(JSON.stringify(e))})">
        <svg viewBox="0 0 24 24" fill="currentColor"><path d="M20.84 4.61a5.5 5.5 0 00-7.78 0L12 5.67l-1.06-1.06a5.5 5.5 0 00-7.78 7.78l1.06 1.06L12 21.23l7.78-7.78 1.06-1.06a5.5 5.5 0 000-7.78z"/></svg>
      </button>`;
    item.addEventListener('click',()=>openModal(e,m));
    frag.appendChild(item);
  });
  list.innerHTML='';list.appendChild(frag);
  list.scrollTop=0;
}
$('s-input').addEventListener('input',e=>{
  if(dictFav){ dictFav=false; updateDictTabs(); $('idx-bar').style.display=''; }
  sq=e.target.value;
  $('s-clear').classList.toggle('show',sq.length>0);
  renderDictList();
  saveState();
});
function clearSearch(){
  // Find the top-ranked search result (same ranking as renderDictList) before clearing
  let targetLtr=null;
  if(sq.trim()){
    const data=dictMode==='ej'?EJ:JE;
    const q=sq.trim().toLowerCase();
    let best=null;
    data.forEach(e=>{
      const head=(dictMode==='ej'?(e.en||''):(e.ja||'')).toLowerCase();
      const body=(dictMode==='ej'?(e.ja||''):((e.en||'')+' '+(e.def||''))).toLowerCase();
      let rank=0,sub=0;
      if(head.includes(q)){rank=1;sub=head.startsWith(q)?0:1;}
      else if(body.includes(q)) rank=2;
      else return;
      const len=head.length;
      if(!best||rank<best.rank||(rank===best.rank&&sub<best.sub)||(rank===best.rank&&sub===best.sub&&len<best.len))
        best={e,rank,sub,len};
    });
    if(best){
      const h=dictMode==='ej'?(best.e.sk||best.e.en):best.e.ja;
      targetLtr=dictMode==='ej'?h[0].toUpperCase():kanaRow(h[0]);
    }
  }
  sq='';$('s-input').value='';
  $('s-clear').classList.remove('show');
  renderDictList();
  if(targetLtr) scrollToLetter(targetLtr);
  $('s-input').focus();
}

// ── Index bar ─────────────────────────────────────────────────────────────────
function renderIdxBar(){
  const bar=$('idx-bar');
  const isJA=dictMode==='je';
  const letters=isJA?JA_KANA_INDEX:'ABCDEFGHIJKLMNOPQRSTUVWXYZ'.split('');
  bar.classList.toggle('ja-full',isJA);
  bar.innerHTML=letters.map(l=>{
    const isHead=isJA&&_KANA_ROW_HEADS.includes(l);
    return `<span class="idx-l${isHead?' row-head':''}" ontouchstart="this.classList.add('hit')" ontouchend="this.classList.remove('hit')" onclick="scrollToLetter('${l}')">${l}</span>`;
  }).join('');
}
function scrollToLetter(l){
  // 検索中なら先にクリアしてインデックス表示に戻す
  if(sq){
    sq=''; $('s-input').value=''; $('s-clear').classList.remove('show');
    renderDictList();
    requestAnimationFrame(()=>setTimeout(()=>_scrollToLetterDirect(l),80));
    return;
  }
  _scrollToLetterDirect(l);
}
function _scrollToLetterDirect(l){
  const list=$('dict-list');
  const el=list.querySelector(`[data-letter="${l}"]`);
  if(!el) return;
  // sec-hdr は position:sticky のため、そのままだと offsetTop/rect が
  // 固定表示位置にズレる（特に上方向ジャンプが効かない）。一時的に
  // sticky を解除して本来のコンテンツ内オフセットを取得する。
  const prev=el.style.position;
  el.style.position='static';
  const top=el.offsetTop;
  el.style.position=prev;
  list.scrollTop=top;
}

// ── Render dict list ──────────────────────────────────────────────────────────
function scoreEntries(data,mode,q){
  // ハイフン正規化: 「pinch hitter」で「pinch-hitter」も、逆も検索できるように
  const qN=q.replace(/-/g,' ');
  const scored=[];
  data.forEach(e=>{
    const head=(mode==='ej'?(e.en||''):(e.ja||'')).toLowerCase();
    const headN=head.replace(/-/g,' ');
    const body=(mode==='ej'?(e.ja||''):((e.en||'')+' '+(e.def||''))).toLowerCase();
    const bodyN=body.replace(/-/g,' ');
    let rank=0;
    if(head.includes(q)||headN.includes(qN))      rank=1;
    else if(body.includes(q)||bodyN.includes(qN)) rank=2;
    else return;
    const sub=rank===1?((head.startsWith(q)||headN.startsWith(qN))?0:1):0;
    scored.push({e,mode,rank,sub,len:head.length});
  });
  scored.sort((a,b)=>a.rank!==b.rank?a.rank-b.rank:a.sub!==b.sub?a.sub-b.sub:a.len-b.len);
  return scored;
}
function makeEntryItem(entry,mode,showBadge){
  const head=mode==='ej'?entry.en:entry.ja;
  // 拡充済み(short)があれば一覧の説明も短文を使う（cf.等の生テキストを出さない）
  const sub =mode==='ej'?(entry.short||entry.ja)
                        :(entry.short||(entry.en+(entry.def?' — '+entry.def:'')));
  const favd=isFav(entry);
  const item=document.createElement('div');
  item.className='entry-item';
  const badge=showBadge?`<span class="entry-badge ${mode}">${mode==='ej'?'英和':'和英'}</span>`:'';
  item.innerHTML=`
    <div class="entry-text">
      <div class="entry-head">${esc(head)}${badge}</div>
      <div class="entry-sub">${esc(formatDef(sub))}</div>
    </div>
    <button class="fav-btn${favd?' on':''}" onclick="event.stopPropagation();toggleFavBtn(this,${JSON.stringify(JSON.stringify(entry))})">
      <svg viewBox="0 0 24 24" fill="${favd?'currentColor':'none'}"><path d="M20.84 4.61a5.5 5.5 0 00-7.78 0L12 5.67l-1.06-1.06a5.5 5.5 0 00-7.78 7.78l1.06 1.06L12 21.23l7.78-7.78 1.06-1.06a5.5 5.5 0 000-7.78z"/></svg>
    </button>`;
  item.addEventListener('click',()=>openModal(entry,mode));
  return item;
}
function renderDictList(){
  const q=sq.trim().toLowerCase();
  const list=$('dict-list');
  const frag=document.createDocumentFragment();
  if(!q){
    // 通常表示: 現在のdictModeのデータをアルファベット/かな順で
    const data=dictMode==='ej'?EJ:JE;
    let lastL=null;
    data.forEach(entry=>{
      const sortHead=dictMode==='ej'?(entry.sk||entry.en):(entry.yomi||entry.ja);
      const ltr=dictMode==='ej'?sortHead[0].toUpperCase():kanaRow(sortHead[0]);
      if(ltr!==lastL){
        lastL=ltr;
        const h=document.createElement('div');
        h.className='sec-hdr';h.dataset.letter=ltr;h.textContent=ltr;
        frag.appendChild(h);
      }
      frag.appendChild(makeEntryItem(entry,dictMode));
    });
  } else {
    // 検索時: 英和・和英を統合してランク順に表示
    const ejResults=scoreEntries(EJ,'ej',q);
    const jeResults=scoreEntries(JE,'je',q);
    const all=[...ejResults,...jeResults];
    all.sort((a,b)=>a.rank!==b.rank?a.rank-b.rank:a.sub!==b.sub?a.sub-b.sub:a.len-b.len);
    if(all.length>0){
      all.forEach(x=>frag.appendChild(makeEntryItem(x.e,x.mode,true)));
    } else {
      const d=document.createElement('div');
      d.className='no-results';d.textContent='該当なし';
      frag.appendChild(d);
    }
  }
  list.innerHTML='';list.appendChild(frag);
  list.scrollTop=0;
  requestAnimationFrame(cacheHeaders);
}
function toggleFavBtn(btn,ejson){
  const e=JSON.parse(ejson);
  // 辞書の♥表示中に外したら、アニメーション後に一覧から消す
  if(tab==='dict' && dictFav && isFav(e)){
    const item=btn.closest('.entry-item');
    if(item){
      item.classList.add('fav-removing');
      item.addEventListener('animationend',()=>{toggleFav(e);renderDictFav();},{once:true});
      return;
    }
  }
  toggleFav(e);
  const on=isFav(e);
  btn.classList.toggle('on',on);
  btn.querySelector('svg').setAttribute('fill',on?'currentColor':'none');
  if(tab==='dict' && dictFav) renderDictFav();
}

// ── Abbr tab ──────────────────────────────────────────────────────────────────
let abbrQ = '';
function renderAbbrList(){
  const q = abbrQ.trim().toLowerCase();
  const data = q ? ABBR.filter(e => e.en.toLowerCase().includes(q) || (e.ja||'').toLowerCase().includes(q) || (e.def||'').includes(q)) : ABBR;
  const list = $('abbr-list');
  const bar = $('abbr-idx-bar');
  if(!data.length){
    list.innerHTML = '<div class="no-results" style="padding:24px 14px;color:var(--c-muted)">該当なし</div>';
    if(bar) bar.innerHTML = '';
    return;
  }
  const frag = document.createDocumentFragment();
  const seenLetters = new Set();
  data.forEach(entry => {
    const letter = entry.en[0].toUpperCase();
    if(!seenLetters.has(letter)){
      seenLetters.add(letter);
      const hdr = document.createElement('div');
      hdr.className = 'sec-hdr';
      hdr.id = 'abbr-sec-' + letter;
      hdr.textContent = letter;
      frag.appendChild(hdr);
    }
    const item = document.createElement('div');
    item.className = 'entry-item';
    const favd = isFav(entry);
    item.innerHTML = `
      <div class="entry-text">
        <div class="entry-head">${esc(entry.en)}</div>
        <div class="entry-sub">${esc(entry.ja)}${entry.def ? ' — ' + esc(entry.def) : ''}</div>
      </div>
      <button class="fav-btn${favd?' on':''}" onclick="event.stopPropagation();toggleFavBtn(this,${JSON.stringify(JSON.stringify(entry))})">
        <svg viewBox="0 0 24 24" fill="${favd?'currentColor':'none'}"><path d="M20.84 4.61a5.5 5.5 0 00-7.78 0L12 5.67l-1.06-1.06a5.5 5.5 0 00-7.78 7.78l1.06 1.06L12 21.23l7.78-7.78 1.06-1.06a5.5 5.5 0 000-7.78z"/></svg>
      </button>`;
    item.addEventListener('click', () => openModal(entry, 'ej'));
    frag.appendChild(item);
  });
  list.innerHTML = '';
  list.appendChild(frag);
  if(bar){
    const letters = [...seenLetters].sort();
    bar.innerHTML = letters.map(l =>
      `<span class="idx-l" onclick="scrollToAbbrLetter('${l}')">${l}</span>`
    ).join('');
  }
}
function scrollToAbbrLetter(l){
  const list=$('abbr-list');
  const t=document.getElementById('abbr-sec-'+l);
  if(list&&t){
    // sec-hdr の sticky を一時解除して本来のオフセットで正確にジャンプ
    const prev=t.style.position;
    t.style.position='static';
    const top=t.offsetTop;
    t.style.position=prev;
    list.scrollTop=top;
  }
}
function clearAbbrSearch(){
  abbrQ = ''; $('abbr-search').value = '';
  $('abbr-clear').classList.remove('show');
  renderAbbrList();
}
// ── アクセス解析（Firebase / GA4）──────────────────────────────────────────
// 匿名の集計のみを送信（性別・年代帯・イベント。氏名/メール等の個人情報は送らない）。
// GA_ID が空の間は完全に無効＝何も読み込まず・何も送りません。
const GA_ID='G-X7WJVPL5HQ'; // FirebaseウェブアプリのMeasurement ID（mbl-dict）
let _gaReady=false;
function initAnalytics(){
  try{
    if(!GA_ID) return; // 未設定なら無効
    if(location.protocol==='file:'||/^(localhost$|127\.|192\.168\.)/.test(location.hostname)) return; // ローカルは計測しない
    const s=document.createElement('script'); s.async=true;
    s.src='https://www.googletagmanager.com/gtag/js?id='+GA_ID; document.head.appendChild(s);
    window.dataLayer=window.dataLayer||[]; window.gtag=function(){dataLayer.push(arguments);};
    gtag('js',new Date()); gtag('config',GA_ID,{anonymize_ip:true});
    _gaReady=true;
    gaSetUserProps(); // 既存ユーザーは属性をユーザープロパティにも設定（得点等と横断分析するため）
  }catch(e){}
}
// 性別・年代をユーザープロパティとして設定（全イベントに紐づき、年代別の得点分布などが見られる）
function gaSetUserProps(){
  try{ const p=getProfile(); if(_gaReady&&window.gtag&&p) gtag('set','user_properties',{gender:p.gender||'', age_band:p.age||''}); }catch(e){}
}
// イベント送信（無効時・失敗時は黙って無視）
function track(name,params){ try{ if(_gaReady&&window.gtag) gtag('event',name,params||{}); }catch(e){} }

document.addEventListener('DOMContentLoaded', () => {
  initAnalytics();
  track('app_open',{runs_today:(typeof runsToday==='number'?runsToday:0), runs_total:(typeof runsTotal==='number'?runsTotal:0)}); // 起動時の今日/累計得点スナップショット
  const inp = $('abbr-search');
  if(inp){
    inp.addEventListener('input', e => {
      abbrQ = e.target.value;
      $('abbr-clear').classList.toggle('show', abbrQ.length > 0);
      renderAbbrList();
    });
  }
  if(typeof syncSwingToggle==='function') syncSwingToggle();
  if(typeof syncNotifToggle==='function') syncNotifToggle();
  if(typeof syncPremium==='function') syncPremium();
  if(typeof syncSfxToggle==='function') syncSfxToggle();
  if(typeof loadSample==='function'){ loadSample('bat','sounds/bat.mp3?v=wood'); loadSample('cheer','sounds/cheer.mp3?v=1'); }
  initPullToRefresh();
  // 初回タッチでWeb Audioを起動（suspended解除＋無音再生）→ 以後の効果音が遅延なく鳴る
  const _unlockAudio=()=>{
    try{ const c=actx(); if(c){ const b=c.createBuffer(1,1,22050); const s=c.createBufferSource(); s.buffer=b; s.connect(c.destination); s.start(0); } }catch(e){}
    document.removeEventListener('touchstart',_unlockAudio); document.removeEventListener('mousedown',_unlockAudio);
  };
  document.addEventListener('touchstart',_unlockAudio,{passive:true});
  document.addEventListener('mousedown',_unlockAudio);
});
// ── プルダウン更新（スクロール先頭で下に引くと最新に更新）──────────────────────
function initPullToRefresh(){
  try{
    const ind=document.createElement('div'); ind.id='ptr-ind';
    ind.innerHTML='<span id="ptr-txt">↓ 引っ張って更新</span>';
    document.body.appendChild(ind);
    const txt=()=>document.getElementById('ptr-txt');
    const TH=70;
    let startY=0, scroller=null, dist=0, pulling=false;
    function scrollableAncestor(el){
      while(el && el!==document.body && el!==document.documentElement){
        const s=getComputedStyle(el).overflowY;
        if((s==='auto'||s==='scroll') && el.scrollHeight>el.clientHeight+1) return el;
        el=el.parentElement;
      }
      return document.querySelector('.tab-panel.active');
    }
    function overlayOpen(){
      const ob=document.getElementById('onboard'); if(ob && ob.style.display!=='none') return true;
      return !!document.querySelector('#ad-overlay.show, #versus-board.show, #versus-result.show, #versus-toast.show, #modal.open, #modal-backdrop.open');
    }
    document.addEventListener('touchstart',e=>{
      if(e.touches.length!==1 || overlayOpen()){ scroller=null; return; }
      scroller=scrollableAncestor(e.target);
      startY=e.touches[0].clientY; dist=0; pulling=false;
    },{passive:true});
    document.addEventListener('touchmove',e=>{
      if(!scroller) return;
      const atTop=(scroller.scrollTop||0)<=0;
      const dy=e.touches[0].clientY-startY;
      if(atTop && dy>6){
        pulling=true; dist=Math.min(dy*0.5,110);
        ind.classList.add('show');
        ind.style.transform='translateX(-50%) translateY('+(dist-52)+'px)';
        txt().textContent = dist>=TH ? '離して更新 🔄' : '↓ 引っ張って更新';
        if(e.cancelable) e.preventDefault();
      } else if(pulling && dy<=0){
        pulling=false; ind.classList.remove('show'); ind.style.transform='';
      }
    },{passive:false});
    const end=()=>{
      if(pulling && dist>=TH){
        txt().textContent='更新中…';
        ind.style.transform='translateX(-50%) translateY(12px)';
        setTimeout(()=>{ const b=location.origin+location.pathname; location.replace(b+'?r='+Date.now()); },150);
      }else{
        ind.classList.remove('show'); ind.style.transform='';
      }
      pulling=false; scroller=null; dist=0;
    };
    document.addEventListener('touchend',end);
    document.addEventListener('touchcancel',end);
  }catch(e){}
}

// ── Modal (bottom sheet) ──────────────────────────────────────────────────────
function formatDef(text){
  if(!text) return '';
  let s=text.replace(/[０-９]/g,c=>String.fromCharCode(c.charCodeAt(0)-0xFF10+48));
  // 数字＋ピリオドを全角に変換（スコア等 "20 – 6." は除外：直前が数字/スペース/–）
  s=s.replace(/(?<![0-9\-– ])(\d+)\.(?![0-9a-zA-Z])/g,'$1．');
  // 全角ピリオドの項目番号の前に改行を挿入（文末句点・閉じ括弧・閉じ引用符の後のみ）
  s=s.replace(/([。）」」\)])\s*(\d+．)/g,'$1\n$2');
  return s.trim();
}
// EJ逆引きマップ（= cross-ref リンク用）
let EJ_MAP=null;
function getEJMap(){
  if(!EJ_MAP){
    EJ_MAP=new Map();
    EJ.forEach((e,i)=>{
      const k=e.en.toLowerCase();
      EJ_MAP.set(k,i);
      // "A/B" や "A, B" 形式のエントリは各バリアントでも引けるようにする
      const seps=k.includes('/')?k.split('/'):k.includes(', ')?k.split(', '):null;
      if(seps) seps.forEach(part=>{
        const p=part.trim();
        if(p && !EJ_MAP.has(p)) EJ_MAP.set(p,i);
      });
    });
  }
  return EJ_MAP;
}
// JE逆引きマップ（〇〇参照 リンク用）
let JE_MAP=null;
function getJEMap(){
  if(!JE_MAP){
    JE_MAP=new Map();
    JE.forEach((e,i)=>{
      JE_MAP.set(e.ja,i);
      // 中点なしのバリアントも登録
      const noDot=e.ja.replace(/・/g,'');
      if(noDot!==e.ja && !JE_MAP.has(noDot)) JE_MAP.set(noDot,i);
    });
  }
  return JE_MAP;
}
// 英語→JEエントリ（JEの英訳から引く。EJ見出しに無い英語用語のリンク先解決用）
let JE_EN_MAP=null;
function getJEEnMap(){
  if(!JE_EN_MAP){
    JE_EN_MAP=new Map();
    JE.forEach((e,i)=>{
      if(!e.en) return;
      // 「A, B」「A/B」等の区切りで複数英訳がある場合は各語を登録
      e.en.split(/[,、\/]/).forEach(part=>{
        const k=part.trim().toLowerCase().replace(/\s*\(.*?\)\s*/g,'').trim();
        if(k.length>=2 && !JE_EN_MAP.has(k)) JE_EN_MAP.set(k,i);
        const kAlt=k.includes('-')?k.replace(/-/g,' '):k.replace(/ /g,'-');
        if(kAlt!==k && !JE_EN_MAP.has(kAlt)) JE_EN_MAP.set(kAlt,i);
      });
    });
  }
  return JE_EN_MAP;
}
function defToHtml(text,mode){
  const s=formatDef(text);
  if(!s) return '';
  // JEモード: 〇〇参照 / 〇〇ともいう パターンをリンク化
  if(mode==='je'){
    const jeMap=getJEMap();
    const ejMap=getEJMap();
    let out='',last=0;
    // グループ: 1=「」参照term, 2=裸参照term, 3=英語ともいう, 4=日本語ともいう, 5=対立語は, 6=対義語は
    const reJA=/[「｢]([^」｣]+)[」｣]参照|([ぁ-んァ-ヶー一-龯・（）()ａ-ｚＡ-Ｚa-zA-Z0-9]+(?:[ 　][ぁ-んァ-ヶー一-龯・（）()ａ-ｚＡ-Ｚa-zA-Z0-9]+)*)参照|([A-Za-z][A-Za-z0-9 \-\/\(\)\']*[A-Za-z0-9])(とも(?:いう|言う))|([ぁ-んァ-ヶー一-龯ー・（）]+)(とも(?:いう|言う))|(対立語|対義語)は([ぁ-んァ-ヶー一-龯ー・（）]+)|cf\.[ ]?([ぁ-んァ-ヶー一-龯・（）()ａ-ｚＡ-Ｚa-zA-Z0-9\-]+(?:[ 　][ぁ-んァ-ヶー一-龯・（）()ａ-ｚＡ-Ｚa-zA-Z0-9\-]+)*)/g;
    let m;
    while((m=reJA.exec(s))!==null){
      out+=esc(s.slice(last,m.index));
      if(m[3]!==undefined){
        // 英語 + ともいう
        const term=m[3].trim();
        const key=term.toLowerCase();
        const hit=ejMap.has(key)?ejMap.get(key):ejMap.has(key+' (the)')?ejMap.get(key+' (the)'):-1;
        if(hit>=0){
          out+=`<a class="def-xref" onclick="openXRef(${hit},'ej')">${esc(term)}</a>`+esc(m[4]);
        } else {
          const q=term.replace(/\\/g,'\\\\').replace(/'/g,"\\'");
          out+=`<a class="def-xref def-xref-search" onclick="xrefSearch('${q}')">${esc(term)}</a>`+esc(m[4]);
        }
      } else if(m[5]!==undefined){
        // 日本語 + ともいう
        const term=m[5].trim();
        const idx=jeMap.has(term)?jeMap.get(term):undefined;
        if(idx!==undefined){
          out+=`<a class="def-xref" onclick="openXRef(${idx},'je')">${esc(term)}</a>`+esc(m[6]);
        } else {
          out+=esc(m[0]);
        }
      } else if(m[7]!==undefined){
        // 対立語は〇〇 / 対義語は〇〇
        const label=m[7]; // 「対立語」または「対義語」
        const term=m[8].trim();
        const termNoDot=term.replace(/・/g,'');
        const idx=jeMap.has(term)?jeMap.get(term):jeMap.has(termNoDot)?jeMap.get(termNoDot):undefined;
        if(idx!==undefined){
          out+=esc(label)+'は'+`<a class="def-xref" onclick="openXRef(${idx},'je')">${esc(term)}</a>`;
        } else {
          const q=term.replace(/\\/g,'\\\\').replace(/'/g,"\\'");
          out+=esc(label)+'は'+`<a class="def-xref def-xref-search" onclick="xrefSearch('${q}')">${esc(term)}</a>`;
        }
      } else if(m[9]!==undefined){
        // cf.〇〇（JEモード内のcf.参照：日本語・英語どちらも）
        const term=m[9].trim();
        const termNoDot=term.replace(/・/g,'');
        const idx=jeMap.has(term)?jeMap.get(term):jeMap.has(termNoDot)?jeMap.get(termNoDot):undefined;
        if(idx!==undefined){
          out+='cf.'+`<a class="def-xref" onclick="openXRef(${idx},'je')">${esc(term)}</a>`;
        } else {
          const ejKey=term.toLowerCase();
          const ejKeyAlt=ejKey.includes('-')?ejKey.replace(/-/g,' '):ejKey.replace(/ /g,'-');
          const ejHit=ejMap.has(ejKey)?ejMap.get(ejKey):ejMap.has(ejKeyAlt)?ejMap.get(ejKeyAlt):-1;
          const jeEn=getJEEnMap();
          const jeEnHit=jeEn.has(ejKey)?jeEn.get(ejKey):jeEn.has(ejKeyAlt)?jeEn.get(ejKeyAlt):-1;
          if(ejHit>=0){
            out+='cf.'+`<a class="def-xref" onclick="openXRef(${ejHit},'ej')">${esc(term)}</a>`;
          } else if(jeEnHit>=0){
            out+='cf.'+`<a class="def-xref" onclick="openXRef(${jeEnHit},'je')">${esc(term)}</a>`;
          } else {
            const q=term.replace(/\\/g,'\\\\').replace(/'/g,"\\'");
            out+='cf.'+`<a class="def-xref def-xref-search" onclick="xrefSearch('${q}')">${esc(term)}</a>`;
          }
        }
      } else {
        // 〇〇参照 パターン
        const inBracket=m[1]!==undefined;
        const rawTerm=(inBracket?m[1]:m[2]).trim();
        // 末尾の助詞・余分な語を除去して検索
        let term=rawTerm
          .replace(/^詳細は/, '')      // 「詳細はXXX」→「XXX」
          .replace(/^combination/, '') // 「combinationXXX」→「XXX」
          .replace(/[をはがのへでとにもや]$|から$|まで$|より$/, '') // 末尾助詞を先に除去
          .replace(/の項目$/, '')       // 「XXXの項目」→「XXX」
          .replace(/の項$/, '')         // 「XXXの項」→「XXX」
          .replace(/試合$/, '')         // 「ノーヒットノーラン試合」→「ノーヒットノーラン」
          .replace(/[をはがのへでとにもや]$|から$|まで$|より$/, '') // の項目除去後も再チェック
          .trim();
        const termNorm=term.replace(/\s*\(\s*/g,'(').replace(/\s*\)\s*/g,')');
        const ob=inBracket?(m[0][0]==='｢'?'｢':'「'):'';
        const cb=inBracket?(m[0][0]==='｢'?'｣':'」'):'';
        // JEマップで検索（元の語・助詞除去・中点除去・正規化の順）
        const termNoDot=term.replace(/・/g,'');
        const idx=jeMap.has(rawTerm)?jeMap.get(rawTerm)
                 :jeMap.has(term)?jeMap.get(term)
                 :jeMap.has(termNoDot)?jeMap.get(termNoDot)
                 :jeMap.has(termNorm)?jeMap.get(termNorm)
                 :undefined;
        if(idx!==undefined){
          const link=`<a class="def-xref" onclick="openXRef(${idx},'je')">${esc(term)}</a>`;
          out+=esc(ob)+link+esc(cb)+'参照';
        } else {
          // JEにない場合はEJマップも試す（英語名など: Babe Ruth参照）
          const ejKey=term.toLowerCase();
          const ejHit=ejMap.has(ejKey)?ejMap.get(ejKey):ejMap.has(ejKey.replace(/-/g,' '))?ejMap.get(ejKey.replace(/-/g,' ')):-1;
          if(ejHit>=0){
            out+=esc(ob)+`<a class="def-xref" onclick="openXRef(${ejHit},'ej')">${esc(term)}</a>`+esc(cb)+'参照';
          } else {
            out+=esc(m[0]);
          }
        }
      }
      last=m.index+m[0].length;
    }
    out+=esc(s.slice(last));
    return out.replace(/\n/g,'<br>');
  }
  const map=getEJMap();
  function _lookupKey(key){
    const keyAlt=key.includes('-')?key.replace(/-/g,' '):key.replace(/ /g,'-');
    const keyS=key.endsWith('s')&&key.length>2?key.slice(0,-1):null; // games→game
    return map.has(key)?map.get(key)
           :map.has(keyAlt)?map.get(keyAlt)
           :map.has(key+' (the)')?map.get(key+' (the)')
           :map.has(key+' (a)')?map.get(key+' (a)')
           :(keyS&&map.has(keyS))?map.get(keyS)
           :-1;
  }
  const jeEnMap=getJEEnMap();
  function linkOneTerm(raw,fallback,_noSplit){
    const trimmed=raw.trim();
    if(!trimmed) return esc(raw);
    const words=trimmed.split(/\s+/);
    for(let n=words.length;n>=1;n--){
      const base=words.slice(0,n).join(' ');
      const hit=_lookupKey(base.toLowerCase());
      if(hit>=0){
        const rest=trimmed.slice(base.length);
        return `<a class="def-xref" onclick="openXRef(${hit},'ej')">${esc(base)}</a>${esc(rest)}`;
      }
    }
    // EJ見出しに無ければ、JEエントリの英訳から探して和英エントリへ飛ばす
    for(let n=words.length;n>=1;n--){
      const base=words.slice(0,n).join(' ');
      const k=base.toLowerCase();
      const jeHit=jeEnMap.has(k)?jeEnMap.get(k):(jeEnMap.has(k.replace(/-/g,' '))?jeEnMap.get(k.replace(/-/g,' ')):-1);
      if(jeHit>=0){
        const rest=trimmed.slice(base.length);
        return `<a class="def-xref" onclick="openXRef(${jeHit},'je')">${esc(base)}</a>${esc(rest)}`;
      }
    }
    // スラッシュ区切りを「A/B」→「Aリンク/Bリンク」に分解（1段のみ、再帰防止）
    if(!_noSplit && trimmed.includes('/')){
      const parts=trimmed.split('/');
      const linked=parts.map(p=>linkOneTerm(p.trim(),fallback,true));
      // 少なくとも1つがリンクになっていれば分割版を返す
      if(linked.some(l=>l.includes('def-xref')))
        return linked.join(esc('/'));
    }
    // 辞書に見つからなくても cf./= 文脈では検索リンクにする
    if(fallback){
      const q=trimmed.replace(/\\/g,'\\\\').replace(/'/g,"\\'");
      return `<a class="def-xref def-xref-search" onclick="xrefSearch('${q}')">${esc(trimmed)}</a>`;
    }
    return esc(trimmed);
  }
  let out='',last=0;
  // cf. の後はカンマ/読点区切りで複数ターム、= の後は単一ターム（数字始まりの略語も対応）
  const re=/(=|[Cc]f\.)\s*([A-Za-z0-9](?:(?! [Cc]f\.)[A-Za-z0-9 \-\/\(\)\'])*(?:[,、]\s*[A-Za-z0-9](?:(?! [Cc]f\.)[A-Za-z0-9 \-\/\(\)\'])*)*)/g;
  let m;
  while((m=re.exec(s))!==null){
    out+=esc(s.slice(last,m.index));
    const prefix=m[1];
    if(prefix==='='){
      const parts=m[2].split(/[,、]/);
      out+=esc('= ');
      out+=parts.map((p,i)=>(i>0?esc(', '):'')+linkOneTerm(p.trim().toLowerCase(),true)).join('');
    } else {
      const parts=m[2].split(/[,、]/);
      out+=esc('cf. ');
      out+=parts.map((p,i)=>(i>0?esc(', '):'')+linkOneTerm(p,true)).join('');
    }
    last=m.index+m[0].length;
  }
  out+=esc(s.slice(last));
  return out.replace(/\n/g,'<br>');
}
function openXRef(idx,mode){
  const arr=mode==='ej'?EJ:JE;
  if(idx>=0&&idx<arr.length) openModal(arr[idx],mode);
}
function openModal(entry,mode){
  modalEntry=entry;modalMode=mode;
  const head=mode==='ej'?entry.en:entry.ja;
  $('m-term').textContent=head;
  // 拡充データ（short/long）があれば画像パターンで表示、無ければ従来表示
  if(entry.short||entry.long){
    $('m-def').innerHTML=enrichHtml(entry,mode);
    bindRelJumps($('m-def'));
  } else {
    let def;
    if(entry.cat==='abbr'){
      def=(entry.ja||'')+(entry.def?'\n\n'+entry.def:'');
    } else {
      def=mode==='ej'?entry.ja:(entry.en+'\n\n'+(entry.def||'')).trim();
    }
    $('m-def').innerHTML=defToHtml(def,mode);
  }
  const on=isFav(entry);
  $('mfav-btn').classList.toggle('on',on);
  $('mfav-svg').setAttribute('fill',on?'currentColor':'none');
  $('modal-backdrop').classList.add('open');
  $('modal').classList.add('open');
  document.body.style.overflow='hidden';
}
// 拡充エントリを画像パターン（短文/長文/例/関連語）でHTML化
// mode='ej'|'je'（関連語のジャンプ先の解決に使う）
function enrichHtml(e,mode){
  mode=mode||'ej';
  let h='';
  if(e.short) h+='<div class="m-short">'+esc(e.short)+'</div>';
  if(e.long)  h+='<div class="m-long">'+esc(e.long)+'</div>';
  if(e.ex&&e.ex.en){
    h+='<div class="m-ex"><span class="lbl">例：</span><span class="en">'+esc(e.ex.en)+'</span>'
      +(e.ex.ja?'（'+esc(e.ex.ja)+'）':'')+'</div>';
  }
  if(e.rel&&e.rel.length){
    const links=e.rel.map(r=>'<span class="def-xref rel-jump" data-mode="'+mode+'" data-head="'+esc(r)+'">'+esc(r)+'</span>').join('、');
    h+='<div class="m-rel"><span class="lbl">関連語：</span>'+links+'</div>';
  }
  return h;
}
// 関連語タップでその見出しへジャンプ（英和/和英でジャンプ先を切替）
function bindRelJumps(root){
  root.querySelectorAll('.rel-jump').forEach(el=>{
    el.addEventListener('click',()=>{
      if(el.getAttribute('data-mode')==='je') jumpToHeadJE(el.getAttribute('data-head'));
      else jumpToHead(el.getAttribute('data-head'));
    });
  });
}
function jumpToHead(h){
  const m=getEJMap();
  const i=m.get((h||'').toLowerCase());
  if(i!=null&&EJ[i]){ openModal(EJ[i],'ej'); return; }
  xrefSearch(h); // 念のためのフォールバック（通常は到達しない）
}
function jumpToHeadJE(h){
  const m=getJEMap();
  const i=m.get((h||'').trim());
  if(i!=null&&JE[i]){ openModal(JE[i],'je'); return; }
  xrefSearch(h);
}
function closeModal(){
  $('modal').classList.remove('open');
  $('modal-backdrop').classList.remove('open');
  document.body.style.overflow='';
  const list=$('dict-list');
  const savedTop=list.scrollTop;
  if(tab==='dict'){ if(dictFav) renderDictFav(); else renderDictList(); }
  list.scrollTop=savedTop;
  // dailyタブで回答済みならnextボタンを有効に戻す
  if(tab==='daily' && dailyAns){
    const nb=$('daily-next-btn');
    if(nb) nb.disabled=false;
  }
}
function xrefSearch(q){
  closeModal();
  // 辞書タブ + EJモードに切り替えてから検索
  document.querySelectorAll('.tab-panel').forEach(p=>p.classList.remove('active'));
  document.querySelectorAll('.bnav-btn').forEach(b=>b.classList.remove('active'));
  $('tab-dict').classList.add('active');
  $('bnav-dict').classList.add('active');
  tab='dict';
  dictMode='ej';
  ['ej','je'].forEach(x=>$('dict-tab-'+x).classList.toggle('active',x==='ej'));
  sq=q; $('s-input').value=q;
  $('s-clear').classList.add('show');
  $('dict-list').scrollTop=0;
  renderDictList(); renderIdxBar();
}
function toggleMFav(){
  if(!modalEntry) return;
  toggleFav(modalEntry);
  const on=isFav(modalEntry);
  $('mfav-btn').classList.toggle('on',on);
  $('mfav-svg').setAttribute('fill',on?'currentColor':'none');
}

// ── Quiz ──────────────────────────────────────────────────────────────────────
// クイズ状態: 連続正解数・アウト数・塁の走者[1塁,2塁,3塁]・得点
let quizStreak=0, quizOuts=0, bases=[false,false,false];
let runsTotal=parseInt(localStorage.getItem('mlb_runs')||'0',10)||0; // 累計得点
let swingAuto=localStorage.getItem('mlb_swing_auto')==='1'; // スイング操作: true=オート
function vibrate(p){try{if(navigator.vibrate)navigator.vibrate(p);}catch(e){}}
// ── 効果音（Web Audioで合成。音声ファイル不要） ──────────────────────────────
let _sfxOn=localStorage.getItem('mlb_sfx')!=='0'; // 既定ON
let _actx=null;
function actx(){
  try{
    if(!_actx) _actx=new (window.AudioContext||window.webkitAudioContext)();
    if(_actx.state==='suspended') _actx.resume();
  }catch(e){ _actx=null; }
  return _actx;
}
// 実音源（mp3）: あれば優先再生、無ければ合成音にフォールバック
const _samp={};
const _sampBuf={}; // Web Audioへデコード済みのバッファ（タップとほぼ同時に鳴らす）
function loadSample(key,src){
  try{ const a=new Audio(src); a.preload='auto'; _samp[key]=a; }catch(e){}
  // 低遅延再生用にデコードしてバッファ保持（HTML5 Audioのplay()遅延を回避）
  try{
    fetch(src).then(r=>r.arrayBuffer()).then(buf=>{
      const c=actx(); if(!c) return;
      c.decodeAudioData(buf, ab=>{ _sampBuf[key]=ab; }, ()=>{});
    }).catch(()=>{});
  }catch(e){}
}
function playSample(key,vol){
  // 1) Web Audioバッファがあれば即時再生（遅延ほぼゼロ）
  const ab=_sampBuf[key];
  if(ab){
    const c=actx();
    if(c){
      try{
        const src=c.createBufferSource(); src.buffer=ab;
        const g=c.createGain(); g.gain.value=(vol==null?1:vol);
        src.connect(g).connect(c.destination); src.start();
        return true;
      }catch(e){}
    }
  }
  // 2) フォールバック: HTML5 Audio
  const a=_samp[key]; if(!a) return false;
  try{ const n=a.cloneNode(true); n.volume=(vol==null?1:vol); const p=n.play(); if(p&&p.catch)p.catch(()=>{}); return true; }
  catch(e){ return false; }
}
function _tone(freq,t0,dur,type,gain,glideTo){
  const c=actx(); if(!c) return;
  const o=c.createOscillator(), g=c.createGain();
  o.type=type||'sine'; o.frequency.setValueAtTime(freq,t0);
  if(glideTo) o.frequency.exponentialRampToValueAtTime(Math.max(20,glideTo),t0+dur);
  g.gain.setValueAtTime(0.0001,t0);
  g.gain.exponentialRampToValueAtTime(gain||0.2,t0+0.006);
  g.gain.exponentialRampToValueAtTime(0.0001,t0+dur);
  o.connect(g).connect(c.destination);
  o.start(t0); o.stop(t0+dur+0.03);
}
function _noise(t0,dur,gain,freq,type,Q){
  const c=actx(); if(!c) return;
  const n=Math.max(1,Math.floor(c.sampleRate*dur));
  const buf=c.createBuffer(1,n,c.sampleRate); const d=buf.getChannelData(0);
  for(let i=0;i<n;i++) d[i]=Math.random()*2-1;
  const src=c.createBufferSource(); src.buffer=buf;
  const f=c.createBiquadFilter(); f.type=type||'bandpass'; f.frequency.value=freq||2000; f.Q.value=Q||1;
  const g=c.createGain();
  g.gain.setValueAtTime(gain||0.3,t0);
  g.gain.exponentialRampToValueAtTime(0.0001,t0+dur);
  src.connect(f).connect(g).connect(c.destination);
  src.start(t0); src.stop(t0+dur+0.02);
}
// カキーン: 実音源(bat.mp3)を優先、無ければ合成音
function sfxCrack(big){
  if(!_sfxOn) return;
  if(playSample('bat', big?1:0.85)) return;
  _synthCrack(big);
}
// 合成版カキーン（フォールバック）: 鋭い衝撃＋非整数倍音の金属リング
function _synthCrack(big){
  const c=actx(); if(!c||!_sfxOn) return; const t=c.currentTime;
  // 1) 「カッ」＝高域ノイズの極短バースト（打撃の立ち上がり）
  _noise(t,0.012,big?0.6:0.45,5200,'highpass',0.7);
  _noise(t,0.03,big?0.28:0.2,3200,'bandpass',1.4);
  // 2) 「キーン」＝金属的な非整数倍音（1 : 1.34 : 1.79 : 2.4）を速めに減衰
  const base=big?1850:1650, dur=big?0.6:0.42;
  [[1,big?0.32:0.24],[1.34,big?0.2:0.15],[1.79,big?0.12:0.09],[2.4,big?0.06:0.045]].forEach(([r,g])=>{
    const o=c.createOscillator(), gg=c.createGain();
    o.type='sine'; o.frequency.setValueAtTime(base*r,t);
    gg.gain.setValueAtTime(0.0001,t);
    gg.gain.exponentialRampToValueAtTime(g,t+0.004);
    gg.gain.exponentialRampToValueAtTime(0.0001,t+dur);
    o.connect(gg).connect(c.destination);
    o.start(t); o.stop(t+dur+0.03);
  });
  // 3) 「コッ」＝ボールを捉えた低域の衝撃
  _tone(180,t,0.05,'sine',0.18,90);
}
// コツン（単打のミート音）
function sfxPock(){
  const c=actx(); if(!c||!_sfxOn) return; const t=c.currentTime;
  _noise(t,0.03,0.32,1300,'bandpass',1); _tone(420,t,0.09,'sine',0.22,190);
}
// 空振り（風切り音）
function sfxWhiff(){
  const c=actx(); if(!c||!_sfxOn) return; const t=c.currentTime;
  _noise(t,0.2,0.26,1600,'lowpass',0.7); _tone(500,t,0.2,'sine',0.06,180);
}
// アウト（がっかり下降音）
function sfxOut(){
  const c=actx(); if(!c||!_sfxOn) return; const t=c.currentTime;
  _tone(320,t,0.4,'sawtooth',0.18,120);
}
// 歓声: 実音源(cheer.mp3)を優先、無ければ合成音
function sfxCheer(){
  if(!_sfxOn) return;
  if(playSample('cheer',1)) return;
  const c=actx(); if(!c) return; const t=c.currentTime;
  [330,392,523].forEach((f,i)=>_tone(f,t+i*0.11,0.24,'sine',0.16)); // E4→G4→C5
}
// ホームラン: カキーンのあとに歓声を足す
function sfxHomerun(){
  sfxCrack(true);
  setTimeout(sfxCheer,360); // 少し後に歓声（フル再生で長めに鳴る）
}
// 正解の小気味よいピコッ
function sfxSelectOK(){
  const c=actx(); if(!c||!_sfxOn) return; const t=c.currentTime;
  _tone(880,t,0.09,'sine',0.14,1320);
}
// 不正解のブブー
function sfxWrong(){
  const c=actx(); if(!c||!_sfxOn) return; const t=c.currentTime;
  _tone(200,t,0.28,'square',0.16,150); _tone(150,t+0.02,0.28,'square',0.12,120);
}
// スリーアウト（下降トロンボーン）
function sfxThreeOut(){
  const c=actx(); if(!c||!_sfxOn) return; const t=c.currentTime;
  _tone(392,t,0.5,'sawtooth',0.18,196);
}
// 打球結果（塁数）に応じた打撃音
function sfxForVal(v){
  if(v>=4) sfxHomerun();          // カキーン＋歓声
  else if(v>=2) sfxCrack(false);
  else if(v>=1) sfxPock();
  else sfxWhiff();
}
function toggleSfx(){
  _sfxOn=!_sfxOn; localStorage.setItem('mlb_sfx',_sfxOn?'1':'0');
  syncSfxToggle();
  if(_sfxOn) sfxSelectOK(); // 確認音
}
function syncSfxToggle(){
  const b=$('sfx-toggle'); if(!b) return;
  b.textContent=_sfxOn?'オン':'オフ';
  b.classList.toggle('auto',!_sfxOn);
}
// ── 難易度（進む塁数 1=単打 / 2=二塁打 / 3=三塁打）──
// 主要語の手調整（自動判定を上書き）。enの小文字・括弧除去で照合
const DIFF_OVERRIDE={
  // やさしい（1）
  'hit':1,'run':1,'out':1,'ball':1,'strike':1,'base':1,'walk':1,'catch':1,'pitcher':1,'batter':1,
  'home run':1,'single':1,'double':1,'triple':1,'inning':1,'error':1,'save':1,'win':1,'loss':1,
  // 難しい（3）＝慣用句・スラング
  'can of corn':3,'ducks on the pond':3,'around the horn':3,'in the hole':3,'on deck':2,
  'texas leaguer':3,'frozen rope':3,'cup of coffee':3,'high and tight':3,'seventh inning stretch':3,
  'grand slam':2,'squeeze play':3,'hit and run':2,'sacrifice fly':2,'double play':2,'triple play':3,
  'pinch hitter':2,'designated hitter':2,'utility player':2,'quality start':3,'walk-off':3,'walkoff':3,
};
function difficultyOf(e){
  if(!e) return 1;
  const key=(e.en||'').toLowerCase().replace(/\(.*?\)/g,'').replace(/[=＝].*$/,'').trim();
  if(DIFF_OVERRIDE[key]!=null) return DIFF_OVERRIDE[key];
  // 訳語の先頭（説明の前）＝実際の答え部分だけでカタカナ率を見る
  const head=(e.ja||'').split(/[：:；;。\n、]/)[0];
  const kata=(head.match(/[ァ-ヶーヴ・]/g)||[]).length;
  const headLen=head.replace(/[\s（）()「」・]/g,'').length;
  const isLoan=headLen>0 && kata/headLen>=0.7; // カタカナ外来語＝やさしい
  if(isLoan) return 1;
  // 英語の単語数（単打を基本に、本当に難しいものだけ上げる）
  const words=key.split(/[\s\-\/]+/).filter(Boolean);
  const wc=words.length;
  if(wc>=4) return 3;        // 4語以上のフレーズ＝三塁打
  if(wc>=3) return 2;        // 3語＝二塁打
  if(wc===2) return 1;       // 2語の複合語は基本やさしい（単打）
  if(key.length>=14) return 2; // 非常に長い1語のみ二塁打
  return 1;
}
// 無料ユーザーの1日のプレイ回数制限
const FREE_DAILY_PLAYS=2;
function _playsToday(){ try{ return JSON.parse(localStorage.getItem('mlb_plays_today'))||{}; }catch(e){ return {}; } }
function playsLeft(){
  if(isPremium()) return Infinity;
  const c=_playsToday(); const n=(c.date===todayStr())?(c.n||0):0;
  return Math.max(0,FREE_DAILY_PLAYS-n);
}
function recordPlay(){
  if(isPremium()) return;
  const t=todayStr(), c=_playsToday();
  const n=((c.date===t)?(c.n||0):0)+1;
  localStorage.setItem('mlb_plays_today',JSON.stringify({date:t,n}));
}
function startQuizPlay(){
  // 無料は1日2回まで。超えたらプレミアム案内
  if(playsLeft()<=0){
    if(confirm('本日の無料回数（'+FREE_DAILY_PLAYS+'回）を使い切りました。\nプレミアム（月額200円）なら回数無制限・広告なしで遊べます。\n設定を開きますか？')){
      goTab('settings');
    }
    return;
  }
  recordPlay();
  lastQuizSection='play';
  $('quiz-menu').style.display='none';
  $('quiz-play').style.display='';
  const _hr=$('hud-runs'); if(_hr) _hr.style.display=''; // フリーバッティングでは本日/通算を表示
  quizStreak=0;quizOuts=0;quizAsked=0;quizGameCorrect=0;
  resetSwing();
  updateOuts();clearBases();updateRuns();
  cq=null;nextQ();
}
function resetSwing(){
  try{ sw.active=false; cancelAnimationFrame(sw.raf); }catch(e){}
  const ui=$('swing-ui'); if(ui) ui.style.display='none';
}
// 累計得点の表示更新
// 本日の得点（日付が変わると自動リセット）
function _todayRuns(){
  try{ return JSON.parse(localStorage.getItem('mlb_runs_today'))||{}; }catch(e){ return {}; }
}
let runsToday=(function(){ const c=_todayRuns(); return (c.date===todayStr())?(c.n||0):0; })();
function updateRuns(){
  const c=_todayRuns();
  runsToday=(c.date===todayStr())?(c.n||0):0;
  const a=$('quiz-runs'),b=$('menu-runs'),cc=$('daily-runs'),d=$('quiz-runs-today'),e=$('menu-runs-today'),f=$('daily-runs-today');
  if(a) a.textContent=runsTotal;
  if(b) b.textContent=runsTotal;
  if(cc) cc.textContent=runsTotal;
  if(d) d.textContent=runsToday;
  if(e) e.textContent=runsToday;
  if(f) f.textContent=runsToday;
  recordRankHist(); // 過去ランキング用に本日の得点を日付付きで保存
}
// 日別の本日得点を記録（過去ランキング表示用。直近10日分だけ保持）
function recordRankHist(){
  try{
    const hk='mlb_rank_hist'; let hist=JSON.parse(localStorage.getItem(hk)||'{}');
    hist[todayStr()]=runsToday;
    const ks=Object.keys(hist).sort(); while(ks.length>10){ delete hist[ks.shift()]; }
    localStorage.setItem(hk,JSON.stringify(hist));
  }catch(e){}
}
function addRuns(n){
  if(n<=0) return;
  // 対戦中は得点を試合側に集計（確定は試合終了時。途中離脱は没収）
  if(versusActive && vs){
    vs.inningRuns+=n; vs.userScore+=n;
    updateVersusBar();
    // 右上HUDは試合中も暫定表示（確定は試合終了時。途中離脱で没収されると元に戻る）
    const at=$('quiz-runs-today'); if(at) at.textContent=(runsToday+vs.userScore);
    const ac=$('quiz-runs');       if(ac) ac.textContent=(runsTotal+vs.userScore);
    ['quiz-runs','quiz-runs-today'].forEach(id=>{const el=$(id);if(el){el.classList.remove('run-bump');void el.offsetWidth;el.classList.add('run-bump');}});
    return;
  }
  const prevRankName=currentRank(runsTotal).r.name;
  runsTotal+=n;
  localStorage.setItem('mlb_runs',String(runsTotal));
  const newRank=currentRank(runsTotal).r;
  if(newRank.name!==prevRankName) rankUpToast(newRank.name);
  // 本日分
  const t=todayStr(), c=_todayRuns();
  runsToday=((c.date===t)?(c.n||0):0)+n;
  localStorage.setItem('mlb_runs_today',JSON.stringify({date:t,n:runsToday}));
  updateRuns();
  track('score_update',{gained:n, runs_today:runsToday, runs_total:runsTotal}); // 今日の得点・累計得点
  ['quiz-runs','quiz-runs-today'].forEach(id=>{const el=$(id);if(el){el.classList.remove('run-bump');void el.offsetWidth;el.classList.add('run-bump');}});
}
let _rankupTimer=null;
function rankUpToast(name){
  const t=$('rankup-toast'); if(!t) return;
  t.innerHTML='<div class="ru-sub">🎉 称号アップ！</div><div class="ru-name">'+esc(name)+'</div><div class="ru-sub">に昇格しました</div>';
  t.classList.add('show');
  if(typeof sfxCheer==='function'){ try{ sfxCheer(); }catch(e){} }
  clearTimeout(_rankupTimer);
  _rankupTimer=setTimeout(()=>{ t.classList.remove('show'); },2600);
}
function quizBack(){
  if(versusActive){
    // 延長（4〜8回）中の放棄は没収／3回まで（未延長）は得点を加算して終了
    if(vs && vs.maxInning===9){ versusQuitConfirm(); }
    else { versusStopEarly(); }
    return;
  }
  quizToMenu();
}
function versusStopEarly(){
  if(confirm('ここで試合を終えますか？（ここまでの得点は加算されます）')){
    showVersusResult(false);
  }
}
function quizToMenu(){
  lastQuizSection='menu';
  resetSwing();
  // 出題セッションを完全終了（裏で制限時間が切れて✖・音が鳴るのを防ぐ）
  stopPitchClock(); cq=null; ans=true;
  const play=$('quiz-play'),menu=$('quiz-menu');
  if(play) play.style.display='none';
  if(menu) menu.style.display='';
  const _hr=$('hud-runs'); if(_hr) _hr.style.display=''; // 試合で隠した本日/通算HUDを復帰
  updateRuns();
}
function updateOuts(){
  const el=$('quiz-outs');
  if(el) el.textContent='OUT '+'●'.repeat(quizOuts)+'○'.repeat(Math.max(0,3-quizOuts));
}
function clearBases(){
  bases=[false,false,false];
  renderBases();
}
// bases配列に従って走者を表示
function renderBases(){
  [1,2,3].forEach(n=>{const e=$('dchar-'+n);if(e)e.style.display=bases[n-1]?'':'none';});
}
// 塁座標（diamond-wrap内の%）: 0=ホーム,1=一塁,2=二塁,3=三塁
const BASE_POS={0:[50,80],1:[82,48],2:[50,15],3:[18,48]};
const SEG_MS=420, RUN_STAGGER=280; // 1区間の所要時間・走者ごとの出発ずらし
// 区間の向き（出発塁で決まる）: ホーム/三塁発＝右向き, 一塁/二塁発＝左向き
function segFacing(from){ return (from===0||from===3)?CHAR_IMG.runR:CHAR_IMG.runL; }
function prefersReduced(){
  try{ return window.matchMedia && window.matchMedia('(prefers-reduced-motion: reduce)').matches; }
  catch(e){ return false; }
}
// ヒット（value塁進む）: 塁上の走者＋打者を、ダイヤ上を実際に走らせて進める。{runs}を返す
function advanceRunners(value){
  // 進塁プラン（リード走者＝3塁側から順に、最後に打者）
  const journeys=[];
  for(let i=2;i>=0;i--){ if(bases[i]) journeys.push({start:i+1, end:i+1+value}); }
  journeys.push({start:0, end:value}); // 打者はホーム発
  // 新しい塁状態と得点
  const nb=[false,false,false]; let runs=0;
  journeys.forEach(j=>{ if(j.end>=4) runs++; else nb[j.end-1]=true; });
  bases=nb;
  // モーション抑制設定なら瞬間反映
  if(prefersReduced()){ renderBases(); return {runs}; }
  // アニメ中は静的走者を隠し、到着時に表示
  [1,2,3].forEach(n=>{const e=$('dchar-'+n); if(e) e.style.display='none';});
  journeys.forEach((j,idx)=>{
    const mv=createMover(j.start);
    const pathEnd=Math.min(j.end,4); // 本塁(4)で打ち止め＝生還
    setTimeout(()=>stepMover(mv, j.start, pathEnd), idx*RUN_STAGGER + 40);
  });
  // 最終状態を反映（保険）
  const maxSeg=Math.max.apply(null, journeys.map(j=>Math.min(j.end,4)-j.start));
  setTimeout(renderBases, (journeys.length-1)*RUN_STAGGER + maxSeg*SEG_MS + 220);
  return {runs};
}
// 走者エレメントを start塁 に生成
function createMover(start){
  const wrap=document.querySelector('#quiz-char .diamond-wrap');
  if(!wrap) return null;
  const mv=document.createElement('img');
  mv.className='dchar dchar-mover';
  const p=BASE_POS[start%4];
  mv.src=segFacing(start);
  mv.style.left=p[0]+'%'; mv.style.top=p[1]+'%';
  wrap.appendChild(mv);
  return mv;
}
// 走者を seg塁 から end塁 まで1区間ずつ塁パスに沿って進める
function stepMover(mv, seg, end){
  if(!mv) return;
  if(seg>=end){
    mv.remove();
    if(end<4){ const st=$('dchar-'+end); if(st) st.style.display=''; } // 到着（生還=4は表示しない）
    return;
  }
  const tp=BASE_POS[(seg+1)%4];
  mv.src=segFacing(seg%4);
  requestAnimationFrame(()=>{ mv.style.left=tp[0]+'%'; mv.style.top=tp[1]+'%'; });
  setTimeout(()=>stepMover(mv, seg+1, end), SEG_MS);
}
function threeOut(){
  const el=$('three-out');
  vibrate(250); sfxThreeOut();
  if(el){el.classList.remove('fade');el.classList.add('show');}
  // しばらく表示 → フェードアウト → 広告（非プレミアム）→ 目次
  setTimeout(()=>{
    if(el) el.classList.add('fade');
    quizStreak=0;quizOuts=0;updateOuts();clearBases();
    quizToMenu();
    // 3アウト終了後は広告を挟む（プレミアムはスキップ）
    showAd();
  },1600);
  setTimeout(()=>{if(el) el.classList.remove('show','fade');},2350);
}
// ── プレミアム＆広告 ─────────────────────────────────────────────────────────
// プレミアム（月額200円）: 広告非表示＆制限なし。実際の課金はApp Store連携で設定する。
function isPremium(){ return localStorage.getItem('mlb_premium')==='1'; }
// AdMob等のネイティブ広告プラグイン（あれば使用）
function capAdMob(){
  try{ return window.Capacitor && window.Capacitor.Plugins && window.Capacitor.Plugins.AdMob; }
  catch(e){ return null; }
}
let _adCb=null,_adTimer=null;
// 広告を表示し、閉じたら cb を呼ぶ。プレミアムは即 cb。
function showAd(cb){
  if(isPremium()){ if(cb) cb(); return; }
  const admob=capAdMob();
  if(admob && admob.prepareInterstitial && admob.showInterstitial){
    // 実広告（ネイティブ）: 表示→完了で cb
    Promise.resolve()
      .then(()=>admob.prepareInterstitial({adId: (window.MLB_AD_INTERSTITIAL||'')}))
      .then(()=>admob.showInterstitial())
      .then(()=>{ if(cb) cb(); })
      .catch(()=>{ showAdPlaceholder(cb); });
    return;
  }
  showAdPlaceholder(cb);
}
// Web／プラグイン未導入時のプレースホルダ広告（5秒後に閉じられる）
function showAdPlaceholder(cb){
  _adCb=cb||null;
  const ov=$('ad-overlay'); if(!ov){ if(cb)cb(); return; }
  ov.classList.add('show');
  const btn=$('ad-close');
  let n=5;
  if(btn){ btn.disabled=true; btn.innerHTML='とじる（<span id="ad-count">'+n+'</span>）'; }
  clearInterval(_adTimer);
  _adTimer=setInterval(()=>{
    n--; const cnt=$('ad-count'); if(cnt) cnt.textContent=n;
    if(n<=0){ clearInterval(_adTimer); if(btn){btn.disabled=false; btn.textContent='とじる';} }
  },1000);
}
function closeAd(){
  clearInterval(_adTimer);
  const ov=$('ad-overlay'); if(ov) ov.classList.remove('show');
  const cb=_adCb; _adCb=null; if(cb) cb();
}
// プレミアム購入（プレースホルダ）。ネイティブではStoreKit/課金プラグインに差し替える。
function buyPremium(){
  const iap=(window.Capacitor&&window.Capacitor.Plugins&&(window.Capacitor.Plugins.Purchases||window.Capacitor.Plugins.InAppPurchase));
  if(iap){ alert('課金処理を開始します（ストア連携）'); return; } // ネイティブで実装
  // 開発／テスト用: 実課金が未連携のため、確認のうえ有効化（リリース時はStoreKitに置換）
  if(confirm('【テスト】プレミアムを有効にしますか？\n※実際の課金はApp Store連携後に有効になります')){
    localStorage.setItem('mlb_premium','1'); syncPremium();
  }
}
function cancelPremium(){
  if(confirm('プレミアムを解除しますか？（テスト用）')){ localStorage.removeItem('mlb_premium'); syncPremium(); }
}
function restorePremium(){ alert('購入の復元はアプリ版（ストア連携後）でご利用いただけます。'); }
function syncPremium(){
  const st=$('premium-status'), row=$('premium-actions');
  const on=isPremium();
  if(st) st.textContent=on?'プレミアム会員（広告なし・制限なし）':'未登録（無料）';
  if(row) row.innerHTML=on
    ? '<button class="swing-toggle auto" onclick="cancelPremium()">解除（テスト）</button>'
    : '<button class="swing-toggle" onclick="buyPremium()">登録する（月額200円）</button>';
  if(typeof updateBanner==='function') updateBanner();
}
function isPersonName(e){
  if(e.cat==='abbr') return false; // 略語エントリは除外しない
  // 定義がカタカナ・中黒・長音だけ → 人名とみなす
  const raw=(qMode==='ej'?e.ja:e.en)||'';
  const s=raw.replace(/[。参照\s]/g,'');
  return s.length>0&&/^[ァ-ヶーｦ-ﾟ・＝]+$/.test(s);
}
function nextQ(_depth){
  _depth=_depth||0;
  // 英和・和英ミックス出題: 和英は簡単めなので約7問に1問（残りは英和）
  qMode=Math.random()<(1/7)?'je':'ej';
  const data=qMode==='ej'?EJ:JE;
  if(data.length<5) return;
  const st=qs[qMode];
  const rq=retryQueue[qMode];
  // 出題対象プール（人名エントリを除外）
  const pool=data.filter(e=>!isPersonName(e));
  let correct;
  // 不正解キューに「今が出題タイミング」のものがあれば優先
  const due=rq.findIndex(r=>st.n>=r.retryAt);
  if(due>=0){
    correct=rq[due].entry;
    rq.splice(due,1);
  } else {
    // 全語制覇型: まだ出していない語を優先して出題（一巡でリセット）
    correct=(qMode==='ej'?EJ:JE)[coverDrawIdx(qMode)];
  }
  // 正解ラベルが空、またはEJモードで日本語を含まない場合は別の問題を引き直す
  const _cLbl=choiceLbl(correct);
  if(!_cLbl.trim()){ if(_depth<25) nextQ(_depth+1); return; }
  // 「記号」等、略語表記の残骸だけになった問題は出題しない（別の問題を引き直す）
  if(/^(記号|符号|略号|略称|略語)$/.test(_cLbl.trim())){ if(_depth<25) nextQ(_depth+1); return; }
  if(qMode==='ej'&&!/[ぁ-んァ-ヶー一-龯]/.test(_cLbl)){ if(_depth<25) nextQ(_depth+1); return; }
  // 選択肢として表示できるか判定
  const _jaRe=/[ぁ-んァ-ヶー一-龯]/;
  function hasGoodLbl(e){
    if(isPersonName(e)) return false;
    if(e.cat==='abbr') return !!(e.def&&e.def.length>=4);
    // 実際に表示するラベルで判定（cf.のみ・括弧残骸などを除外）
    const lbl=choiceLbl(e);
    if(!lbl || lbl.trim().length<2) return false;
    if(/^(記号|符号|略号|略称|略語)$/.test(lbl.trim())) return false;
    if(qMode==='ej' && !_jaRe.test(lbl)) return false;
    return true;
  }
  const used=new Set([correct]);const wrongs=[];let t=0;
  const cLblC=choiceLbl(correct);const cLen=cLblC.length;const cFirst=cLblC[0]||'';
  const cEnFirst=(correct.en||'').split(/[,\/]/)[0].trim().toLowerCase();
  while(wrongs.length<4&&t++<600){
    const c=pool[Math.floor(Math.random()*pool.length)];
    if(!used.has(c)&&hasGoodLbl(c)){
      const lbl=choiceLbl(c);
      if(lbl.trim().length<2) continue;
      // JEモード: 正解と英語先頭語が同じ項目は選択肢に出さない（紛らわしすぎる）
      if(qMode==='je'&&cEnFirst.length>=2){const eF=(c.en||'').split(/[,\/]/)[0].trim().toLowerCase();if(eF===cEnFirst) continue;}
      // 最初の350回は「似た選択肢」を優先（長さが近い or 頭文字一致）
      if(t<350){
        const lenOk=cLen>0&&Math.min(cLen,lbl.length)/Math.max(cLen,lbl.length)>0.45;
        const firstOk=cFirst&&lbl[0]===cFirst;
        if(!lenOk&&!firstOk&&Math.random()<0.75) continue;
      }
      used.add(c);wrongs.push(c);
    }
  }
  cq={correct,choices:[...wrongs,correct].sort(()=>Math.random()-.5),hit:difficultyOf(correct)};
  ans=false;
  quizAsked++;   // 出題数（進むほど制限時間を短縮）
  const qdc=$('quiz-def-card');if(qdc)qdc.style.display='none';
  renderQ();
}
function renderQ(){
  if(!cq) return;
  const{correct,choices}=cq;
  setChar('quiz-char-img',qMode==='ej'?'bat':'pit');
  const modeLbl=qMode==='ej'?'英語 → 日本語':'日本語 → 英語';
  const hit=cq.hit||1;
  // 塁は打撃タイミング（往復回数）で決まるため塁打バッジは表示しない。
  // 難問（復路で三塁打がねらえる）だけ印を付ける。
  const badge=(hit>=3)?' <span class="hit-badge hit-3">🔥 難問（三塁打ねらい）</span>':'';
  const qq=$('qq-lbl');
  qq.innerHTML=esc(modeLbl)+badge;
  $('qq-term').textContent=qMode==='ej'?correct.en:correct.ja;
  $('next-btn').disabled=true; $('next-btn').textContent='次の問題 →';
  const grid=$('choices');grid.innerHTML='';
  choices.forEach(ch=>{
    const lbl=choiceLbl(ch);
    if(!lbl.trim()&&ch!==correct) return;
    const btn=document.createElement('button');
    btn.className='choice';btn.textContent=lbl;
    btn.onclick=()=>pick(btn,ch===correct,correct);
    grid.appendChild(btn);
  });
  startPitchClock();
}
// ── ピッチクロック（制限時間） ──
let quizTimer=null, quizAsked=0, quizGameCorrect=0;
// 制限時間は「正解数（実力）」連動。苦戦中は長め、正解を重ねるほど短縮
function pitchMs(){ return Math.max(8000, 16000 - Math.floor(quizGameCorrect/3)*1000); } // 16→…→8秒
function startPitchClock(){
  const fill=$('pitch-clock-fill'); if(!fill) return;
  clearTimeout(quizTimer); clearTimeout(fill._warn);
  const ms=pitchMs();
  fill.style.transition='none'; fill.style.width='100%'; fill.classList.remove('warn');
  void fill.offsetWidth;
  fill.style.transition='width '+ms+'ms linear';
  fill.style.width='0%';
  fill._warn=setTimeout(()=>fill.classList.add('warn'),Math.max(0,ms-3000));
  quizTimer=setTimeout(onTimeUp,ms);
}
function stopPitchClock(){
  clearTimeout(quizTimer);quizTimer=null;
  const fill=$('pitch-clock-fill');
  if(fill){ clearTimeout(fill._warn); const w=getComputedStyle(fill).width; fill.style.transition='none'; fill.style.width=w; }
}
function onTimeUp(){
  if(ans||!cq) return;
  // 時間切れ＝見逃し三振（1アウト）。不正解として処理
  pick(null,false,cq.correct);
}
function choiceLbl(entry){
  // 略語エントリはdef（日本語説明）を選択肢に使う
  if(entry.cat==='abbr'){
    const d=(entry.def||'').replace(/\s*cf\..*$/i,'').replace(/[。、\s]+$/,'').trim().substring(0,80);
    if(d) return d;
    // defがない場合はjaから抽出（"= WORD の略。" を除去して最初の文）
    return (entry.ja||'').replace(/^=\s*\S+\s+の略[。]?\s*/,'').split(/[。\n]/)[0].trim().substring(0,40);
  }
  // 拡充済みエントリは short（簡潔な答え）を選択肢に使う
  if(entry.short && entry.short.trim()){
    let s=entry.short.trim();
    // 「(… は和製/和用法/和略 …)」等の注記カッコを除去
    s=s.replace(/\s*[（(][^）)]*(和製|和用法|和略|和)[^）)]*[）)]/g,'').trim();
    // 英→日では、英字を含む注記カッコ（＝英語併記）を除去して答えバレを防ぐ
    if(qMode==='ej') s=s.replace(/\s*[（(][^）)]*[A-Za-z][^）)]*[）)]/g,'').trim();
    // 末尾の句点「。」を除去
    s=s.replace(/[。、\s]+$/,'').trim();
    if(s.length>60) s=s.substring(0,60)+'…';
    if(s) return s;
  }
  const raw=(qMode==='ej'?entry.ja:entry.en)||'';
  let s=raw;
  // 先頭の「= word1 word2 の略（記号/称/語）。」形式（複数語略語）を除去してから残りを処理
  // ※「の略記号」を「の略」だけ消すと「記号」が誤って答えになるため記号/称/語まで消す
  s=s.replace(/^=\s*(?:[A-Za-z0-9\-\/]+\s+)+の略(?:記号|符号|号|称|語)?[。]?\s*/,'');
  // 先頭の「= word 」形式（同義語・一語）を除去
  s=s.replace(/^=\s*\S+\s+/,'');
  // 先頭の「（追記）」を除去
  s=s.replace(/^（追記）/,'');
  // 先頭の番号「1. 」「１．」などを除去
  s=s.replace(/^[0-9０-９]+[.．。、]\s*/,'');
  // 先頭が（例文訳）の長い括弧書きなら除去（例文翻訳が選択肢になるのを防ぐ）
  s=s.replace(/^（[^）]{8,}）\s*/,'');
  // 最初の文（句点まで）を取り出す
  const m=s.match(/^([^。\n]+。?)/);
  if(m) s=m[1];
  // 英語の引用符つき文章を除去 "..."
  s=s.replace(/[“”"][^“”"]*[“”"]/g,'').trim();
  // cf. / e.g. 以降を除去
  s=s.replace(/\s*cf\..*$/i,'').replace(/\s*e\.g\..*$/i,'');
  // ；：以降除去
  s=s.replace(/\s*[；：].*$/,'');
  // 先頭・末尾に残った鉤括弧・引用符の残骸を除去
  s=s.replace(/^[「｢『"”'\s]+/,'').replace(/[「｢『」｣』"”'\s]+$/,'').trim();
  // cf./e.g. だけが残った場合は無効化
  if(/^(cf\.|e\.g\.)/i.test(s)) s='';
  // 末尾の句点・読点・空白を除去（読点での途中カットはしない＝意味が途切れるため）
  s=s.replace(/[。、\s]+$/,'').trim();
  // 英→日では、英字を含む注記カッコ（＝英語併記）を除去して答えバレを防ぐ
  if(qMode==='ej'){ s=s.replace(/\s*[（(][^）)]*[A-Za-z][^）)]*[）)]/g,'').replace(/[。、\s]+$/,'').trim(); }
  // 60文字を超えたら省略（読点区切りの完全な文を保持）
  if(s.length>60) s=s.substring(0,60)+'…';
  // EJモードで日本語が含まれない場合はraw中の日本語部分を抽出して使用
  if(qMode==='ej' && s && !/[ぁ-んァ-ヶー一-龯]/.test(s)){
    const jaM=raw.match(/[ぁ-んァ-ヶー一-龯][^。\n]*/);
    if(jaM) s=jaM[0].replace(/\s*cf\..*$/i,'').replace(/[。\s]+$/,'').trim().substring(0,60);
  }
  if(!s){
    // 生テキストから cf./e.g./括弧を除いた最初の断片（無ければ空のまま）
    s=raw.replace(/\s*cf\..*$/i,'').replace(/\s*e\.g\..*$/i,'')
         .replace(/[「｢」｣『』]/g,'').replace(/[。\n].*/,'').trim().substring(0,60);
  }
  return s;
}
// ── キャラクター表情切替 ──
const CHAR_IMG={def:'images/デフォ.png',joy:'images/よろこび.png',bat:'images/バッター.png',pit:'images/ピッチャー.png',runR:'images/右向き.png',runL:'images/左向き.png',sad:'images/うなだれ.png'};
function setChar(imgId,key,pop){
  const img=$(imgId);
  if(!img) return;
  img.src=CHAR_IMG[key]||CHAR_IMG.def;
  // ダイヤモンド上のホームキャラはtranslateで中央寄せしているため専用popを使う
  const cls=img.classList.contains('dchar')?'char-pop-d':'char-pop';
  img.classList.remove('char-pop','char-pop-d');
  if(pop){void img.offsetWidth;img.classList.add(cls);}
}
// ホームイン時に飛び出すポーズをランダム選択
//  ハート/シー=各30回に1回、ガッツ/王冠/でんわ=各5回に1回、残りは よろこび
function homeRunPose(){
  const r=Math.random();
  if(r<1/30) return 'images/ハート.png';
  if(r<2/30) return 'images/シー.png';
  if(r<2/30+1/5) return 'images/ガッツ.png';
  if(r<2/30+2/5) return 'images/王冠.png';
  if(r<2/30+3/5) return 'images/でんわ.png';
  return 'images/よろこび.png';
}
// 今日の10問ボーナス: 飛び出し画像＋大きな「＋N点」
function dailyBonusFlash(points){
  const el=$('quiz-flash');if(!el) return;
  const ci=$('quiz-flash-char'); if(ci) ci.src=homeRunPose();
  const bn=$('quiz-flash-bonus');
  if(bn){
    bn.textContent='＋'+points+'点';
    if(points>=5){ bn.style.background='#ffd23f'; bn.style.color='#1e1e2e'; } // 全問正解=金プレート＋濃紺文字
    else { bn.style.background='#ff5a4e'; bn.style.color='#fff'; }            // 6〜9問=コーラルプレート＋白文字
  }
  const mk=$('quiz-flash-mark'); if(mk) mk.textContent='';
  el.className='show ok char bonus';
  clearTimeout(el._t1);clearTimeout(el._t2);
  el._t1=setTimeout(()=>el.classList.replace('show','hide'),2300);
  el._t2=setTimeout(()=>{el.className='';if(bn)bn.textContent='';},2600);
  vibrate([120,60,220]);
}
function showFlash(ok,withChar){
  const el=$('quiz-flash');const mk=$('quiz-flash-mark');
  let special=false;
  if(ok&&withChar){
    const pose=homeRunPose();
    const ci=$('quiz-flash-char'); if(ci) ci.src=pose;
    special=pose.includes('ハート')||pose.includes('シー'); // レア＝2回飛び出し
  }
  el.className='show '+(ok?'ok':'ng')+(ok&&withChar?' char':'')+(special?' twice':'');
  mk.textContent=ok?'○':'✕';
  clearTimeout(el._t1);clearTimeout(el._t2);
  const hold=special?2600:(ok&&withChar?1600:800);
  el._t1=setTimeout(()=>el.classList.replace('show','hide'),hold);
  el._t2=setTimeout(()=>{el.className='';mk.textContent='';},hold+300);
}
function pick(btn,ok,correct){
  if(ans) return;ans=true;
  // 音は最初に鳴らす（DOM更新より先＝タップと同時に聞こえる）
  if(ok){ if(swingAuto) sfxForVal(1); else sfxSelectOK(); } else { sfxWrong(); }
  stopPitchClock();
  const st=qs[qMode];st.n++;if(ok) st.ok++;
  updateScores();
  setChar('quiz-char-img',ok?'joy':'def',ok);
  let doSwingPhase=false;
  if(ok){
    quizStreak++;
    quizGameCorrect++;   // この試合の正解数（＝球のレベル・制限時間短縮の基準）
    const hit=(cq&&cq.hit)?cq.hit:1;   // 難易度＝甘い球（狙える最大の当たり）
    if(swingAuto){
      // オート: 常に単打（1塁）。長打・本塁打は手動のみ＝オートが有利にならない
      const _p=$('quiz-play'); if(_p){ try{_p.scrollTo({top:0,behavior:'smooth'});}catch(e){_p.scrollTop=0;} }
      applyHit(1);
    } else {
      // 手動: タイミングよく打つ挑戦フェーズへ
      doSwingPhase=true;
    }
  } else {
    quizStreak=0;
    quizOuts++;
    updateOuts();
    showFlash(false);
    // 走者は塁に残す（クリアしない）
  }
  if(btn) btn.classList.add(ok?'ok':'ng');
  if(!ok){
    const cLbl=choiceLbl(correct);
    document.querySelectorAll('.choice').forEach(b=>{
      if(b.textContent===cLbl) b.classList.add('ok');
    });
    retryQueue[qMode].push({entry:correct, retryAt:st.n+20});
    // 復習リストに追加
    const arr=qMode==='ej'?EJ:JE;
    const ci=arr.indexOf(correct);
    if(ci>=0) addToReview(qMode,ci);
  }
  document.querySelectorAll('.choice').forEach(b=>b.disabled=true);
  // 辞書カード表示
  const card=$('quiz-def-card');
  if(card&&correct){
    const isEJ=qMode==='ej';
    const head=isEJ?correct.en:correct.ja;
    const enriched=(correct.short||correct.long);
    if(enriched){
      // 拡充済みは画像パターンで表示（cf.等の生テキストや結合ミスを出さない）
      card.innerHTML=`<div class="ddc-en">${isEJ?'英和':'和英'}</div><div class="ddc-head">${esc(head)}</div>`+
        `<div class="ddc-def">${enrichHtml(correct,qMode)}</div>`;
      card.style.display='';
      bindRelJumps(card);
    } else {
      const sub=isEJ&&correct.cat==='abbr'&&correct.ja?correct.ja:'';
      const defText=isEJ
        ?(correct.cat==='abbr'?correct.def:correct.ja)
        :(correct.en+'\n\n'+(correct.def||'')).trim();
      card.innerHTML=
        `<div class="ddc-en">${isEJ?'英和':'和英'}</div>`+
        `<div class="ddc-head">${esc(head)}</div>`+
        (sub?`<div class="ddc-sub">${defToHtml(sub,isEJ?'ej':'je')}</div>`:'')+
        (defText?`<div class="ddc-def">${defToHtml(defText,isEJ?'ej':'je')}</div>`:'');
      card.style.display='';
    }
  }
  if(doSwingPhase){
    // 打撃フェーズ中は「次の問題」を止め、結果確定後に finishTurn
    $('next-btn').disabled=true;
    const hit=(cq&&cq.hit)?cq.hit:1;
    startSwing(hit);
  } else {
    finishTurn();
  }
}
// ヒット結果を反映（value塁進む。0=アウト）
function applyHit(value){
  if(value<=0){
    quizStreak=0; quizOuts++; updateOuts(); showOutFx();
    setChar('quiz-char-img','sad'); // アウトはうなだれ
    return;
  }
  const r=advanceRunners(value);
  addRuns(r.runs);
  // ホームラン(value>=4)は sfxHomerun 側で歓声を鳴らすので、ここでは非HRの得点時のみ歓声
  if(r.runs>0){ showFlash(true,true); vibrate([120,60,220]); if(value<4) setTimeout(sfxCheer,120); }
  else vibrate(35);
}
// ターン終了処理（次へ有効化・アウト判定）
function outLimit(){ return versusActive?1:3; }
function finishTurn(){
  $('next-btn').disabled=false;
  if(quizOuts>=outLimit()){
    if(versusActive){
      // 対戦：アウトでも解説を見られるよう、ボタンを押してからスコア表示
      const nb=$('next-btn'); if(nb){ nb.disabled=false; nb.textContent='スコアを見る →'; }
    }else{
      $('next-btn').disabled=true;
      setTimeout(threeOut,900);
    }
  }
}
// 「次の問題」ボタン：対戦でアウト後はスコアボードへ、それ以外は次問へ
function quizNext(){
  if(versusActive && quizOuts>=outLimit()){ versusUserHalfDone(); return; }
  nextQ();
}
function showOutFx(){
  vibrate(200);
  const el=$('quiz-outs');
  if(el){el.classList.remove('run-bump');void el.offsetWidth;el.classList.add('run-bump');}
}
// ── CPU対戦モード（中間案C：短縮9回・1アウト制） ─────────────────────────────
let versusActive=false, vs=null;
const VS_WINRATE={1:0.75,2:0.68,3:0.60,4:0.52,5:0.45};
function oppLevelForUser(){
  const runs=(typeof runsTotal==='number')?runsTotal:0;
  const {r}=currentRank(runs);
  const idx=RANKS.findIndex(x=>x.name===r.name);
  return Math.max(1,Math.min(5,idx+1));
}
function genOppName(){
  const my=(getProfile()&&getProfile().nick)||'';
  for(let i=0;i<25;i++){ const c=genNickname(''); if(c&&c!==my) return c; }
  return 'ライバル';
}
function _versusToday(){ try{ return JSON.parse(localStorage.getItem('mlb_versus_today'))||{}; }catch(e){ return {}; } }
function startVersus(){
  // 対戦は1日1回無料。2回目以降は広告を挟む（プレミアムは広告なし）
  const t=todayStr(), c=_versusToday();
  const n=(c.date===t)?(c.n||0):0;
  if(n>=1 && !isPremium()){
    showAd(()=>_beginVersus());
    return;
  }
  _beginVersus();
}
function _beginVersus(){
  const t=todayStr(), c=_versusToday();
  const n=((c.date===t)?(c.n||0):0)+1;
  localStorage.setItem('mlb_versus_today',JSON.stringify({date:t,n}));
  const lvl=oppLevelForUser(), wr=VS_WINRATE[lvl]||0.6;
  vs={ name:genOppName(), level:lvl, winRate:wr, userHome:Math.random()<0.5,
       inning:1, userLine:[], cpuLine:[], userScore:0, cpuScore:0, inningRuns:0,
       maxInning:3, finished:false, userWins:(Math.random()<wr) };
  versusActive=true; lastQuizSection='play';
  quizGameCorrect=0;
  track('versus_start',{level:vs.level});
  $('quiz-menu').style.display='none';
  $('quiz-play').style.display='';
  $('versus-bar').style.display='';
  const _hr=$('hud-runs'); if(_hr) _hr.style.display='none'; // 試合中は本日/通算を隠す（試合の得点だけに集中）
  updateVersusBar();
  versusToast('🆚 対戦相手\n'+vs.name+'（Lv'+vs.level+'）\n'+(vs.userHome?'あなたは後攻':'あなたは先攻'), ()=>versusStartInning());
}
function updateVersusBar(){
  const b=$('versus-bar'); if(!b||!vs) return;
  const myNick=(getProfile()&&getProfile().nick)||'あなた';
  b.innerHTML='🆚 <b><span style="color:#1d6fd6">'+esc(myNick)+'</span> '+(vs.userScore)+' - '+(vs.cpuScore)+' '+esc(vs.name)+'（Lv'+vs.level+'）</b>　'+vs.inning+'回';
}
function versusStartInning(){
  if(!versusActive) return;
  // 広告は4回に入る時と8回に入る時の2回だけ（プレミアムはスキップ）
  if(vs.inning===4 || vs.inning===8){ showAd(()=>_versusUserHalf()); }
  else _versusUserHalf();
}
function _versusUserHalf(){
  quizOuts=0; quizStreak=0; vs.inningRuns=0;
  updateOuts(); clearBases(); updateVersusBar();
  const bn=$('versus-inning'); if(bn){ bn.textContent=vs.inning+'回：あなたの攻撃（1アウトで交代）'; bn.style.display=''; setTimeout(()=>{if(bn)bn.style.display='none';},1400); }
  cq=null; ans=false; nextQ();
}
// 相手（CPU）の1イニング得点をシミュレート
function simCpuInning(isFinal){
  if(isFinal){
    // 最終回はユーザー得点が確定済み → 勝敗を勝率に合わせて確定（逆転サヨナラ演出込み）
    if(vs.userScore<=0) return Math.max(1, vs.userScore-vs.cpuScore+1); // 0点なら勝てない
    if(vs.userWins){
      const need=vs.userScore-vs.cpuScore; // 追いつくのに必要
      if(need<=0) return 0;
      return Math.max(0,Math.min(need-1, Math.floor(Math.random()*need))); // 届かない
    } else {
      return Math.max(1, vs.userScore-vs.cpuScore+1); // 1点上回って逆転
    }
  }
  // 通常回：接戦になるよう調整
  const diff=vs.userScore-vs.cpuScore;
  let p = diff>1?0.85 : (diff<-1?0.2 : 0.5);
  let runs=0; if(Math.random()<p) runs=1+((Math.random()<0.28)?1:0);
  if(vs.userWins && vs.cpuScore+runs>=vs.userScore) runs=Math.max(0,vs.userScore-1-vs.cpuScore); // 勝ち予定なら相手を先に行かせない
  return runs;
}
function versusUserHalfDone(){
  const i=vs.inning-1;
  vs.userLine[i]=vs.inningRuns; vs.userScore=vs.userLine.reduce((a,b)=>a+(b||0),0);
  const isFinal=(vs.maxInning===9 && vs.inning===9);
  const cpu=simCpuInning(isFinal);
  vs.cpuLine[i]=cpu; vs.cpuScore=vs.cpuLine.reduce((a,b)=>a+(b||0),0);
  updateVersusBar();
  renderVersusBoard();
  $('versus-board').classList.add('show');
}
function renderVersusBoard(){
  const wrap=$('versus-board-inner'); if(!wrap) return;
  const cells=n=>{ let s=''; for(let k=0;k<9;k++){ const v=(n[k]==null)?'':n[k]; s+='<td>'+v+'</td>'; } return s; };
  const myName=(getProfile()&&getProfile().nick)||'あなた';
  // 表の行ラベルは「先攻/後攻」で短く（横幅節約）。名前は表の上に別行で表示（あなたは青）
  const meBlue='<span style="color:#1d6fd6">'+esc(myName)+'</span>';
  const namesLine=vs.userHome
    ? '<div class="vs-names">先攻：'+esc(vs.name)+'　後攻：'+meBlue+'</div>'
    : '<div class="vs-names">先攻：'+meBlue+'　後攻：'+esc(vs.name)+'</div>';
  const userRow='<tr><th style="color:#1d6fd6">'+(vs.userHome?'後攻':'先攻')+'</th>'+cells(vs.userLine)+'<td class="vs-r">'+vs.userScore+'</td></tr>';
  const cpuRow='<tr><th>'+(vs.userHome?'先攻':'後攻')+'</th>'+cells(vs.cpuLine)+'<td class="vs-r">'+vs.cpuScore+'</td></tr>';
  const rows=vs.userHome ? (cpuRow+userRow) : (userRow+cpuRow); // 先攻が上
  const msg=versusInningMsg();
  const last=(vs.maxInning===9 && vs.inning===9);
  const checkpoint=(vs.maxInning===3 && vs.inning===3);
  let btns;
  if(last){ btns='<button class="vs-btn" onclick="showVersusResult(true)">結果を見る ▶</button>'; }
  else if(checkpoint){
    btns='<button class="vs-btn" onclick="versusExtend()">9回まで続ける（広告）▶</button>'+
         '<button class="vs-btn vs-btn-sub" onclick="showVersusResult(false)">ここで終える（得点は加算）</button>';
  } else {
    btns='<button class="vs-btn" onclick="versusNextInning()">次の回 ▶</button>';
  }
  wrap.innerHTML=
    '<div class="vs-board-ttl">'+vs.inning+'回 終了</div>'+
    namesLine+
    '<table class="vs-score"><tr class="vs-hdr"><th></th>'+[1,2,3,4,5,6,7,8,9].map(n=>'<th>'+n+'</th>').join('')+'<th class="vs-r">計</th></tr>'+
    rows+'</table>'+
    (msg?'<div class="vs-msg">'+msg+'</div>':'')+
    '<div class="vs-btns">'+btns+'</div>';
}
function versusInningMsg(){
  const i=vs.inning-1;
  const cpu=vs.cpuLine[i]||0, usr=vs.userLine[i]||0;
  let m='この回：あなた'+usr+'点／相手'+cpu+'点';
  if(vs.maxInning===9 && vs.inning===9){
    // 逆転・サヨナラ判定
    const beforeU=vs.userScore-usr, beforeC=vs.cpuScore-cpu;
    const flipped=(Math.sign(beforeU-beforeC)!==Math.sign(vs.userScore-vs.cpuScore))&&(vs.userScore!==vs.cpuScore);
    const winner=vs.userScore>vs.cpuScore?'user':'cpu';
    const homeIsWinner=(vs.userHome&&winner==='user')||(!vs.userHome&&winner==='cpu');
    if(flipped&&homeIsWinner) m+= winner==='user'?'　🎉 逆転サヨナラ勝ち！':'　💥 逆転サヨナラ負け…';
    else if(flipped) m+= winner==='user'?'　🎉 逆転勝ち！':'　💥 逆転負け…';
  }
  return m;
}
function versusNextInning(){
  $('versus-board').classList.remove('show');
  vs.inning++;
  versusStartInning();
}
function versusExtend(){
  vs.maxInning=9;
  $('versus-board').classList.remove('show');
  vs.inning++;
  versusStartInning();
}
// fullGame=true: 9回完走（勝てば+10ボーナス）／false: 3回で終える（得点は加算・ボーナスなし）
function showVersusResult(fullGame){
  vs.finished=true;
  const win=vs.userScore>vs.cpuScore, tie=vs.userScore===vs.cpuScore;
  // 得点はどちらでも加算（3回で終えても有効）。ボーナス+10は9回完走・勝利のみ
  versusActive=false;
  track('versus_result',{full_game:!!fullGame, win:win, user_score:vs.userScore, cpu_score:vs.cpuScore, level:vs.level});
  if(vs.userScore>0) addRuns(vs.userScore);
  if(fullGame && win) addRuns(10);
  // 対戦成績は9回完走のみ記録
  if(fullGame){ try{ let rec=JSON.parse(localStorage.getItem('mlb_vs_rec'))||{w:0,l:0}; if(win)rec.w++; else rec.l++; localStorage.setItem('mlb_vs_rec',JSON.stringify(rec)); }catch(e){} }
  const ov=$('versus-board'); if(ov) ov.classList.remove('show');
  const r=$('versus-result');
  if(r){
    r.querySelector('.vs-res-ttl').textContent = fullGame ? (win?'勝利！ 🎉':(tie?'引き分け':'敗戦…')) : (vs.inning+'回で終了');
    r.querySelector('.vs-res-score').textContent='あなた '+vs.userScore+' - '+vs.cpuScore+' 相手';
    r.querySelector('.vs-res-char').src = (win||tie)?CHAR_IMG.joy:CHAR_IMG.sad;
    r.querySelector('.vs-res-bonus').textContent = (fullGame&&win)?'累計ボーナス ＋10点！':(vs.userScore>0?'得点 ＋'+vs.userScore+'点を加算':'');
    r.classList.add('show');
  }
}
function closeVersusResult(){
  const r=$('versus-result'); if(r) r.classList.remove('show');
  vs=null; $('versus-bar').style.display='none';
  quizToMenu();
}
function versusQuitConfirm(){
  if(confirm('対戦を途中でやめると、この試合の得点は無効（累計・本日ともに加算なし）になります。やめますか？')){
    versusActive=false; vs=null;
    $('versus-board').classList.remove('show');
    $('versus-bar').style.display='none';
    updateRuns(); // 暫定表示を確定値に戻す（没収）
    quizToMenu();
  }
}
function versusToast(text,cb){
  const t=$('versus-toast'); if(!t){ if(cb)cb(); return; }
  t.textContent=text; t.classList.add('show');
  setTimeout(()=>{ t.classList.remove('show'); if(cb)cb(); },1800);
}
// ── スイングタイミング（打撃）ミニゲーム ──
// 連続正解数に応じて球のレベルが上がる（0-4:やさしい, 5-9, 10-19, 20-39, 40+:最難関）
// cfg: spd=マーカー速度(%/frame), perfect/great/good=中央からの半幅(%), brk=変化球度(0-1)
// 赤の滞在時間をなだらかに: 310→260→215→175→140→110→80→50ms（変化球はLv2から）
const SWING_TIERS=[
  {spd:0.86, perfect:8,  great:16, good:27, brk:0,   label:'Lv1 スローボール'},
  {spd:0.86, perfect:7.5,great:15, good:25, brk:0.12,label:'Lv2 変化球'},
  {spd:0.94, perfect:7,  great:13, good:22, brk:0.15,label:'Lv3 変化球'},
  {spd:1.03, perfect:6.5,great:12, good:20, brk:0.2, label:'Lv4 速い変化球'},
  {spd:1.14, perfect:6,  great:11, good:18, brk:0.25,label:'Lv5 快速＋変化'},
  {spd:1.28, perfect:5.5,great:10, good:16, brk:0.3, label:'Lv6 剛速＋変化'},
  {spd:1.49, perfect:5,  great:9,  good:15, brk:0.4, label:'Lv7 鋭い変化球'},
  {spd:2.0,  perfect:4.5,great:9,  good:14, brk:0.5, label:'Lv8 魔球（最難関）'},
];
function pitchTier(){
  const c=quizGameCorrect;
  if(c<=3) return 0; if(c<=7) return 1; if(c<=12) return 2; if(c<=18) return 3;
  if(c<=25) return 4; if(c<=33) return 5; if(c<=42) return 6; return 7;
}
let sw={active:false,pos:0,dir:1,raf:0,cfg:null,hit:1,foulUsed:false,pW:0,gW:0,dW:0,pass:1,center:50};
function startSwing(hit){
  sw.hit=hit; sw.foulUsed=false;
  beginPitch();
  // スイングバーが画面に見えるようスクロール（小さい画面対策）
  const ui=$('swing-ui'); if(ui){ try{ ui.scrollIntoView({behavior:'smooth',block:'center'}); }catch(e){ const p=$('quiz-play'); if(p) p.scrollTop=0; } }
}
function beginPitch(){
  const cfg=SWING_TIERS[pitchTier()];
  sw.cfg=cfg; sw.pos=0; sw.dir=1; sw.active=true; sw.pass=1;
  // 判定幅は球のレベルのみで決まる（塁は打つタイミング＝往復回数で決まる）
  sw.pW=cfg.perfect; sw.gW=cfg.great; sw.dW=cfg.good;
  // 3問目（連続正解3）から赤（判定中心）をランダムに移動
  if(quizGameCorrect>=3){
    const lo=sw.dW+4, hi=100-sw.dW-4;
    sw.center=lo+Math.random()*(hi-lo);
  }else{
    sw.center=50;
  }
  setZone('sz-perfect',sw.pW); setZone('sz-great',sw.gW); setZone('sz-good',sw.dW);
  const pl=$('swing-pitch'); if(pl) pl.textContent=cfg.label;
  const gd=$('swing-guide'); if(gd) gd.textContent='赤 → 往路:本塁打 / 復路:'+(sw.hit>=3?'三塁打':'二塁打')+' / 以降:単打';
  const rs=$('swing-result'); if(rs){rs.textContent=''; rs.className='swing-result';}
  const bt=$('swing-btn'); if(bt) bt.disabled=false;
  const ui=$('swing-ui'); if(ui) ui.style.display='';
  loopSwing();
}
function setZone(id,half){
  const e=$(id); if(!e) return;
  e.style.left=(sw.center-half)+'%'; e.style.width=(2*half)+'%';
}
function loopSwing(){
  if(!sw.active) return;
  let step=sw.cfg.spd;
  if(sw.cfg.brk) step=sw.cfg.spd*(1+sw.cfg.brk*Math.sin(sw.pos/100*Math.PI)); // 変化球＝中央付近で加速
  sw.pos+=sw.dir*step;
  // 端で折り返すたびに往復回数を+1（1=往路, 2=復路, 3以降=それ以降の往復）
  if(sw.pos>=100){sw.pos=100;sw.dir=-1;sw.pass++;} else if(sw.pos<=0){sw.pos=0;sw.dir=1;sw.pass++;}
  const m=$('swing-marker'); if(m) m.style.left=sw.pos+'%';
  sw.raf=requestAnimationFrame(loopSwing);
}
function doSwing(){
  if(!sw.active) return;
  sw.active=false; cancelAnimationFrame(sw.raf);
  const bt=$('swing-btn'); if(bt) bt.disabled=true;
  const err=Math.abs(sw.pos-sw.center);
  const hard=sw.hit>=3;
  let res;
  if(err<=sw.pW){
    // 赤ゾーン＝往復回数で塁が決まる（長打系）
    if(sw.pass<=1)       res={val:4, label:'ジャスト！ホームラン！'};
    else if(sw.pass===2) res={val:hard?3:2, label:hard?'ジャスト！三塁打！':'ジャスト！二塁打！'};
    else                 res={val:1, label:'ジャスト！ヒット！'};
  }else if(err<=sw.gW){
    // 黄ゾーン＝単打
    res={val:1, label:'ナイスミート！ヒット！'};
  }else {
    // 青ゾーン＝詰まってアウト、白ゾーン＝空振りアウト
    const inBlue=err<=sw.dW;
    // ファウル打ち直し判定（10問目まで・30%・1打席1回。青＝詰まりの当たりのみ猶予）
    const canFoul=(quizGameCorrect<=10)&&(!sw.foulUsed)&&inBlue&&(Math.random()<0.30);
    if(canFoul){
      sw.foulUsed=true;
      showSwingResult('ファウル！打ち直し',null);
      sfxPock(); vibrate(60);
      setTimeout(beginPitch,950);
      return;
    }
    res={val:0, label:(inBlue?'詰まった当たり…アウト':'空振り…アウト'), out:true};
  }
  sfxForVal(res.val); // カキーン／コツン／空振り
  showSwingResult(res.label, !res.out);
  const wait=res.out?950:750;
  setTimeout(()=>{
    const ui=$('swing-ui'); if(ui) ui.style.display='none';
    applyHit(res.val);   // val<=0 は applyHit 内でアウト処理
    finishTurn();
  }, wait);
}
function showSwingResult(text,isHit){
  const rs=$('swing-result'); if(!rs) return;
  rs.textContent=text;
  rs.className='swing-result'+(isHit===true?' hit':(isHit===false?' out':''));
}
// 設定: スイング操作トグル
function toggleSwingAuto(){
  swingAuto=!swingAuto;
  localStorage.setItem('mlb_swing_auto',swingAuto?'1':'0');
  syncSwingToggle();
}
function syncSwingToggle(){
  const b=$('swing-toggle');
  if(b){ b.textContent=swingAuto?'オート':'手動'; b.classList.toggle('auto',swingAuto); }
  // 点の入り方の説明もモードで切り替え
  const man=$('howto-manual'), aut=$('howto-auto');
  if(man) man.style.display=swingAuto?'none':'';
  if(aut) aut.style.display=swingAuto?'':'none';
  const hint=$('howto-swing-hint');
  if(hint) hint.textContent=swingAuto?'▶ コツコツ学習したい人向け':'▶ 長打・本塁打を狙いたい人向け';
}

function updateScores(){
  $('sc-ej').textContent=qs.ej.ok+' / '+EJ_TOTAL;
  $('sc-je').textContent=qs.je.ok+' / '+JE_TOTAL;
  const ejRate=$('sc-ej-rate'),jeRate=$('sc-je-rate');
  if(ejRate) ejRate.textContent=qs.ej.n>0?Math.round(qs.ej.ok/qs.ej.n*100)+'%':'--%';
  if(jeRate) jeRate.textContent=qs.je.n>0?Math.round(qs.je.ok/qs.je.n*100)+'%':'--%';
  const sub=$('quiz-score-sub');
  if(sub) sub.textContent='英和: '+qs.ej.ok+'正解 / '+qs.ej.n+'問　和英: '+qs.je.ok+'正解 / '+qs.je.n+'問';
}
function askResetQuiz(){
  $('quiz-reset-confirm').classList.add('show');
  $('quiz-reset-btn').disabled=true;
}
function cancelResetQuiz(){
  $('quiz-reset-confirm').classList.remove('show');
  $('quiz-reset-btn').disabled=false;
}
function resetQuizScore(){
  qs.ej={ok:0,n:0};qs.je={ok:0,n:0};
  updateScores();
  $('quiz-reset-confirm').classList.remove('show');
  const btn=$('quiz-reset-btn');
  if(btn){btn.disabled=false;btn.textContent='リセット済';setTimeout(()=>{btn.textContent='リセット';},1500);}
}

// ── Daily Quiz ────────────────────────────────────────────────────────────────
let dailyState=null;
let dailyAns=false;
let dailyCurIdx=0;
let calYear=new Date().getFullYear();
let calMonth=new Date().getMonth(); // 0-based

function todayStr(){const d=new Date();return d.getFullYear()+'-'+String(d.getMonth()+1).padStart(2,'0')+'-'+String(d.getDate()).padStart(2,'0');}

function makeDailyQuestions(){
  const ejPool=EJ.filter(e=>!isPersonName2(e,'ej'));
  const jePool=JE.filter(e=>!isPersonName2(e,'je'));
  // 和英は2割（10問中2問）に固定、残り8問は英和。全語制覇型の共有袋から未出題を優先して選ぶ
  const usedEj=new Set(),usedJe=new Set();const ejSel=[],jeSel=[];let guard=0;
  while(ejSel.length<8&&guard++<800){const i=coverDrawIdx('ej');if(usedEj.has(i))continue;usedEj.add(i);if(dailyChoiceLbl(EJ[i],'ej').trim().length<2)continue;ejSel.push(i);}
  guard=0;
  while(jeSel.length<2&&guard++<800){const i=coverDrawIdx('je');if(usedJe.has(i))continue;usedJe.add(i);if(dailyChoiceLbl(JE[i],'je').trim().length<2)continue;jeSel.push(i);}
  const chosen=[...ejSel.map(i=>({type:'ej',idx:i})),...jeSel.map(i=>({type:'je',idx:i}))].sort(()=>Math.random()-.5);
  return chosen.map(q=>{
    // q.idx は EJ/JE 本体配列のインデックス（描画も arr[q.correct]/arr[ci] で本体を参照）
    const fullArr=q.type==='ej'?EJ:JE;
    const correct=fullArr[q.idx];
    const pool=q.type==='ej'?ejPool:jePool;
    const wrongs=[];const used=new Set([q.idx]);let t=0;
    const dLbl=dailyChoiceLbl(correct,q.type);const dLen=dLbl.length;const dFirst=dLbl[0]||'';
    const dEnFirst=(correct.en||'').split(/[,\/]/)[0].trim().toLowerCase();
    while(wrongs.length<3&&t++<600){
      const cand=pool[Math.floor(Math.random()*pool.length)];
      const orig=fullArr.indexOf(cand);
      if(orig>=0&&!used.has(orig)&&hasGoodLbl2(cand,q.type)){
        const lbl=dailyChoiceLbl(cand,q.type);
        if(lbl.trim().length<2) continue;
        // JEモード: 正解と英語先頭語が同じ項目は選択肢に出さない
        if(q.type==='je'&&dEnFirst.length>=2){const eF=(cand.en||'').split(/[,\/]/)[0].trim().toLowerCase();if(eF===dEnFirst) continue;}
        if(t<350){
          const lenOk=dLen>0&&Math.min(dLen,lbl.length)/Math.max(dLen,lbl.length)>0.45;
          const firstOk=dFirst&&lbl[0]===dFirst;
          if(!lenOk&&!firstOk&&Math.random()<0.75) continue;
        }
        used.add(orig);wrongs.push(orig);
      }
    }
    const choices=[...wrongs,q.idx].sort(()=>Math.random()-.5);
    return {type:q.type,idx:q.idx,choices,correct:q.idx};
  });
}
function isPersonName2(e,mode){
  if(e.cat==='abbr') return false;
  const raw=(mode==='ej'?e.ja:e.en)||'';
  const s=raw.replace(/[。参照\s]/g,'');
  return s.length>0&&/^[ァ-ヶーｦ-ﾟ・＝]+$/.test(s);
}
// ── 全語制覇型の出題（今日の10問・フリーバッティング・試合で共有）─────────────
// 出した語を「袋」から消化し、一巡（全語出しきる）したら自動リセット。ej/je別に保持。
const _COVER_KEY='mlb_cover_bags';
let _coverBags=null;
function _coverLoad(){ if(_coverBags) return _coverBags; try{_coverBags=JSON.parse(localStorage.getItem(_COVER_KEY))||{};}catch(e){_coverBags={};} return _coverBags; }
function _coverSave(){ try{localStorage.setItem(_COVER_KEY,JSON.stringify(_coverBags));}catch(e){} }
function _coverRefill(mode){
  _coverLoad();
  const arr=mode==='ej'?EJ:JE; const idxs=[];
  for(let i=0;i<arr.length;i++){ if(!isPersonName2(arr[i],mode)) idxs.push(i); }
  for(let i=idxs.length-1;i>0;i--){const j=Math.floor(Math.random()*(i+1));[idxs[i],idxs[j]]=[idxs[j],idxs[i]];}
  _coverBags[mode]=idxs; _coverBags[mode+'Len']=arr.length; _coverSave();
}
// 未出題を優先して1問引く（袋が空＝一巡完了なら詰め直す）。EJ/JE本体のインデックスを返す
function coverDrawIdx(mode){
  _coverLoad();
  const arr=mode==='ej'?EJ:JE;
  if(!_coverBags[mode]||!_coverBags[mode].length||_coverBags[mode+'Len']!==arr.length) _coverRefill(mode);
  const idx=_coverBags[mode].pop(); _coverSave();
  return (idx==null)?Math.floor(Math.random()*arr.length):idx;
}
const _jaRe2=/[ぁ-んァ-ヶー一-龯]/;
function hasGoodLbl2(e,mode){
  if(isPersonName2(e,mode)) return false;
  if(e.cat==='abbr') return !!(e.def&&e.def.length>=4);
  const raw=(mode==='ej'?e.ja:e.en)||'';
  const s=raw.replace(/^=\s*\S+\s+/,'').replace(/^（追記）/,'').replace(/^（[^）]{8,}）\s*/,'').trim();
  if(s.length<4) return false;
  if(mode==='ej' && !_jaRe2.test(s)) return false;
  return true;
}
function _hideDailyAll(){
  ['daily-start','daily-char','daily-prog-bar','daily-quiz-area','daily-result','daily-cal-area']
    .forEach(id=>{const el=$(id);if(el)el.style.display='none';});
}
function startDailyQuiz(){
  // 最初の未回答問題から開始
  dailyCurIdx=dailyState?dailyState.ans.findIndex(a=>a===null):0;
  if(dailyCurIdx<0) dailyCurIdx=0;
  _hideDailyAll();
  $('daily-char').style.display='';
  $('daily-prog-bar').style.display='';
  $('daily-quiz-area').style.display='';
  renderDailyQ();
}
function initDaily(){
  updateRuns(); // 得点表示を同期
  const today=todayStr();
  const raw=localStorage.getItem('mlb_daily_state');
  if(raw){
    try{dailyState=JSON.parse(raw);}catch(e){dailyState=null;}
  }
  if(!dailyState||dailyState.date!==today){
    // 新しい日 → スタート画面を挟まず即開始
    dailyState={date:today,qs:makeDailyQuestions(),ans:Array(10).fill(null)};
    localStorage.setItem('mlb_daily_state',JSON.stringify(dailyState));
    startDailyQuiz();
    return;
  }
  // find first unanswered
  dailyCurIdx=dailyState.ans.findIndex(a=>a===null);
  if(dailyCurIdx<0){
    _hideDailyAll();
    showDailyResult();
  } else {
    // 途中再開 → スタート画面を挟まず続きから
    _hideDailyAll();
    $('daily-char').style.display='';
    $('daily-prog-bar').style.display='';
    $('daily-quiz-area').style.display='';
    renderDailyQ();
  }
}
function renderDailyQ(){
  if(dailyCurIdx<0||dailyCurIdx>=10) return;
  const q=dailyState.qs[dailyCurIdx];
  const arr=q.type==='ej'?EJ:JE;
  const correct=arr[q.correct];
  // progress
  const done=dailyState.ans.filter(a=>a!==null).length;
  $('daily-prog-fill').style.width=(done/10*100)+'%';
  $('daily-prog-lbl').textContent=done+' / 10';
  setChar('daily-char-img',q.type==='ej'?'bat':'pit');
  $('daily-qq-lbl').textContent=q.type==='ej'?'英語 → 日本語':'日本語 → 英語';
  $('daily-qq-term').textContent=q.type==='ej'?correct.en:correct.ja;
  $('daily-next-btn').disabled=true;
  dailyAns=false;
  const grid=$('daily-choices');grid.innerHTML='';
  q.choices.forEach(ci=>{
    const entry=arr[ci];
    const lbl=dailyChoiceLbl(entry,q.type);
    if(!lbl.trim()&&ci!==q.correct) return;
    const btn=document.createElement('button');
    btn.className='choice';
    btn.textContent=lbl;
    btn.onclick=()=>dailyPick(btn,ci===q.correct,q,correct,ci);
    grid.appendChild(btn);
  });
}
function dailyChoiceLbl(entry,mode){
  if(entry.cat==='abbr'){
    const d=(entry.def||'').replace(/\s*cf\..*$/i,'').replace(/[。、\s]+$/,'').trim().substring(0,40);
    if(d) return d;
    return (entry.ja||'').replace(/^=\s*\S+\s+の略[。]?\s*/,'').split(/[。\n]/)[0].trim().substring(0,40);
  }
  const raw=(mode==='ej'?entry.ja:entry.en)||'';
  let s=raw;
  s=s.replace(/^=\s*(?:[A-Za-z0-9\-\/]+\s+)+の略[。]?\s*/,'').replace(/^=\s*\S+\s+/,'').replace(/^（追記）/,'').replace(/^[0-9０-９]+[.．。、]\s*/,'');
  s=s.replace(/^（[^）]{8,}）\s*/,'');
  const m=s.match(/^([^。\n]+。?)/);if(m) s=m[1];
  s=s.replace(/[“”„][^””„]*[“”„]/g,'').trim().replace(/\s*cf\..*$/i,'').replace(/\s*e\.g\..*$/i,'').replace(/\s*[；：].*$/,'').replace(/[。、\s]+$/,'').trim();
  if(s.length>60) s=s.substring(0,60)+'…';
  return s||raw.replace(/[。\n].*/,'').trim().substring(0,60)||raw.substring(0,60);
}
function dailyPick(btn,ok,q,correct,chosenIdx){
  if(dailyAns) return;
  dailyAns=true;
  // 効果音は最初に（タップと同時に聞こえるように）
  if(ok) sfxSelectOK(); else sfxWrong();
  dailyState.ans[dailyCurIdx]=ok;
  localStorage.setItem('mlb_daily_state',JSON.stringify(dailyState));
  // add to review if wrong
  if(!ok){
    addToReview(q.type,q.idx);
  }
  // 10問目の正解はクイズのホームインと同じ飛び出す演出
  const isLast=dailyCurIdx===9;
  if(ok&&isLast){ showFlash(true,true); vibrate([120,60,220]); }
  else { showFlash(ok); if(ok) vibrate(35); }
  setChar('daily-char-img',ok?'joy':'def',ok);
  btn.classList.add(ok?'ok':'ng');
  if(!ok){
    const cLbl=dailyChoiceLbl(correct,q.type);
    document.querySelectorAll('#daily-choices .choice').forEach(b=>{
      if(b.textContent===cLbl) b.classList.add('ok');
    });
  }
  document.querySelectorAll('#daily-choices .choice').forEach(b=>b.disabled=true);
  const nextBtn=$('daily-next-btn');
  nextBtn.disabled=false;
  nextBtn.textContent=dailyCurIdx===9?'今日の10問達成 ☆':'次の問題 →';
  // update progress
  const done=dailyState.ans.filter(a=>a!==null).length;
  $('daily-prog-fill').style.width=(done/10*100)+'%';
  $('daily-prog-lbl').textContent=done+' / 10';
  // 正解エントリの辞書カードを表示
  const arr2=q.type==='ej'?EJ:JE;
  const entry=arr2[q.idx];
  const card=$('daily-def-card');
  if(entry){
    const isEJ=q.type==='ej';
    const head=isEJ?entry.en:entry.ja;
    const enriched=(entry.short||entry.long);
    if(enriched){
      card.innerHTML=`<div class="ddc-en">${isEJ?'英和':'和英'}</div><div class="ddc-head">${esc(head)}</div>`+
        `<div class="ddc-def">${enrichHtml(entry,q.type)}</div>`;
      card.style.display='';
      bindRelJumps(card);
    } else {
      const sub=isEJ&&entry.cat==='abbr'&&entry.ja?entry.ja:'';
      const defText=isEJ
        ?(entry.cat==='abbr'?entry.def:entry.ja)
        :(entry.en+'\n\n'+(entry.def||'')).trim();
      card.innerHTML=
        `<div class="ddc-en">${isEJ?'英和':'和英'}</div>`+
        `<div class="ddc-head">${esc(head)}</div>`+
        (sub?`<div class="ddc-sub">${defToHtml(sub,isEJ?'ej':'je')}</div>`:'')+
        (defText?`<div class="ddc-def">${defToHtml(defText,isEJ?'ej':'je')}</div>`:'');
      card.style.display='';
    }
  }
}
function dailyNextQ(){
  $('daily-def-card').style.display='none';
  $('daily-next-btn').textContent='次の問題 →';
  dailyCurIdx++;
  if(dailyCurIdx>=10){
    showDailyResult();
  } else {
    renderDailyQ();
  }
}
function showDailyResult(){
  // 結果画面では下部ナビ「クイズ」で目次に戻れるようにする
  lastQuizSection='menu';
  _hideDailyAll();
  $('daily-result').style.display='';
  const correct=dailyState.ans.filter(a=>a===true).length;
  $('daily-result-score').textContent=correct+' / 10';
  track('daily_complete',{correct:correct}); // 今日の10問の正解数
  // mark done（1日1回）＋正解数で総得点ボーナス（10問=+4 / 6〜9問=+2）
  const bonus=correct===10?5:(correct>=6?3:0);
  const today=todayStr();
  let done=JSON.parse(localStorage.getItem('mlb_daily_done')||'[]');
  let bonusMsg='';
  let flashShown=false;
  if(!done.includes(today)){
    done.push(today);localStorage.setItem('mlb_daily_done',JSON.stringify(done));
    if(bonus>0){
      addRuns(bonus);
      bonusMsg=(correct===10?'全問正解！':'6問以上正解！')+' ＋'+bonus+'点';
      setTimeout(()=>dailyBonusFlash(bonus),400); // 飛び出し画像＋大きな「＋N点」
      flashShown=true;
    }
  } else if(bonus>0){
    bonusMsg=(correct===10?'全問正解！':'6問以上正解！')+' ＋'+bonus+'点 獲得済み';
  }
  const bn=$('daily-bonus');
  if(bn){ bn.textContent=bonusMsg; bn.style.display=bonusMsg?'':'none'; }
  updateRuns();
  // progress bar to 100%
  $('daily-prog-fill').style.width='100%';
  $('daily-prog-lbl').textContent='10 / 10';
  renderCalendar();
  // 今日の10問が終わったら広告（プレミアムはスキップ）。
  // ボーナス演出（＋N点の飛び出し）が出た時は、それが読み終わってから広告を出す
  setTimeout(()=>showAd(), flashShown?3300:1200);
}
function renderCalendar(){
  const done=new Set(JSON.parse(localStorage.getItem('mlb_daily_done')||'[]'));
  const today=todayStr();
  const title=calYear+'年'+(calMonth+1)+'月';
  const tid=$('cal-title'),tid2=$('cal-title2');
  if(tid) tid.textContent=title;
  if(tid2) tid2.textContent=title;
  const dows=['日','月','火','水','木','金','土'];
  const firstDay=new Date(calYear,calMonth,1).getDay();
  const daysInMonth=new Date(calYear,calMonth+1,0).getDate();
  let html=dows.map(d=>`<div class="cal-dow">${d}</div>`).join('');
  for(let i=0;i<firstDay;i++) html+=`<div class="cal-day empty"></div>`;
  for(let d=1;d<=daysInMonth;d++){
    const ds=calYear+'-'+String(calMonth+1).padStart(2,'0')+'-'+String(d).padStart(2,'0');
    const isToday=ds===today;
    const isDone=done.has(ds);
    html+=`<div class="cal-day${isToday?' today':''}">${d}${isDone?'<span class="cal-heart"><svg viewBox="0 0 24 24" style="width:70%;height:70%;fill:rgba(244,160,181,0.45);stroke:none"><path d="M20.84 4.61a5.5 5.5 0 00-7.78 0L12 5.67l-1.06-1.06a5.5 5.5 0 00-7.78 7.78l1.06 1.06L12 21.23l7.78-7.78 1.06-1.06a5.5 5.5 0 000-7.78z"/></svg></span>':''}</div>`;
  }
  // 常に6行（42マス）になるよう末尾を空白で埋める
  const totalCells=firstDay+daysInMonth;
  const remainder=42-totalCells;
  for(let i=0;i<remainder;i++) html+=`<div class="cal-day empty"></div>`;
  const g=$('cal-grid'),g2=$('cal-grid2');
  if(g) g.innerHTML=html;
  if(g2) g2.innerHTML=html;
}
function calPrev(){calMonth--;if(calMonth<0){calMonth=11;calYear--;}renderCalendar();}
function calNext(){calMonth++;if(calMonth>11){calMonth=0;calYear++;}renderCalendar();}

// ── Review List ───────────────────────────────────────────────────────────────
function addToReview(type,idx){
  let review=JSON.parse(localStorage.getItem('mlb_review')||'[]');
  const exists=review.some(r=>r.type===type&&r.idx===idx);
  if(!exists){review.unshift({type,idx});localStorage.setItem('mlb_review',JSON.stringify(review));}
}
function removeFromReview(type,idx){
  let review=JSON.parse(localStorage.getItem('mlb_review')||'[]');
  review=review.filter(r=>!(r.type===type&&r.idx===idx));
  localStorage.setItem('mlb_review',JSON.stringify(review));
}
function askClearReview(){ const c=$('review-clear-confirm'); if(c) c.classList.add('show'); }
function cancelClearReview(){ const c=$('review-clear-confirm'); if(c) c.classList.remove('show'); }
function clearReview(){
  localStorage.setItem('mlb_review','[]');
  const c=$('review-clear-confirm'); if(c) c.classList.remove('show');
  renderReviewList();
}
function renderReviewList(){
  const review=JSON.parse(localStorage.getItem('mlb_review')||'[]');
  const list=$('review-list');
  if(!review.length){
    list.innerHTML=`<div class="empty"><img src="images/よろこび.png" alt=""><p>間違えた問題はありません</p></div>`;
    return;
  }
  const frag=document.createDocumentFragment();
  review.forEach(r=>{
    const arr=r.type==='ej'?EJ:JE;
    if(r.idx>=arr.length) return;
    const entry=arr[r.idx];
    const head=r.type==='ej'?entry.en:entry.ja;
    const sub=r.type==='ej'?entry.ja:entry.en;
    const item=document.createElement('div');
    item.className='review-item';
    item.innerHTML=`
      <div class="entry-text" style="flex:1;min-width:0">
        <div class="entry-head">${esc(head)}</div>
        <div class="entry-sub">${esc(formatDef(sub))}</div>
      </div>
      <button class="review-del" onclick="event.stopPropagation();doRemoveReview('${r.type}',${r.idx})">✕</button>`;
    item.addEventListener('click',()=>openXRef(r.idx,r.type));
    frag.appendChild(item);
  });
  list.innerHTML='';list.appendChild(frag);
}
function doRemoveReview(type,idx){
  removeFromReview(type,idx);
  renderReviewList();
}

// ── Scroll position persistence ───────────────────────────────────────────────
let _headerCache=[];  // [{letter, top}] — キャッシュ済みオフセット
function cacheHeaders(){
  const list=$('dict-list');
  if(!list||!list.clientHeight) return; // 非表示中は計測しない（top=0の不正キャッシュを防ぐ）
  const lr=list.getBoundingClientRect();
  _headerCache=Array.from(list.querySelectorAll('[data-letter]')).map(el=>({
    letter:el.getAttribute('data-letter'),
    top:el.getBoundingClientRect().top-lr.top+list.scrollTop
  }));
}
function saveScrollPos(){
  const list=$('dict-list');
  if(dictFav) return; // お気に入り表示中は辞書位置を上書きしない
  if(!list||!_headerCache.length) return;
  const st=list.scrollTop+4;
  let cur=null;
  for(let i=0;i<_headerCache.length;i++){
    if(_headerCache[i].top<=st) cur=_headerCache[i].letter;
    else break;
  }
  if(cur) localStorage.setItem('mlb_pos', dictMode+'|'+cur);
}
let _scrollTimer=null;

// ── State save/restore ────────────────────────────────────────────────────────
// 保存対象: タブのみ（辞書モードは常にEJ、検索クエリは保存しない）
function saveState(){
  localStorage.setItem('mlb_state',JSON.stringify({tab:tab}));
}

// ── Init ──────────────────────────────────────────────────────────────────────
const _st=JSON.parse(localStorage.getItem('mlb_state')||'null');
const _savedPos=localStorage.getItem('mlb_pos');
// タブ復元は下のrequestAnimationFrame内でgoTab()に一本化（ナビ非表示タブでも正しく描画）

// 初期描画（スプラッシュが覆っている間に実施）。setTimeoutでbfcache/背景タブでも確実に実行
setTimeout(()=>{
  renderDictList();
  renderIdxBar();
  renderHome(); // ホーム（連続ログイン・今日の一語・豆知識）を初期化
  // 最後に使った画面を復元（初期表示はホーム。home以外ならgoTabで切替）
  if(_st && _st.tab && _st.tab!=='home'){
    try{ goTab(_st.tab); }catch(e){}
  }
  // EJモードのスクロール位置のみ復元
  const _savedParts=_savedPos?_savedPos.split('|'):null;
  if(_savedParts && _savedParts[0]==='ej' && _savedParts[1] && !sq){
    setTimeout(()=>scrollToLetter(_savedParts[1]),80);
  }
  // 初回のみ: 性別・年代・ニックネームのオンボーディングを表示
  maybeShowOnboard();
},0);

// タブ切替・スクロール時にstate保存
$('dict-list').addEventListener('scroll',()=>{
  clearTimeout(_scrollTimer);
  _scrollTimer=setTimeout(()=>{saveScrollPos();saveState();},200);
});
const _origGoTab=goTab;
goTab=function(t){_origGoTab(t);saveState();};
</script>
</body>
</html>
"""

html = (HTML
    .replace('__EJ_DATA__', EJ_JSON)
    .replace('__JE_DATA__', JE_JSON)
    .replace('__TRIVIA_DATA__', TRIVIA_JSON)
    .replace('__EJ_COUNT__', str(EJ_COUNT))
    .replace('__JE_COUNT__', str(JE_COUNT)))

out = r'C:\Users\s3104\OneDrive\ドキュメント\Claude\MajorLeagueBaseballBilingualDictionary\index.html'
with open(out, 'w', encoding='utf-8') as f:
    f.write(html)

print(f"Done!  {len(html)//1024} KB  →  {out}", file=sys.stderr)
